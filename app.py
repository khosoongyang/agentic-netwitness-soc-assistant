"""
SOC Platform v4
───────────────
• Auto-fetches incidents every 30 seconds once token is verified
• Dashboard tab: severity bars, status bars, latest incidents, assignee chart
• Cleaner sidebar with connection status + refresh countdown
• Incidents tab with severity filter + auto-refresh indicator
• Chat stub (wire LangChain inside chat_respond)
• ChromaDB semantic search & sync
"""

# ── Streamlit Community Cloud: ChromaDB needs sqlite3 >= 3.35 ──────────────────
# Debian images on Streamlit Cloud ship a system sqlite3 older than 3.35, which
# chromadb rejects ("unsupported version of sqlite3"). pysqlite3-binary bundles a
# newer sqlite; swap it in for the stdlib module BEFORE anything imports chromadb.
# No-op locally (pysqlite3-binary isn't installed on the Windows venv, and that
# platform's own sqlite3 is already new enough) — this block is deploy-only.
try:
    __import__("pysqlite3")
    import sys as _sys
    _sys.modules["sqlite3"] = _sys.modules.pop("pysqlite3")
except Exception:
    pass

import re
import streamlit as st
import streamlit.components.v1 as components
import requests
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import time
import os
import base64
from datetime import datetime, timedelta
from typing import Optional
from collections import Counter
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

try:
    from dotenv import load_dotenv, set_key, find_dotenv
    DOTENV_OK = True
except ImportError:
    DOTENV_OK = False

try:
    import chromadb
    from chromadb.utils import embedding_functions as _chroma_embedding_functions
    CHROMA_OK = True
except ImportError:
    CHROMA_OK = False


def _openai_ef():
    """Shared OpenAI embedding function for every ChromaDB collection this
    app creates. Replaces chromadb's bundled local ONNX MiniLM default."""
    return _chroma_embedding_functions.OpenAIEmbeddingFunction(
        api_key=os.environ.get("OPENAI_API_KEY", ""),
        model_name="text-embedding-3-small",
    )

# ── Streamlit Community Cloud: secrets → environment ──────────────────────────
# Locally the app reads credentials from .env (via python-dotenv). On Streamlit
# Cloud there is no .env; secrets are entered in the dashboard and exposed as
# st.secrets. Copy any scalar secrets into os.environ (without clobbering a value
# already set) so every existing os.environ.get(...) call site keeps working
# unchanged. Guarded: a silent no-op locally when no secrets are configured.
def _bridge_cloud_secrets() -> None:
    try:
        for _k, _v in st.secrets.items():
            if isinstance(_v, (str, int, float, bool)) and _k not in os.environ:
                os.environ[_k] = str(_v)
    except Exception:
        pass

_bridge_cloud_secrets()

# ── SOC Triage Agent (LangChain) ──────────────────────────────────────────────
# Streamlit re-executes this script on every rerun but NEVER re-imports
# modules already in sys.modules — edits to the agent files were invisible
# until a full process restart. This shim reloads them when their file
# mtime changes, so agent upgrades apply on the next page refresh.
def _maybe_reload_agent_modules():
    import importlib
    import sys as _sys
    watched = {
        "soc_triage_agent.soc_triage_agent":
            Path(__file__).parent / "soc_triage_agent" / "soc_triage_agent.py",
        "soc_workflow": Path(__file__).parent / "soc_workflow.py",
    }
    try:
        for mod_name, path in watched.items():
            m = _sys.modules.get(mod_name)
            if m is None:
                continue
            mtime = path.stat().st_mtime
            # Stamp the module object itself (survives across sessions) —
            # session-scoped bookkeeping missed changes made BEFORE a new
            # browser session started on an old process.
            loaded = getattr(m, "__loaded_mtime__", None)
            if loaded is None or mtime > loaded:
                for name in (mod_name, "soc_triage_agent"):
                    mm = _sys.modules.get(name)
                    if mm is not None:
                        importlib.reload(mm)
                _sys.modules[mod_name].__loaded_mtime__ = mtime
    except Exception:
        pass

_maybe_reload_agent_modules()

from soc_triage_agent import (CiscoLLMConfig, soc_triage_chat_respond,
                              _TRIAGE_TRIGGER, render_triage_trace,
                              format_ticket_display)
from workflow_state_store import (db_connect, db_init, get_state as wss_get_state,
                                  save_triage_result as wss_save_triage_result,
                                  WorkflowAlreadyRunningError)

# ── Multi-agent workflow (triage → investigation → reporting) ────────────────
try:
    from soc_workflow import (
        run_until_triage_approval as wf_run_until_triage_approval,
        generate_triage_ai_summary as wf_generate_triage_ai_summary,
        render_triage_thinking_plain as wf_render_triage_thinking_plain,
        needs_investigation       as wf_needs_investigation,
        handoff_to_investigation  as wf_handoff_to_investigation,
        run_investigation         as wf_run_investigation,
        handoff_to_reporting      as wf_handoff_to_reporting,
        run_reporting             as wf_run_reporting,
    )
    WORKFLOW_OK = True
except Exception:
    WORKFLOW_OK = False


# ── Background workflow engine ───────────────────────────────────────────────
# Streamlit interrupts the running script at its next UI call whenever the
# user interacts (clicking "View" on the agent board, sending a message…).
# Running investigation+reporting inline therefore died mid-run on any click.
# They now run in a daemon worker thread that survives all interactions; the
# UI polls the shared run record and renders its state live.

_ANSI_STRIP = re.compile(r"\x1b\[[0-9;]*m")


@st.cache_resource
def _workflow_store() -> dict:
    """Process-global store for the active background workflow run."""
    return {"run": None}


def _resolve_full_incident(sel_id: str, fallback_inc: dict) -> tuple[dict, bool]:
    """Prefer the in-memory, live-fetched incident (st.session_state.incidents
    — populated by the existing NetWitness alert-retrieval cycle and includes
    full per-alert data) over the SQLite-cached raw_json, which
    db_upsert_incidents deliberately strips alerts/journalEntries from before
    storage (alerts are refetched live and would otherwise make that DB grow
    unbounded). Returns (incident, is_full)."""
    for i in st.session_state.get("incidents") or []:
        if str(i.get("id") or i.get("incidentId") or "") == str(sel_id):
            return i, True
    return fallback_inc, False


def _board_touch(agent: str, *, status: str | None = None, think: str | None = None,
                 output: str | None = None, progress: int | None = None) -> None:
    """Page-independent state update for st.session_state.agent_board — a
    minimal version of the Agent Board page's local _board_set() closure
    (which also live-refreshes that page's on-screen slots; those slots
    don't exist on other pages, so this only updates the shared state that
    every page reads)."""
    panel = st.session_state.agent_board[agent]
    if status is not None:
        panel["status"] = status
        if status in ("done", "cached", "skipped"):
            panel["progress"] = 100
    if progress is not None:
        panel["progress"] = max(0, min(100, int(progress)))
    if think is not None:
        panel["thinking"].append(f"[{datetime.now().strftime('%H:%M:%S')}] {think}")
        panel["thinking"] = panel["thinking"][-60:]
    if output is not None:
        panel["output"] = output
    panel["updated"] = datetime.now().strftime("%H:%M:%S")


def _run_triage_workflow_with_ui(incident: dict, *, allow_retry: bool = False) -> dict | None:
    """Single UI-facing entry point for Parsing+Triage — used by BOTH the
    Start Process button and the chat trigger, so there is exactly one
    implementation of 'run Parsing+Triage', not two independently-sequenced
    ones. Shows live per-stage progress via st.status(). Returns the
    workflow ctx dict, or None if a run was already active and the caller
    didn't request a retry (the analyst sees the existing result instead)."""
    inc_id = str(incident.get("id") or incident.get("incidentId") or "")

    existing = wss_get_state(inc_id)
    if (existing and existing.get("workflow_status") in ("Processing", "Awaiting Approval")
            and not allow_retry):
        st.info(f"Triage is already **{existing['workflow_status']}** for this "
               f"incident (run {existing.get('run_id')}). Showing the existing "
               f"result — use Retry to start a fresh run.")
        return None

    ALL_PHASES = ["Parsing and Normalisation", "IOC — Availability",
                 "IOC — Confidentiality", "IOC — Integrity",
                 "Risk Rating", "SOC Classification"]
    PHASE_DESC = {
        "Parsing and Normalisation": "Extracting and normalising alert fields",
        "IOC — Availability": "Scanning availability indicators — CPU spikes, reboots, buffer saturation",
        "IOC — Confidentiality": "Scanning confidentiality indicators — exfil, geo anomalies, permission changes",
        "IOC — Integrity": "Scanning integrity indicators — unknown binaries, hash anomalies, processes",
        "Risk Rating": "Calculating risk across initiation, occurrence & adverse impact dimensions",
        "SOC Classification": "Classifying incident severity and generating recommended actions",
    }

    st.session_state.agent_board["triage"].update({"thinking": [], "output": ""})
    _board_touch("triage", status="running", think="Workflow started")

    with st.status("Initialising workflow…", expanded=True) as wf_status:
        _prog = {"done": 0}

        def on_progress(event, label, text=""):
            key = next((p for p in ALL_PHASES if label == p or label in p or p in label), label)
            if event == "phase_start":
                desc = PHASE_DESC.get(key, label)
                wf_status.update(label=f"{desc}…", expanded=True)
                _board_touch("triage", think=f"▶ {desc}")
            elif event == "phase_complete":
                result = f" — {text}" if text else ""
                st.write(f"**{key}**{result}")
                _prog["done"] += 1
                _board_touch("triage", think=f"{key}{result}",
                            progress=round(_prog["done"] / max(1, len(ALL_PHASES)) * 100))
            elif event == "phase_error":
                st.write(f"**{key}** — failed: {text}")

        try:
            result = wf_run_until_triage_approval(
                incident, progress_fn=on_progress, allow_retry=allow_retry)
        except WorkflowAlreadyRunningError as exc:
            wf_status.update(label="Already running", state="complete", expanded=False)
            st.info(f"Triage is already {exc.state.get('workflow_status')} for this "
                   f"incident — showing the existing result.")
            return None
        except Exception as exc:
            wf_status.update(label="Workflow failed", state="error", expanded=True)
            _board_touch("triage", status="failed", think=str(exc)[:150], output=f"Error: {exc}")
            st.error(f"Workflow failed: {exc}")
            return None

        _err = result["errors"].get("parsing") or result["errors"].get("triage")
        if _err:
            wf_status.update(label="Workflow failed", state="error", expanded=True)
            _board_touch("triage", status="failed", think=str(_err)[:150], output=f"Error: {_err}")
            st.error(f"Workflow failed: {_err}")
            return None

        wf_status.update(label="Triage complete — awaiting your approval",
                         state="complete", expanded=False)
        _tri = result["triage"]
        # Board panel stays a step-by-step progress log (phase name +
        # cached/timing marker) — the connected-reasoning narrative lives
        # only on the case-detail page's Thinking Process card, sourced
        # live from triage_result_json (see the "Triage" stage branch
        # above), not duplicated here.
        _board_touch("triage",
                     status=("cached" if _tri.get("cached") else "done"),
                     think="Triage complete — awaiting SOC analyst approval",
                     output=format_ticket_display(_tri["ticket"]))

    st.session_state.triage_in_flight = None
    return result


def _workflow_worker(run: dict, tri: dict, incident: dict) -> None:
    """Investigation + reporting stages, off the Streamlit script thread.
    NO st.* calls in here — the UI polls `run` and renders its state."""
    import soc_workflow as _wfm

    panels = run["panels"]
    wf_md  = run["wf_md"]

    def bset(agent, status=None, think=None, output=None, progress=None):
        p = panels[agent]
        if status is not None:
            p["status"] = status
            if status in ("done", "cached", "skipped"):
                p["progress"] = 100
        if progress is not None:
            p["progress"] = max(0, min(100, int(progress)))
        if think is not None:
            clean = _ANSI_STRIP.sub("", str(think))
            p["thinking"].append(
                f"[{datetime.now().strftime('%H:%M:%S')}] {clean}")
            p["thinking"] = p["thinking"][-60:]
        if output is not None:
            p["output"] = output
        p["updated"] = datetime.now().strftime("%H:%M:%S")

    inc_id = run["incident_id"]
    title  = run["title"]
    cls    = run["cls"]
    unc    = run["unc"]
    ticket = tri.get("ticket") or {}
    inv = None
    rep: dict = {}
    try:
        # ── Stage 2: Investigation ─────────────────────────────────────────
        if run["investigate"]:
            bset("investigation", status="running", think="Investigation started",
                 progress=10)

            def _fb_event(event: str, detail: str) -> None:
                # Board choreography for the feedback loop: the work visibly
                # returns to the triage card for the deep-dive, then hands
                # back to investigation for the second pass.
                if event == "gaps_detected":
                    bset("investigation", think=f"{detail}")
                elif event == "triage_deep_dive_start":
                    bset("triage", status="running", think=f"{detail}")
                elif event == "triage_deep_dive_done":
                    bset("triage", status="done", think=f"{detail}")
                    bset("investigation", think=f"{detail}")
                elif event == "second_pass_start":
                    bset("investigation", status="running", think=f"{detail}")
                elif event == "supplement_error":
                    bset("triage", status="done",
                         think=f"deep-dive failed: {detail}")
                else:
                    bset("investigation", think=detail)

            inv = _wfm.investigate_with_feedback(
                tri, incident, inc_id,
                line_cb=lambda ln: bset("investigation", think=ln)
                if ln.strip() else None,
                feedback_cb=_fb_event)
            _fbmeta = inv.get("feedback_loop") or {}
            if _fbmeta.get("triggered"):
                # Honest reporting — a crashed deep-dive or failed second
                # pass must never be presented as a successful loop.
                gap_ids = ", ".join(g.split(":")[0]
                                    for g in (_fbmeta.get("gaps") or []))
                if _fbmeta.get("supplement_error"):
                    wf_md.append(
                        f"- Feedback loop: gaps detected (`{gap_ids}`) but "
                        f"the triage deep-dive failed — "
                        f"`{str(_fbmeta['supplement_error'])[:120]}` — "
                        f"pass-1 findings kept")
                elif _fbmeta.get("second_pass_failed"):
                    wf_md.append(
                        f"- Feedback loop: triage deep-dive answered "
                        f"{_fbmeta.get('gaps_answered', 0)} gap(s) but the "
                        f"re-investigation failed — pass-1 findings kept")
                else:
                    _extra = ""
                    if _fbmeta.get("playbook_redirect"):
                        _extra += (" · playbook redirected: "
                                   + ", ".join(_fbmeta["playbook_redirect"]
                                               .values()))
                    if _fbmeta.get("suggested_classification"):
                        _extra += (f" · deep-dive suggests classification "
                                   f"**{_fbmeta['suggested_classification']}**"
                                   f" (analyst to review)")
                    wf_md.append(
                        f"- Feedback loop (pass {_fbmeta.get('passes', 1)}): "
                        f"gaps `{gap_ids}` → triage deep-dive answered "
                        f"**{_fbmeta.get('gaps_answered', 0)}** → "
                        f"investigation re-ran with the supplement{_extra}")
            if inv.get("status") == "failed":
                try:
                    (SOC_DB_DIR / "last_investigation_error.json").write_text(
                        _json.dumps(inv, indent=2, default=str), encoding="utf-8")
                except Exception:
                    pass
                _ie = str(inv.get("error") or "")
                _cold = ("503" in _ie or "unavailable" in _ie.lower())
                wf_md.append(f"- Investigation: failed — "
                             f"`{_ie[:150]}` "
                             + ("**(the LLM endpoint was asleep — wait a few "
                                "minutes for it to boot, then re-run)** "
                                if _cold else "")
                             + f"(full details: `soc_db/last_investigation_error.json`)")
                bset("investigation", status="failed",
                     think=f"{str(inv.get('error') or '')[:150]}",
                     output=f"Investigation failed:\n\n```\n"
                            f"{str(inv.get('error') or '')[:800]}\n```")
                inv = None
            else:
                bset("investigation", status="done",
                     think=f"Complete — {inv.get('incident_folder')}",
                     output=inv.get("narrative_report") or inv.get("summary")
                            or "Investigation completed.")
                wf_md.append(f"- Investigation: {inv.get('status')} — "
                             f"folder `{inv.get('incident_folder')}`")
                try:
                    rec = _wfm.build_post_investigation_record(
                        inv, ticket, title, run_stamp=run["run_id"])
                    for _attempt in (1, 2, 3):
                        try:
                            _wfm.pipeline_insert("post_investigation", rec)
                            break
                        except Exception:
                            if _attempt == 3:
                                raise
                            time.sleep(0.8)   # ride out a brief sqlite lock
                    run["chroma_queue"].append(("post_investigation", rec))
                    wf_md.append("- Findings recorded — **Pipeline DB** "
                                 "tab → *Post-Investigation*")
                except Exception as exc:
                    # Never swallow silently — a missing findings record looks
                    # like "no output" to the analyst.
                    bset("investigation",
                         think=f"findings record insert failed: {str(exc)[:120]}")
                    wf_md.append(f"- Findings record insert failed: "
                                 f"`{str(exc)[:150]}`")
        else:
            bset("investigation", status="skipped",
                 think=f"Skipped — classification {cls} below routing threshold",
                 output=f"Investigation skipped: classification **{cls}** is "
                        "below the routing threshold (critical/high/medium).")
            wf_md.append(f"- Investigation: skipped "
                         f"(classification {cls} below routing threshold)")

        # ── HITL gate: pause before reporting until the analyst approves ────
        # The Event is pre-set in non-manual mode (see spawn site), so this is a
        # no-op for the normal auto-chain and only blocks when manual review is
        # on. gate_decision (set by the Reject button before .set()) is how a
        # threading.Event — which can't carry "which way" on its own — tells
        # this thread the analyst said no instead of yes.
        gate = run.get("gate_report")
        _gate_rejected = False
        if gate is not None and not gate.is_set():
            run["awaiting"] = "report"
            bset("reporting", status="queued", progress=0,
                 think="Investigation complete — awaiting analyst approval to "
                       "generate the report…")
            gate.wait()
            run["awaiting"] = None
            if run.get("gate_decision") == "rejected":
                _gate_rejected = True
                _reason = run.get("gate_reason") or "no reason given"
                bset("reporting", status="rejected",
                     think=f"Rejected by analyst — {_reason}",
                     output=f"**Rejected by analyst"
                            f"{(' (' + run['analyst'] + ')') if run.get('analyst') else ''}"
                            f".** Reason: {_reason}\n\nNo report was generated. "
                            "Re-run triage on this incident to try again.")
                wf_md.append(f"- Reporting: rejected by analyst — {_reason}")
            else:
                bset("reporting", think="Approved by analyst — starting report")

        # ── Stage 3: Reporting ─────────────────────────────────────────────
        if not _gate_rejected:
            bset("reporting", status="running", think="Reporting started", progress=10)
            try:
                tid = _wfm.handoff_to_reporting(tri, incident, inv)
                bset("reporting", think="Triage + investigation context handed over",
                     progress=35)
                rep = _wfm.run_reporting(
                    tid, run_stamp=run["run_id"],
                    line_cb=lambda ln: bset("reporting", think=ln)
                    if ln.strip() else None)
            except Exception as exc:
                rep = {"status": "failed", "error": str(exc)}
            if rep.get("status") == "failed":
                try:
                    (SOC_DB_DIR / "last_reporting_error.json").write_text(
                        _json.dumps(rep, indent=2, default=str), encoding="utf-8")
                except Exception:
                    pass
                bset("reporting", status="failed",
                     think=f"{str(rep.get('error') or '')[:150]}",
                     output=f"Reporting failed:\n\n```\n"
                            f"{str(rep.get('error') or '')[:800]}\n```")
                wf_md.append(f"- Reporting: failed — "
                             f"`{str(rep.get('error') or '')[:150]}` "
                             f"(full details: `soc_db/last_reporting_error.json`)")
            else:
                rec = {
                    "id": f"final_{unc}@{run['run_id']}",
                    "incident_id": inc_id, "ticket_unc": unc,
                    "title": f"[FINAL] {title}", "severity": cls,
                    "summary": str(rep.get("summary")
                                   or "Incident report generated.")[:500],
                    "approved_by": run.get("analyst") or "—",
                    "report": {k: v for k, v in rep.items()
                               if k not in ("subprocess", "orchestrator_subprocess")}}
                _wfm.pipeline_insert("finalized_report", rec)
                run["chroma_queue"].append(("finalized_report", rec))
                exports = rep.get("document_exports") or {}
                # Structured data + export paths for the Generated Files panel
                # (see _render_generated_files) — kept on the panel dict itself
                # so it survives the disk fallback (last_workflow_result.json).
                panels["reporting"]["exports"] = exports
                panels["reporting"]["reporting_data"] = rec["report"]
                rep_out = [
                    f"**Status:** {rep.get('status')} "
                    f"(mode: {rep.get('reporting_mode', 'standard')})",
                    f"**LLM:** {rep.get('llm_status') or '—'}",
                    "",
                    f"**Summary:** {rep.get('summary') or '—'}",
                    "",
                ]
                _generated_fmts = [f.upper() for f in ("docx", "pdf") if exports.get(f)]
                if _generated_fmts:
                    rep_out.append(f"- Documents generated: {' · '.join(_generated_fmts)} "
                                   "— see **Generated Files** below")
                else:
                    rep_out.append("- No documents were exported (see errors below)")
                rep_out += ["", "Full report suite: **Pipeline DB** tab "
                                "→ *Finalized Report*."]
                bset("reporting", status="done", think="Report finalised",
                     output="\n".join(rep_out))
                wf_md.append(f"- Reporting: {rep.get('status')} "
                             f"(mode: {rep.get('reporting_mode', 'standard')})")
                for fmt, icon in (("docx", "▤"), ("pdf", "▤")):
                    if exports.get(fmt):
                        wf_md.append(f"- {icon} {fmt.upper()} report generated")
                    else:
                        err = str(exports.get(f"{fmt}_error") or exports.get("error")
                                  or "no file produced")[:150]
                        wf_md.append(f"- {icon} {fmt.upper()} report: "
                                     f"export failed — `{err}`")
                wf_md.append("")
                wf_md.append("Download the Word/PDF report from the "
                             "**Pipeline DB** tab → *Finalized Report*.")
    except Exception as exc:
        wf_md.append(f"- Workflow worker crashed: `{str(exc)[:200]}`")
        for ag in ("investigation", "reporting"):
            if panels[ag]["status"] in ("running", "queued"):
                bset(ag, status="failed",
                     think=f"worker crash: {str(exc)[:150]}")
    finally:
        # Audit trail row — one NEW record per workflow execution.
        try:
            dur = int(time.time() - run["started_ts"])
            stages = {
                "triage": "cached" if run.get("cached_triage") else "fresh",
                "investigation": panels["investigation"]["status"],
                "reporting": panels["reporting"]["status"],
            }
            _wfm.pipeline_insert("workflow_runs", {
                "id": f"run_{run['run_id']}_{inc_id[:20]}",
                "incident_id": inc_id,
                "title": f"Run {run['started_hms']} — {title}",
                "severity": cls,
                "summary": " · ".join(f"{k}: {v}" for k, v in stages.items())
                           + f" · ticket {unc} · {dur}s"
                           + (f" · analyst {run['analyst']}"
                              if run.get("analyst") else ""),
                "stages": stages, "ticket_unc": unc,
                "duration_seconds": dur,
                "analyst": run.get("analyst") or "—",
                "gate_decision": run.get("gate_decision") or "",
                "gate_reason": run.get("gate_reason") or ""})
        except Exception:
            pass
        run["finished_at"] = time.time()
        run["done"] = True
        # Persist the finished run so ANY session can surface the results on
        # its next interaction — even if the browser poll loop died mid-run
        # (exception in a rerun, page refresh, process restart).
        try:
            (SOC_DB_DIR / "last_workflow_result.json").write_text(
                _json.dumps(run, default=str, indent=1), encoding="utf-8")
        except Exception:
            pass

# ── Destructive-action confirmation ──────────────────────────────────────────
# Every delete in this app used to fire on a single click with no "are you
# sure" step (Clear Stage, per-record delete, credential clear). This is a
# two-click, inline confirmation — no st.dialog dependency, so it works the
# same whether it's called once (a global button) or many times inside a
# loop with a unique per-row key (the per-record delete list).
def _confirm_action(action_key: str, trigger_label: str, warn_body: str,
                    confirm_label: str = "Yes, delete", *,
                    use_container_width: bool = True) -> bool:
    """Renders its own trigger button. First click arms it (shows the warning
    + Confirm/Cancel); returns True only on the run where Confirm is clicked."""
    armed_key = f"_confirm_armed_{action_key}"
    if not st.session_state.get(armed_key):
        if st.button(trigger_label, key=f"_confirm_trig_{action_key}",
                     use_container_width=use_container_width):
            st.session_state[armed_key] = True
            st.rerun()
        return False
    st.warning(warn_body)
    cc1, cc2 = st.columns(2)
    do_confirm = cc1.button(f"⚠ {confirm_label}", key=f"_confirm_yes_{action_key}",
                            type="primary", use_container_width=True)
    if cc2.button("Cancel", key=f"_confirm_no_{action_key}", use_container_width=True):
        st.session_state[armed_key] = False
        st.rerun()
    if do_confirm:
        st.session_state[armed_key] = False
        return True
    return False


def _reject_with_reason(action_key: str, trigger_label: str = "Reject") -> str | None:
    """Two-step reject control requiring a reason — the counterpart to
    Approve that the HITL gates used to be missing entirely. Mirrors
    _confirm_action's arm/cancel pattern but collects required free text
    instead of a yes/no, since a rejection needs an audit-trail reason.
    Returns the reason string on the run the analyst submits it; None
    otherwise (including while still armed/mid-entry)."""
    armed_key = f"_reject_armed_{action_key}"
    if not st.session_state.get(armed_key):
        if st.button(trigger_label, key=f"_reject_trig_{action_key}",
                     use_container_width=True):
            st.session_state[armed_key] = True
            st.rerun()
        return None
    reason = st.text_area(
        "Reason for rejection (required)",
        key=f"_reject_reason_{action_key}",
        placeholder="e.g. Triage misclassified this — it's a known false positive.")
    rc1, rc2 = st.columns(2)
    submit = rc1.button("Confirm rejection", key=f"_reject_go_{action_key}",
                        type="primary", use_container_width=True)
    if rc2.button("Cancel", key=f"_reject_cancel_{action_key}", use_container_width=True):
        st.session_state[armed_key] = False
        st.rerun()
    if submit:
        if not reason.strip():
            st.error("A reason is required before rejecting.")
            return None
        st.session_state[armed_key] = False
        return reason.strip()
    return None


def _normalise_llm_url(url: str) -> str:
    """
    Ensure the base URL is safe for ChatOpenAI.
    - Strips trailing slashes
    - Removes any embedded model path (e.g. /models/some-model/...)
    - Guarantees the URL ends with /v1
    """
    from urllib.parse import urlparse
    url = url.strip().rstrip("/")
    parsed = urlparse(url)
    # If someone pasted a model-specific URL, reduce it to just host + /v1
    if "/models/" in parsed.path:
        url = f"{parsed.scheme}://{parsed.netloc}/v1"
    elif not parsed.path.endswith("/v1"):
        url = url + "/v1"
    return url

def get_cisco_cfg() -> CiscoLLMConfig:
    """
    Build CiscoLLMConfig from session state at call time so that
    credentials entered in the sidebar are always picked up. Falls back to
    the OPENAI_* env vars (already configured for the reporting agent) when
    no custom endpoint has been entered in the sidebar.
    """
    raw_url = st.session_state.get("cisco_url", "").strip()
    url     = (_normalise_llm_url(raw_url) if raw_url
               else os.environ.get("OPENAI_BASE_URL", "").strip() or "https://api.openai.com/v1")
    return CiscoLLMConfig(
        base_url    = url,
        api_key     = (st.session_state.get("cisco_key", "").strip()
                       or os.environ.get("OPENAI_API_KEY", "").strip() or "changeme"),
        model       = (st.session_state.get("cisco_model", "").strip()
                       or os.environ.get("OPENAI_MODEL", "").strip() or "gpt-4o-mini"),
        temperature = 0.0,
        # 1024 was starving the reasoning model: it spent the whole budget on
        # chain-of-thought and got truncated before emitting the final JSON,
        # which surfaced as "0 IOCs matched" on every triage run.
        max_tokens  = 3072,
        timeout     = 300,
    )

# ══════════════════════════════════════════════════════════════════════════════
# .ENV  — load persisted credentials on every startup
# ══════════════════════════════════════════════════════════════════════════════
ENV_FILE = Path(__file__).parent / ".env"

def _clean(val: str) -> str:
    """Strip whitespace and stray quotes that dotenv sometimes leaves."""
    return val.strip().strip("'\"").strip()

def env_load() -> dict:
    """Read credentials from .env."""
    if DOTENV_OK and ENV_FILE.exists():
        load_dotenv(ENV_FILE, override=True)
    return {
        "host":          _clean(os.environ.get("NW_HOST",       "")),
        "username":      _clean(os.environ.get("NW_USERNAME",   "")),
        "password":      _clean(os.environ.get("NW_PASSWORD",   "")),
        "nw_cert_path":  _clean(os.environ.get("NW_CERT_PATH",  "")),
        # Cisco LLM
        "cisco_url":     _clean(os.environ.get("CISCO_LLM_URL",   "")),
        "cisco_key":     _clean(os.environ.get("CISCO_LLM_KEY",   "")),
        "cisco_model":   _clean(os.environ.get("CISCO_LLM_MODEL", "")),
    }

def env_save(host: str, username: str, password: str) -> None:
    """Persist NetWitness credentials to .env."""
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "NW_HOST",     host.strip())
        set_key(str(ENV_FILE), "NW_USERNAME", username.strip())
        # base64-encode password to avoid special char issues
        set_key(str(ENV_FILE), "NW_PASSWORD",
                base64.b64encode(password.strip().encode()).decode("ascii"))
    except Exception:
        pass  # .env locked on Windows — credentials still active this session

def nw_cert_env_save(cert_path: str) -> None:
    """Persist the TLS cert path to .env.
    Silently skips if .env is locked (common on Windows) — cert path
    is still stored in session state and works for the current session.
    """
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "NW_CERT_PATH", cert_path.strip())
    except Exception:
        pass  # .env locked on Windows — session state still holds the path

def nw_cert_env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "NW_CERT_PATH", "")
    except Exception:
        pass

def cisco_env_save(url: str, key: str, model: str) -> None:
    """Persist Cisco LLM credentials to .env."""
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "CISCO_LLM_URL",   url.strip())
        set_key(str(ENV_FILE), "CISCO_LLM_MODEL", model.strip())
        # base64-encode key to avoid special char issues
        set_key(str(ENV_FILE), "CISCO_LLM_KEY",
                base64.b64encode(key.strip().encode()).decode("ascii"))
    except Exception:
        pass

def env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "NW_HOST",     "")
            set_key(str(ENV_FILE), "NW_USERNAME", "")
            set_key(str(ENV_FILE), "NW_PASSWORD", "")
    except Exception:
        pass

def cisco_env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "CISCO_LLM_URL",   "")
            set_key(str(ENV_FILE), "CISCO_LLM_KEY",   "")
            set_key(str(ENV_FILE), "CISCO_LLM_MODEL", "")
    except Exception:
        pass

def nw_login(host: str, username: str, password: str) -> tuple[bool, str, str]:
    """
    Login with username/password → returns (ok, message, access_token).
    POST /rest/api/auth/userpass with form-encoded credentials.
    Auto-retries with verify=False if cert verification fails.
    """
    if not host.strip():
        return False, "Host URL is empty — enter https://192.168.x.x", ""
    host = host.strip().strip("'\"").strip()  # remove any stray quotes
    if not username.strip():
        return False, "Username is empty.", ""
    if not password.strip():
        return False, "Password is empty.", ""

    def _attempt(verify):
        return requests.post(
            f"{host.rstrip('/')}/rest/api/auth/userpass",
            data={"username": username, "password": password},
            headers={
                "Content-Type": "application/x-www-form-urlencoded; charset=ISO-8859-1",
                "Accept":       "application/json;charset=UTF-8",
            },
            timeout=15,
            verify=verify,
        )

    try:
        verify = nw_tls_verify()
        try:
            r = _attempt(verify)
        except requests.exceptions.SSLError:
            # Cert didn't work — clear it and retry without verification
            st.session_state.nw_cert_path = ""
            r = _attempt(False)

        if r.status_code == 200:
            data  = r.json()
            token = data.get("accessToken") or data.get("access_token") or ""
            if token:
                return True, "NetWitness connected", token
            return False, f"Login OK but no token in response: {str(data)[:100]}", ""
        else:
            return False, f"HTTP {r.status_code}: {r.text[:200]}", ""
    except requests.exceptions.ConnectionError as e:
        return False, f"Cannot reach {host} — check VPN/network: {str(e)[:120]}", ""
    except requests.exceptions.Timeout:
        return False, "Login timed out — check GP VPN is connected.", ""
    except Exception as e:
        return False, f"Login error: {e}", ""

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Aegis",
    page_icon=None,
    layout="wide",
    # Nothing calls st.sidebar anymore (top nav bar replaced it) — collapsed
    # defensively so Streamlit doesn't reserve an empty sidebar strip.
    initial_sidebar_state="collapsed",
)

# ══════════════════════════════════════════════════════════════════════════════
# CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@400;500;600;700&display=swap');

:root {
  /* Aegis design system — re-skin of the SOC dashboard */
  --bg:      #060b13;
  --nav:     #07101b;
  --bg1:     #0e1929;   /* card surface */
  --bg2:     #101d30;   /* raised card */
  --card:    #0e1929;
  --card2:   #101d30;
  --border:  #223149;
  --line:    #223149;
  --accent:  #6f7cff;   /* Aegis blue */
  --blue:    #6f7cff;
  --cyan:    #36c5d3;
  --purple:  #a67af4;
  --green:   #43d28c;
  --warn:    #f4bc5f;
  --danger:  #ff6e7c;
  --orange:  #f4bc5f;
  --muted:   #8b9bb2;   /* subdued text */
  --sub:     #8b9bb2;
  --faint:   #61738d;
  --text:    #f3f6fb;
  --r:       14px;
  --mono:    'IBM Plex Sans', 'Segoe UI', sans-serif;
  --sans:    'IBM Plex Sans', 'Segoe UI', sans-serif;

  /* Type scale — sanctioned sizes going forward. The file has decades of
     one-off inline font-size values (0.48rem-2.3rem); this isn't a full
     retrofit, just a floor so new/edited spots stop going under it. */
  --fs-3xs:  0.68rem;   /* smallest allowed — dense mono status lines */
  --fs-2xs:  0.74rem;
  --fs-xs:   0.82rem;
  --fs-sm:   0.9rem;
  --fs-base: 1rem;
  --fs-lg:   1.1rem;
}

/* Standardized global font enforcement across UI elements while preserving icon fonts */
html, body, .stApp, p, div:not([data-testid*="Icon"]), button, input, select, textarea, label, h1, h2, h3, h4, h5, h6 {
  font-family: 'IBM Plex Sans', 'Segoe UI', sans-serif;
}

[data-testid*="Icon"], [class*="material-symbols"], [class*="material-icons"] {
  font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
}

html { font-size: 87.5%; }
html, body, [class*="css"] {
  background: var(--bg);
  color: var(--text);
  line-height: 1.6;
}
/* was 14px — now 1rem so it tracks the html anchor above and scales with prefs */
body, [class*="css"] { font-size: 1rem; }
.main, .stApp {
  background:
    radial-gradient(circle at 64% -20%, #17233b 0, transparent 33%),
    var(--bg);
}
.main { padding-top: 0.5rem; }
/* Streamlit reserves a large title/header gutter by default. Keep the primary
   navigation close to the top app chrome instead of leaving a blank band. */
[data-testid="stMainBlockContainer"],
.main .block-container {
  padding-top: 0.75rem !important;
  padding-left: 1.5rem !important;
  padding-right: 1.5rem !important;
  width: 100% !important;
  max-width: none !important;
}

/* ── Sidebar ── */
section[data-testid="stSidebar"] > div:first-child {
  background: linear-gradient(180deg, #040810 0%, #06101A 100%);
  border-right: 1px solid #0E1E30;
  padding-top: 1rem;
}

/* ── Aegis Sidebar Vertical Menu (NO FRAMES, NO BUTTON LOOK - EXACT SCREENSHOT MATCH) ── */
.sec-label-overview {
  color: #5d6f88 !important;
  text-transform: uppercase;
  letter-spacing: 0.16em !important;
  font-size: 0.62rem !important;
  font-weight: 800 !important;
  margin: 14px 0 8px 6px !important;
}

section[data-testid="stSidebar"] div.stButton,
section[data-testid="stSidebar"] div.stButton > button,
section[data-testid="stSidebar"] button[data-testid="stBaseButton-secondary"],
section[data-testid="stSidebar"] button[data-testid="stBaseButton-primary"] {
  background: transparent !important;
  background-image: none !important;
  border: 0 !important;
  border-style: none !important;
  border-width: 0 !important;
  outline: none !important;
  box-shadow: none !important;
  border-radius: 6px !important;
  color: #94a3b8 !important;
  font-size: 0.88rem !important;
  font-weight: 500 !important;
  padding: 8px 12px !important;
  height: 38px !important;
  min-height: 38px !important;
  text-align: left !important;
  display: flex !important;
  align-items: center !important;
  justify-content: flex-start !important;
  width: 100% !important;
  margin: 2px 0 !important;
  cursor: pointer !important;
}

section[data-testid="stSidebar"] div.stButton > button p,
section[data-testid="stSidebar"] button div,
section[data-testid="stSidebar"] button span {
  font-size: 0.88rem !important;
  font-weight: 500 !important;
  text-align: left !important;
  justify-content: flex-start !important;
  width: 100% !important;
  margin: 0 !important;
}

section[data-testid="stSidebar"] div.stButton > button:hover,
section[data-testid="stSidebar"] button[data-testid="stBaseButton-secondary"]:hover {
  background: #0f1c30 !important;
  color: #f1f5f9 !important;
  border: 0 !important;
  box-shadow: none !important;
  transform: none !important;
}

/* Active item state — entire option selected in a blue background hue, NO BORDER */
section[data-testid="stSidebar"] button[kind="primary"],
section[data-testid="stSidebar"] button[data-testid="stBaseButton-primary"] {
  background: #142542 !important;
  color: #3880ff !important;
  font-weight: 600 !important;
  border: 0 !important;
  border-radius: 6px !important;
  box-shadow: none !important;
}

section[data-testid="stSidebar"] button[kind="primary"] p,
section[data-testid="stSidebar"] button[data-testid="stBaseButton-primary"] p,
section[data-testid="stSidebar"] button[kind="primary"] span,
section[data-testid="stSidebar"] button[data-testid="stBaseButton-primary"] span {
  color: #3880ff !important;
  font-weight: 600 !important;
}

/* ── Headings (Aegis: white, bold, IBM Plex Sans) ── */
h1,h2,h3 {
  font-family: var(--sans);
  font-weight: 700;
  color: var(--text);
  letter-spacing: -0.2px;
  margin-bottom: 0.4rem;
}
h1 { font-size: 1.55rem !important; letter-spacing: -0.4px; }
h2 { font-size: 1.1rem !important; }
h3 { font-size: 0.95rem !important; }

/* ── Metric cards ── */
[data-testid="metric-container"] {
  background: linear-gradient(135deg, #070D18 0%, #0A1628 100%);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 18px 20px;
  box-shadow: 0 2px 16px rgba(0,0,0,0.4);
  transition: all 0.2s;
}
[data-testid="metric-container"]:hover {
  border-color: #1E4060;
  transform: translateY(-2px);
  box-shadow: 0 4px 24px rgba(0,212,255,0.08);
}
[data-testid="metric-container"] label {
  color: var(--muted);
  font-size: 0.72rem;
  font-weight: 600;
  letter-spacing: 0.5px;
  text-transform: uppercase;
  font-family: var(--sans);
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
  color: var(--text);
  font-family: var(--sans);
  font-size: 2rem;
  font-weight: 700;
  letter-spacing: -0.5px;
}

/* ── Main Area Buttons ── */
.main .stButton > button {
  background: linear-gradient(90deg, #052030, #083050);
  color: var(--accent);
  border: 1px solid #0E4A6A;
  border-radius: 8px;
  font-family: var(--sans);
  font-size: 0.78rem;
  font-weight: 500;
  padding: 8px 16px;
  transition: all 0.15s;
}
.stButton > button:hover {
  background: linear-gradient(90deg, #083050, #0E4A6A);
  box-shadow: 0 0 18px rgba(0,212,255,0.2);
  color: #fff;
  border-color: var(--accent);
  transform: translateY(-1px);
}
.stButton > button:active { transform: translateY(0); }

/* ── Inputs ── */
.stTextInput > div > div > input,
.stTextArea textarea {
  background: #060E1A !important;
  border: 1px solid #1A3050 !important;
  color: var(--text) !important;
  border-radius: 8px;
  font-family: var(--sans);
  font-size: 0.88rem;
  padding: 10px 14px !important;
}
.stTextInput > div > div > input:focus,
.stTextArea textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(0,212,255,0.1) !important;
}
.stTextInput label, .stTextArea label, .stSelectbox label {
  color: var(--text) !important;
  font-size: 0.82rem !important;
  font-weight: 500 !important;
  margin-bottom: 4px !important;
}
.stSelectbox > div > div {
  background: #060E1A !important;
  border: 1px solid #1A3050 !important;
  color: var(--text) !important;
  border-radius: 8px;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
  background: transparent;
  border-bottom: 1px solid var(--border);
  gap: 4px;
}
.stTabs [data-baseweb="tab"] {
  font-family: var(--sans);
  font-size: 0.82rem;
  font-weight: 500;
  color: var(--muted);
  border: none;
  padding: 10px 20px;
  background: transparent;
  border-radius: 8px 8px 0 0;
  transition: all 0.15s;
}
.stTabs [data-baseweb="tab"]:hover {
  color: var(--text);
  background: rgba(0,212,255,0.04);
}
.stTabs [aria-selected="true"] {
  color: var(--accent) !important;
  border-bottom: 2px solid var(--accent) !important;
  background: rgba(0,212,255,0.06) !important;
  font-weight: 600 !important;
}

/* ── Cards ── */
.card {
  background: linear-gradient(135deg, #070D18 0%, #08111E 100%);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 16px 20px;
  margin: 6px 0;
  transition: all 0.15s;
}
.card:hover {
  border-color: #1E4060;
  transform: translateX(3px);
  box-shadow: 0 2px 16px rgba(0,0,0,0.3);
}
.card-critical { border-left: 4px solid var(--danger) !important; }
.card-high     { border-left: 4px solid var(--orange) !important; }
.card-medium   { border-left: 4px solid var(--warn)   !important; }
.card-low      { border-left: 4px solid var(--green)  !important; }

/* ── Badges ── */
.badge {
  padding: 3px 10px;
  border-radius: 20px;
  font-size: 0.7rem;
  font-weight: 600;
  display: inline-block;
  font-family: var(--sans);
}
.badge-critical { background:#321b25; color:#ff99a3; border:1px solid #713744; }
.badge-high     { background:#2b221a; color:#f3c679; border:1px solid #684c2a; }
.badge-medium   { background:#241a34; color:#c9a6f7; border:1px solid #5b3f82; }
.badge-low      { background:#122b21; color:#7fe0ac; border:1px solid #2a6146; }
.badge-info     { background:#192743; color:#aeb7ff; border:1px solid #3b4c81; }

/* ── Chat bubbles ── */
.bubble-user {
  background: #080F1A;
  border-left: 3px solid #005C8A;
  padding: 12px 16px;
  border-radius: 0 10px 10px 0;
  margin: 8px 0;
  font-size: 0.9rem;
  line-height: 1.6;
}
.bubble-agent {
  background: #040A12;
  border-left: 3px solid var(--accent);
  padding: 12px 16px;
  border-radius: 0 10px 10px 0;
  margin: 8px 0;
  font-size: 0.9rem;
  line-height: 1.6;
}
.bubble-label {
  font-family: var(--sans);
  font-size: 0.65rem;
  font-weight: 600;
  letter-spacing: 0.5px;
  margin-bottom: 5px;
  color: var(--accent);
  text-transform: uppercase;
}

/* ── Status dots ── */
.dot { display:inline-block; width:8px; height:8px; border-radius:50%; margin-right:7px; vertical-align:middle; }
.dot-green  { background:var(--green);  box-shadow:0 0 8px var(--green);  animation:pulse 2s infinite; }
.dot-red    { background:var(--danger); box-shadow:0 0 6px var(--danger); }
.dot-yellow { background:var(--warn);   box-shadow:0 0 6px var(--warn);   animation:pulse 1s infinite; }
@keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.35} }

/* ── Section labels ── */
.sec-label {
  font-family: var(--sans);
  font-size: 0.7rem;
  font-weight: 600;
  color: #4A7090;
  text-transform: uppercase;
  letter-spacing: 1px;
  margin: 20px 0 10px;
  padding-bottom: 6px;
  border-bottom: 1px solid #0E1E2E;
}

/* ── Stat mini cards ── */
.stat-mini {
  background: #070D18;
  border: 1px solid var(--border);
  border-radius: 10px;
  padding: 14px 16px;
  text-align: center;
  font-family: var(--sans);
}
.stat-mini .val { font-size: 1.5rem; font-weight: 700; color: var(--accent); }
.stat-mini .lbl { font-size: 0.68rem; font-weight: 500; color: var(--muted); margin-top: 3px; text-transform: uppercase; }

/* ── Tooltips / info boxes ── */
.info-box {
  background: #060E1A;
  border: 1px solid #1A3050;
  border-radius: 10px;
  padding: 14px 18px;
  font-size: 0.82rem;
  color: var(--text);
  line-height: 1.6;
  margin: 8px 0;
}
.info-box .title {
  font-weight: 600;
  color: var(--accent);
  margin-bottom: 6px;
  font-size: 0.85rem;
}

/* ── Expanders ── */
.streamlit-expanderHeader {
  font-family: var(--sans) !important;
  font-size: 0.85rem !important;
  font-weight: 500 !important;
  color: var(--text) !important;
}

/* ── Scrollbar ── */
hr { border-color: #0E1E2E; margin: 1.2rem 0; }
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: #1E3050; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #2A4060; }

/* ── Sidebar logo / app name ── */
.app-logo {
  text-align: center;
  padding: 10px 0 16px;
  border-bottom: 1px solid #0E1E2E;
  margin-bottom: 4px;
}
.app-logo .name {
  font-size: 1rem;
  font-weight: 700;
  color: var(--accent);
  letter-spacing: 1px;
}
.app-logo .sub {
  font-size: 0.65rem;
  color: var(--muted);
  margin-top: 2px;
}

/* ── Keyboard focus — the custom HTML cards/buttons below had no visible
   focus state at all, unlike Streamlit's own widgets. Additive only. ── */
.stButton > button:focus-visible,
.card:focus-visible, .card:has(:focus-visible),
.ag-stat:focus-visible, .ag-am:focus-visible,
[data-testid="stVerticalBlockBorderWrapper"]:has(button:focus-visible) {
  outline: 2px solid var(--accent) !important;
  outline-offset: 2px !important;
}

/* ── Minimal responsiveness — this app is a desktop SOC console and isn't
   being redesigned for mobile, but the custom flex rows below had zero
   @media rules and would just compress/clip on a laptop-width window. ── */
@media (max-width: 900px) {
  .ag-stats, .ag-attn, .ag-metas {
    flex-direction: column;
  }
  .ag-metas { align-items: stretch; }
}
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE  — seed from .env on first load
# ══════════════════════════════════════════════════════════════════════════════
_env = env_load()
# Decode base64 password
if _env["password"]:
    try:
        _env["password"] = base64.b64decode(_env["password"].encode()).decode("utf-8")
    except Exception:
        pass
# Decode base64 Cisco API key
if _env.get("cisco_key"):
    try:
        _env["cisco_key"] = base64.b64decode(_env["cisco_key"].encode()).decode("utf-8")
    except Exception:
        pass

DEFAULTS = {
    "nw_host":          _env["host"],
    "nw_username":      _env["username"],
    "nw_password":      _env["password"],
    "nw_token":         "",       # refreshed on every login
    "nw_verified":      False,
    "nw_msg":           "",
    "nw_working_ep":    "",       # endpoint that worked
    "nw_working_auth":  {},       # auth headers that worked
    "nw_incidents_path":"/rest/api/incidents",   # ← configurable endpoint path
    "nw_auth_style":    "NetWitness-Token",       # ← configurable auth header style
    "endpoint_scan_results": [],   # ← persisted scanner results
    "nw_cert_path":     _env.get("nw_cert_path", ""),  # ← path to CA/server cert for TLS verification
    "incidents":        [],
    "last_fetch":       None,
    "last_full_fetch":  None,   # ← last time a full (non-incremental) fetch ran
    "last_fetch_mode":  None,   # ← "full" | "incremental", shown in diagnostics
    "chat_history":     [],
    "chat_incident":    None,
    "pending_auto_triage": False,   # set by "Triage" button — auto-runs the pipeline
    "jump_to_ask_tab":     False,   # set by "Triage" button — switches to Ask a Question tab
    "nav_page":            "Overview", # ← top-nav active page
    "selected_case_id":    None,    # ← case loaded in My Workspace, if any
    "workspace_chat":      {},      # ← {case_id: [messages]} for Ask Aegis
    "triage_in_flight":    None,    # {incident_id, attempts} — survives an interrupted
                                    # inline triage run so it can auto-restart
    "chroma_client":    None,
    "chroma_col":       None,
    "search_results":   [],
    "_startup_done":    False,
    # ── File upload ──────────────────────────────────────────
    "uploaded_incident": None,
    "uploaded_filename": "",
    # ── Agent board (thinking + outputs per agent) ───────────
    "agent_board": {
        "triage":        {"status": "idle", "thinking": [], "output": "", "updated": "", "progress": 0},
        "investigation": {"status": "idle", "thinking": [], "output": "", "updated": "", "progress": 0},
        "reporting":     {"status": "idle", "thinking": [], "output": "", "updated": "", "progress": 0,
                          "exports": {}, "reporting_data": None},
    },
    "agent_board_sel": None,
    # ── Cisco Foundation LLM ─────────────────────────────────
    "cisco_url":         _env.get("cisco_url",   ""),
    "cisco_key":         _env.get("cisco_key",   ""),
    "cisco_model":       _env.get("cisco_model", ""),
    "cisco_connected":   bool(_env.get("cisco_key", "").strip()),
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

REFRESH_INTERVAL     = 30                      # seconds
INCREMENTAL_OVERLAP  = timedelta(minutes=5)    # clock-skew / indexing-lag buffer
# Full ground-truth resync cadence. At 53k incidents a full resync re-downloads
# EVERYTHING plus per-incident alerts — minutes of churn — so hourly (not every
# 10 min) is the default; override with NW_FULL_RESYNC_MIN. Incremental
# refreshes still run every 30s and stay cheap.
FULL_RESYNC_INTERVAL = timedelta(
    minutes=int(os.environ.get("NW_FULL_RESYNC_MIN", "60") or 60))


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def normalise_sev(inc: dict) -> str:
    raw = str(inc.get("riskScore") or inc.get("severity") or "").upper()
    try:
        s = int(raw)
        return "CRITICAL" if s >= 90 else "HIGH" if s >= 70 else "MEDIUM" if s >= 40 else "LOW"
    except ValueError:
        return raw if raw in ("CRITICAL","HIGH","MEDIUM","LOW") else "LOW"

def gp_is_reachable() -> tuple[bool, str]:
    """
    Check if the NW server is reachable via GlobalProtect VPN.
    Simply tries to open a TCP socket to the host:port.
    Returns (reachable, message).
    """
    import socket
    from urllib.parse import urlparse
    host = st.session_state.nw_host.rstrip("/")
    if not host:
        return False, "No host configured."
    try:
        parsed = urlparse(host)
        hostname = parsed.hostname or host
        port     = parsed.port or (443 if parsed.scheme == "https" else 80)
        sock = socket.create_connection((hostname, port), timeout=5)
        sock.close()
        return True, f"VPN reachable — {hostname}:{port}"
    except OSError:
        return False, (
            "Cannot reach server — GlobalProtect may not be connected.\n"
            "Connect GP VPN and try again."
        )

def nw_headers(include_content_type: bool = False) -> dict:
    """
    Build request headers based on the currently selected auth style.
    include_content_type should only be True for POST/PATCH requests that
    send a JSON body. GET requests must NOT include Content-Type or NW
    returns 400 'Unsupported format was supplied'.
    """
    token = st.session_state.nw_token.strip()
    style = st.session_state.get("nw_auth_style", "NetWitness-Token")

    base = {
        "Accept": "application/json;charset=UTF-8",
    }
    if include_content_type:
        base["Content-Type"] = "application/json;charset=UTF-8"

    if style == "NetWitness-Token":
        base["NetWitness-Token"] = token
    elif style == "Bearer":
        base["Authorization"] = f"Bearer {token}"
    elif style == "Cookie":
        base["Cookie"] = f"access_token={token}"
    elif style == "Both":
        base["Authorization"] = f"Bearer {token}"
        base["Cookie"]        = f"access_token={token}"
    else:
        base["NetWitness-Token"] = token

    return base


def nw_tls_verify():
    """
    Returns the value to pass as requests(verify=...).
    - If a valid cert path is configured AND the file exists → try it.
    - If the cert causes any SSL error → silently fall back to False.
    - Default is False (suppressed via urllib3.disable_warnings above).
    """
    cert_path = st.session_state.get("nw_cert_path", "").strip()
    if cert_path and Path(cert_path).is_file():
        return cert_path
    return False


def nw_incidents_url(host: str) -> str:
    """Build the incidents endpoint URL from the configurable path."""
    path = st.session_state.get("nw_incidents_path", "/rest/api/incidents").strip()
    if not path.startswith("/"):
        path = "/" + path
    return f"{host.rstrip('/')}{path}"

def nw_verify_token() -> tuple[bool, str]:
    host  = st.session_state.nw_host.rstrip("/")
    token = st.session_state.nw_token.strip()
    if not host or not token:
        return False, "Enter both Host URL and token."

    def _attempt(verify):
        return requests.get(
            nw_incidents_url(host),
            headers=nw_headers(),
            params={"pageSize": 1, "pageNumber": 0},
            timeout=15,
            verify=verify,
        )
    try:
        try:
            r = _attempt(nw_tls_verify())
        except requests.exceptions.SSLError:
            st.session_state.nw_cert_path = ""
            r = _attempt(False)

        if r.status_code == 200:
            return True, "NetWitness connected"
        elif r.status_code == 403:
            return False, (
                "Access Denied — add 'integration-server.api.access' "
                "permission to the Administrators role in NW."
            )
        elif r.status_code == 401:
            return False, "Token rejected — login again."
        else:
            return False, f"Unexpected response — HTTP {r.status_code}"
    except requests.exceptions.Timeout:
        return False, "Timed out — check GP VPN is connected."
    except Exception as e:
        return False, f"Cannot reach server: {e}"

def _bounded_get(url: str, *, headers=None, params=None, verify=False,
                 timeout: int = 30, wall_seconds: int = 45):
    """requests.get with a HARD wall-clock cap.

    The `timeout` parameter only bounds the gap between socket reads — a
    server that dribbles a byte every few seconds holds the connection open
    forever (this wedged the whole app for hours). The request runs in a
    daemon thread; if it exceeds wall_seconds in total, TimeoutError is
    raised and the orphaned thread is abandoned to finish/die on its own."""
    import threading as _th
    box: dict = {}

    def _do():
        try:
            box["r"] = requests.get(url, headers=headers, params=params,
                                    timeout=timeout, verify=verify)
        except Exception as exc:
            box["e"] = exc

    t = _th.Thread(target=_do, daemon=True)
    t.start()
    t.join(timeout=wall_seconds)
    if t.is_alive():
        raise TimeoutError(f"request exceeded {wall_seconds}s wall clock "
                           f"(server stalling mid-response)")
    if "e" in box:
        raise box["e"]
    return box["r"]


# ── alerts-response tolerance (NetWitness Respond API shape varies by version) ──
# These pure helpers now live in nw_alerts.py (first slice of the app.py
# modularization). Imported here so every existing call site keeps working
# unchanged. See nw_alerts.py for the docstrings + rationale.
from nw_alerts import (
    _extract_alert_items,
    _alerts_has_more,
    _alerts_error_hint,
    _distill_alerts,
    _merge_alert_digest,
    _alerts_fetch_warning,
)


def nw_fetch_incidents(
    limit: int | None = None, since: str | None = None,
    deadline_seconds: int | None = None,
) -> tuple[bool, list, str]:
    """
    Returns (ok, items, diagnostic_message).
    ok=True even if items is empty (empty just means no incidents in NW right now).
    ok=False means a real error occurred (auth, network, etc).

    limit=None (default) means load every incident the API has — pagination
    keeps going until the API reports hasNext=False. MAX_PAGES is a safety
    valve, not a real-world cap, in case an API bug ever returns hasNext=True
    forever.

    since=None (default) means a full fetch — every incident NetWitness has.
    Pass an ISO8601 cutoff (see incremental_since()) to ask NetWitness for
    only incidents created/updated after that point — used by the periodic
    auto-refresh so it isn't re-fetching + re-enriching the entire incident
    history (and every incident's full alert history) every 30s.
    """
    if not st.session_state.nw_verified or not st.session_state.nw_token:
        return False, [], "Not authenticated."
    host = st.session_state.nw_host.rstrip("/")
    headers = nw_headers()
    all_items = []
    page = 0
    diag = ""
    MAX_PAGES = 1000   # 1000 * pageSize(100) = 100,000 incidents ceiling
    # Wall-clock budget: a stalling NetWitness (slow pages, dribbling
    # responses) must yield PARTIAL results with an honest diagnostic —
    # never an unbounded hang. None -> env NW_FETCH_DEADLINE (default 600s).
    if deadline_seconds is None:
        try:
            deadline_seconds = int(os.environ.get("NW_FETCH_DEADLINE", "600"))
        except ValueError:
            deadline_seconds = 600
    _deadline = time.monotonic() + max(30, deadline_seconds)
    # Include a wide date range by default — some NW versions return 400 without it
    since = since or "2020-01-01T00:00:00.000Z"
    try:
        while True:
            if time.monotonic() > _deadline:
                diag = (f"Fetch deadline ({deadline_seconds}s) reached at "
                        f"page {page} — returning {len(all_items)} incident(s) "
                        f"fetched so far. NetWitness is responding slowly.")
                return True, all_items, diag
            try:
                r = _bounded_get(
                    nw_incidents_url(host),
                    headers=nw_headers(),
                    params={"pageSize": 100, "pageNumber": page, "since": since},
                    timeout=30,
                    wall_seconds=60,
                    verify=False,
                )
            except TimeoutError:
                diag = (f"NetWitness stalled mid-response on page {page} — "
                        f"returning {len(all_items)} incident(s) fetched so far.")
                return True, all_items, diag
            if r.status_code == 200:
                data      = r.json()
                items     = data.get("items", [])
                total_api = data.get("totalItems", "?")
                all_items.extend(items)
                has_next  = data.get("hasNext", False)
                reached_limit = limit is not None and len(all_items) >= limit
                if not has_next or reached_limit or page >= MAX_PAGES:
                    diag = (
                        f"API reports {total_api} total incident(s) — "
                        f"fetched {len(all_items)} across {page+1} page(s)."
                    )
                    break
                page += 1
            elif r.status_code == 401:
                st.session_state.nw_verified = False
                st.session_state.nw_msg = "Token expired — login again."
                return False, [], "401 Unauthorized — token expired or invalid."
            elif r.status_code == 403:
                return False, [], (
                    "403 Forbidden — account lacks integration-server.api.access permission."
                )
            else:
                return False, [], f"HTTP {r.status_code}: {r.text[:200]}"
        # Fetch associated alerts/logs for each incident — in parallel.
        # This used to be a serial for-loop doing one blocking request per
        # incident (up to 15s timeout each), so on a list of N incidents the
        # whole fetch took N * request-time. Since this runs on every
        # auto-refresh and on every rerun after one is due (sending a chat
        # message, uploading a file — any Streamlit interaction reruns the
        # whole script), that serial loop is what made the entire UI look
        # frozen. Fetching concurrently bounds the wall-clock time to
        # roughly one request instead of N.
        clean_path = st.session_state.get("nw_incidents_path", "/rest/api/incidents").strip()
        if clean_path.endswith("/list"):
            clean_path = clean_path[:-5]

        def _fetch_alerts(inc: dict) -> None:
            inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
            if not inc_id:
                inc["alerts"] = []
                return
            alerts_url = f"{host.rstrip('/')}{clean_path}/{inc_id}/alerts"
            # Paginate fully — this used to fetch only pageNumber=0, so any
            # incident with more than 100 alerts silently lost the rest.
            collected: list = []
            a_page = 0
            try:
                while a_page < MAX_PAGES:
                    if time.monotonic() > _deadline:
                        inc["alerts_fetch_error"] = "fetch deadline reached"
                        break
                    r_alerts = _bounded_get(
                        alerts_url,
                        headers=headers,
                        params={"pageSize": 100, "pageNumber": a_page},
                        timeout=10,
                        wall_seconds=20,
                        verify=False,
                    )
                    if r_alerts.status_code != 200:
                        # Remember WHY alerts are missing, with enough to act on:
                        # code + the exact URL + a body snippet + a status-specific
                        # hint. The triage/investigation handoff and the UI surface
                        # this instead of silently passing an incident with no
                        # event data. (A live/VPN run now yields an actionable
                        # error rather than a bare code.)
                        try:
                            _body = (r_alerts.text or "")[:200]
                        except Exception:
                            _body = ""
                        inc["alerts_fetch_error"] = f"HTTP {r_alerts.status_code}"
                        inc["alerts_fetch_diag"] = {
                            "code": r_alerts.status_code, "url": alerts_url,
                            "body": _body, "hint": _alerts_error_hint(r_alerts.status_code)}
                        break
                    try:
                        a_data = r_alerts.json()
                    except Exception:
                        a_data = None
                    # Tolerate response-shape variance between NW versions: a 200
                    # that isn't {items,hasNext} used to yield 0 alerts silently.
                    page_items = _extract_alert_items(a_data)
                    collected.extend(page_items)
                    if not _alerts_has_more(a_data, a_page):
                        break
                    a_page += 1
            except Exception as _exc:
                inc["alerts_fetch_error"] = str(_exc)[:120]
                inc["alerts_fetch_diag"] = {"code": None, "url": alerts_url,
                                            "body": "", "hint": "network/exception — check VPN & host reachability"}
            # Tag every alert with its parent incident ID so alerts are
            # traceable back to the incident they came from.
            for a in collected:
                a["incident_id"] = inc_id
            inc["alerts"] = collected
            # Distill the alerts' rich event fields (endpoint host, users, IPs,
            # MACs, alert titles, MITRE) into alertMeta so triage/investigation
            # and the skills see a real host/user — the incident object alone
            # only carries SourceIp/DestinationIp.
            try:
                _merge_alert_digest(inc)
            except Exception:
                pass

        if all_items:
            if time.monotonic() > _deadline:
                diag += " Alert enrichment skipped (fetch deadline reached)."
                for inc in all_items:
                    inc.setdefault("alerts", [])
                    inc["alerts_fetch_error"] = "fetch deadline reached"
            else:
                with ThreadPoolExecutor(max_workers=min(16, len(all_items))) as pool:
                    list(pool.map(_fetch_alerts, all_items))

        return True, all_items, diag
    except requests.exceptions.Timeout:
        return False, [], "Request timed out — check VPN/network."
    except Exception as e:
        return False, [], f"Exception: {e}"

# ══════════════════════════════════════════════════════════════════════════════
# SQLITE  — permanent incident log
# ══════════════════════════════════════════════════════════════════════════════
import sqlite3
import json as _json

# All SQLite databases now live in soc_db/ (moved 2026-07-09).
# db_connect()/db_init() moved to workflow_state_store.py (imported above) so
# there is exactly one owner of the incidents table schema, shared with
# soc_workflow.py's run_until_triage_approval().
SOC_DB_DIR = Path(__file__).parent / "soc_db"
DB_FILE = SOC_DB_DIR / "soc_incidents.db"

def db_upsert_incidents(incidents: list) -> int:
    """Upsert a list of incident dicts — new rows inserted, existing rows updated."""
    now = datetime.now().isoformat(timespec="seconds")
    rows = []
    for inc in incidents:
        inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if not inc_id:
            continue
        # Store a SLIM raw_json: the alerts arrays are refetched live and made
        # this DB grow to 2.4 GB; alertMeta (small, carries the only incident-
        # level indicators) is kept. In-memory incidents stay untouched.
        slim = {k: v for k, v in inc.items()
                if k not in ("alerts", "journalEntries")}
        if inc.get("alerts"):
            slim["_alerts_stripped"] = len(inc["alerts"])
        rows.append((
            inc_id,
            inc.get("title") or inc.get("name") or "",
            normalise_sev(inc),
            str(inc.get("status") or ""),
            str(inc.get("assignee") or ""),
            int(inc.get("alertCount") or inc.get("numAlerts") or 0),
            str(inc.get("created") or inc.get("createdDate") or "")[:19],
            str(inc.get("updated") or inc.get("lastUpdated") or "")[:19],
            _json.dumps(slim),
            now,   # first_seen  — INSERT only
            now,   # last_seen   — always updated
        ))
    if not rows:
        return 0
    # Chunked commits: one giant transaction over 50k+ rows held the write
    # lock for the whole upsert (the 2.4 GB DB made that minutes) — any
    # other connection in that window died with "database is locked".
    CHUNK = 2000
    with db_connect() as con:
        for _i in range(0, len(rows), CHUNK):
            _chunk = rows[_i:_i + CHUNK]
            con.executemany("""
            INSERT INTO incidents
                (id, title, severity, status, assignee, alert_count,
                 created, updated, raw_json, first_seen, last_seen)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(id) DO UPDATE SET
                title       = excluded.title,
                severity    = excluded.severity,
                status      = excluded.status,
                assignee    = excluded.assignee,
                alert_count = excluded.alert_count,
                updated     = excluded.updated,
                raw_json    = excluded.raw_json,
                last_seen   = excluded.last_seen
            """, _chunk)
            con.commit()   # release the write lock between chunks
        con.execute(
            "INSERT INTO fetch_log (fetched_at, count) VALUES (?,?)",
            (now, len(rows)),
        )
        con.commit()
    return len(rows)

def db_get_incident(inc_id: str) -> dict | None:
    """Exact-id lookup — used by My Workspace to resolve a clicked case."""
    if not inc_id:
        return None
    with db_connect() as con:
        row = con.execute("SELECT * FROM incidents WHERE id = ?",
                          (str(inc_id),)).fetchone()
    return dict(row) if row else None

def db_set_parsing_status(inc_id: str, status: str) -> None:
    with db_connect() as con:
        con.execute("UPDATE incidents SET parsing_status = ? WHERE id = ?",
                    (status, str(inc_id)))
        con.commit()

def db_save_parsing_result(inc_id: str, result: dict) -> None:
    """Persist a compact parsing summary — full normalised_alert/processed_alert
    stay on disk (see run_parsing); the DB only needs enough to drive status UI."""
    parsing_status = "Complete" if result.get("status") == "completed" else "Failed"
    summary = {
        "status": result.get("status"),
        "parser_confidence": result.get("parser_confidence"),
        "recommended_next_action": result.get("recommended_next_action"),
        "output_files": result.get("output_files"),
        "ai_summary": result.get("ai_summary"),
        "ai_thinking": result.get("ai_thinking"),
        "ai_summary_model": result.get("ai_summary_model"),
        "completed_at": datetime.now().isoformat(timespec="seconds"),
    }
    with db_connect() as con:
        con.execute(
            "UPDATE incidents SET parsing_status = ?, parsing_result_json = ? WHERE id = ?",
            (parsing_status, _json.dumps(summary), str(inc_id)),
        )
        con.commit()

def db_load_parsed_context(inc_id: str) -> dict | None:
    """Load the processed_alert this incident's Parsing stage already produced
    (via the Start Process button), for feeding into triage as parsed_context.
    db_save_parsing_result only stores a compact summary + the on-disk paths
    (see its docstring), so this reads the flat processed_alert back off disk.
    Returns None if parsing hasn't completed for this incident — triage falls
    back to its standalone behaviour in that case, same as before this existed."""
    row = db_get_incident(inc_id)
    if not row or row.get("parsing_status") != "Complete":
        return None
    try:
        summary = _json.loads(row.get("parsing_result_json") or "{}")
        path = (summary.get("output_files") or {}).get("processed_alert_flat")
        if not path or not os.path.exists(path):
            return None
        with open(path, "r", encoding="utf-8") as f:
            return _json.load(f)
    except Exception:
        return None

def db_load_incidents(
    severity: str = "ALL",
    status:   str = "ALL",
    search:   str = "",
    limit:    int = 500,
) -> list[dict]:
    """Query incidents from SQLite with optional filters."""
    clauses, params = [], []
    if severity != "ALL":
        clauses.append("severity = ?"); params.append(severity)
    if status != "ALL":
        clauses.append("status = ?");   params.append(status)
    if search.strip():
        clauses.append("(title LIKE ? OR assignee LIKE ? OR id LIKE ?)")
        params += [f"%{search}%", f"%{search}%", f"%{search}%"]
    where = ("WHERE " + " AND ".join(clauses)) if clauses else ""
    with db_connect() as con:
        rows = con.execute(
            f"SELECT * FROM incidents {where} ORDER BY last_seen DESC LIMIT ?",
            params + [limit],
        ).fetchall()
    return [dict(r) for r in rows]

_CLOSED_STATUSES = ("CLOSED", "RESOLVED", "REMEDIATED")

def db_stats() -> dict:
    """Return aggregate stats from the permanent log."""
    with db_connect() as con:
        total   = con.execute("SELECT COUNT(*) FROM incidents").fetchone()[0]
        by_sev  = dict(con.execute(
            "SELECT severity, COUNT(*) FROM incidents GROUP BY severity"
        ).fetchall())
        by_stat = dict(con.execute(
            "SELECT status, COUNT(*) FROM incidents GROUP BY status"
        ).fetchall())
        # "active" = not closed/resolved/remediated — the counts that
        # actually mean "needs attention now", vs. the all-time totals
        # above which just answer "how big is the archive".
        _closed_q = ",".join("?" * len(_CLOSED_STATUSES))
        by_sev_active = dict(con.execute(
            f"SELECT severity, COUNT(*) FROM incidents "
            f"WHERE status NOT IN ({_closed_q}) OR status IS NULL "
            f"GROUP BY severity", _CLOSED_STATUSES
        ).fetchall())
        unassigned_active = con.execute(
            f"SELECT COUNT(*) FROM incidents "
            f"WHERE (assignee IS NULL OR assignee='' OR assignee='Unassigned') "
            f"AND (status NOT IN ({_closed_q}) OR status IS NULL)",
            _CLOSED_STATUSES
        ).fetchone()[0]
        fetches = con.execute("SELECT COUNT(*) FROM fetch_log").fetchone()[0]
        last_f  = con.execute(
            "SELECT fetched_at FROM fetch_log ORDER BY id DESC LIMIT 1"
        ).fetchone()
    return {
        "total":   total,
        "by_sev":  by_sev,
        "by_stat": by_stat,
        "by_sev_active":     by_sev_active,
        "unassigned_active": unassigned_active,
        "fetches": fetches,
        "last_fetch": last_f[0] if last_f else "—",
    }

def db_export_csv() -> str:
    """Return all incidents as a CSV string for download."""
    import csv, io
    rows = db_load_incidents(limit=100_000)
    if not rows:
        return ""
    buf = io.StringIO()
    w   = csv.DictWriter(buf, fieldnames=[
        "id","title","severity","status","assignee",
        "alert_count","created","updated","first_seen","last_seen",
    ])
    w.writeheader()
    for r in rows:
        w.writerow({k: r.get(k,"") for k in w.fieldnames})
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# STARTUP — runs once per session, silently connects & fetches
# ══════════════════════════════════════════════════════════════════════════════

# Initialise DB on every startup (no-op if tables already exist)
db_init()
# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE DATABASE — per-stage SQLite + ChromaDB, fully inline (no 2nd app)
# ══════════════════════════════════════════════════════════════════════════════
PIPELINE_STAGES = [
    "alerts_to_triage",
    "post_triage_investigate",
    "post_triage_no_investigate",
    "post_investigation",
    "initial_ticket",
    "pending_ticket_report",
    "finalized_report",
    "workflow_runs",
]
PIPELINE_LABELS = {
    "alerts_to_triage":           "Alerts to Triage",
    "post_triage_investigate":    "Post-Triage · Needs Investigation",
    "post_triage_no_investigate": "Post-Triage · No Investigation Needed",
    "post_investigation":         "Post-Investigation · Findings",
    "initial_ticket":             "Initial Ticket Generation",
    "pending_ticket_report":      "Pending Ticket / Report Generation",
    "finalized_report":           "Finalized Report",
    "workflow_runs":              "Workflow Runs (Audit)",
}
PIPELINE_ICONS = {
    "alerts_to_triage":           "◎",
    "post_triage_investigate":    "⌕",
    "post_triage_no_investigate": "○",
    "post_investigation":         "▤",
    "initial_ticket":             "▣",
    "pending_ticket_report":      "◧",
    "finalized_report":           "✦",
    "workflow_runs":              "≡",
}
PIPELINE_COLORS = {
    "alerts_to_triage":           "#FF3B3B",
    "post_triage_investigate":    "#FF7700",
    "post_triage_no_investigate": "#0AF0A0",
    "post_investigation":         "#2DD4BF",
    "initial_ticket":             "#00D4FF",
    "pending_ticket_report":      "#FFB700",
    "finalized_report":           "#A78BFA",
    "workflow_runs":              "#8B9DC3",
}
PIPELINE_DB_FILE = SOC_DB_DIR / "soc_pipeline.db"

def _pl_con():
    # Busy-timeout so UI reads and the background worker's writes never
    # collide into silent 'database is locked' failures.
    con = sqlite3.connect(str(PIPELINE_DB_FILE), check_same_thread=False,
                          timeout=15)
    con.row_factory = sqlite3.Row
    return con

def pipeline_db_init():
    with _pl_con() as c:
        try:
            c.execute("PRAGMA journal_mode=WAL")
            c.execute("PRAGMA synchronous=NORMAL")
        except Exception:
            pass
        for s in PIPELINE_STAGES:
            c.execute(f"""CREATE TABLE IF NOT EXISTS {s} (
                id TEXT PRIMARY KEY, incident_id TEXT, title TEXT,
                severity TEXT, stage TEXT, created_at TEXT,
                summary TEXT, raw_json TEXT)""")
        c.commit()

def pipeline_count(stage):
    try:
        with _pl_con() as c:
            return c.execute(f"SELECT COUNT(*) FROM {stage}").fetchone()[0]
    except Exception:
        return 0

def pipeline_load(stage, limit=300):
    try:
        with _pl_con() as c:
            rows = c.execute(
                f"SELECT * FROM {stage} ORDER BY created_at DESC LIMIT ?", (limit,)
            ).fetchall()
        return [dict(r) for r in rows]
    except Exception:
        return []

def pipeline_insert(stage, record):
    import uuid as _uuid
    rec_id = str(record.get("id") or record.get("unc") or _uuid.uuid4())[:64]
    now = datetime.now().isoformat(timespec="seconds")
    try:
        with _pl_con() as c:
            # Same-id inserts REPLACE the row (re-running the same incident
            # reuses its ticket ids), which used to look like "nothing
            # happened". Track a run counter and stamp the summary so a
            # refreshed record is unmistakably new.
            runs = 1
            prev = None
            try:
                prev = c.execute(f"SELECT raw_json FROM {stage} WHERE id=?",
                                 (rec_id,)).fetchone()
                if prev:
                    runs = int((_json.loads(prev[0] or "{}"))
                               .get("workflow_runs_count") or 1) + 1
            except Exception:
                runs = 2 if prev else 1
            record = dict(record)
            record["workflow_runs_count"] = runs
            summary = str(record.get("summary") or record.get("description") or "")
            if runs > 1:
                summary = f"[run {runs} · {now[11:19]}] {summary}"
            c.execute(
                f"INSERT OR REPLACE INTO {stage} "
                "(id,incident_id,title,severity,stage,created_at,summary,raw_json) "
                "VALUES (?,?,?,?,?,?,?,?)",
                (rec_id,
                 str(record.get("incident_id") or record.get("incidentId") or ""),
                 str(record.get("title") or record.get("name") or ""),
                 str(record.get("severity") or record.get("classification") or ""),
                 stage, now,
                 summary[:500],
                 _json.dumps(record)))
            c.commit()
    except Exception:
        pass
    return rec_id


def pipeline_last_write(stage):
    """Most recent created_at in a stage — shown on the stage card so it's
    obvious when the stage last changed (counts alone hide REPLACEs)."""
    try:
        with _pl_con() as c:
            v = c.execute(f"SELECT MAX(created_at) FROM {stage}").fetchone()[0]
        return str(v)[5:16].replace("T", " ") if v else "—"
    except Exception:
        return "—"

def _pl_chroma_col(stage):
    if not CHROMA_OK or not st.session_state.get("chroma_client"):
        return None
    try:
        return st.session_state.chroma_client.get_or_create_collection(
            name=f"pipeline_{stage}", metadata={"hnsw:space": "cosine"},
            embedding_function=_openai_ef())
    except Exception:
        return None

def pipeline_chroma_insert(stage, record):
    col = _pl_chroma_col(stage)
    if col is None:
        return
    import uuid as _uuid
    rec_id = str(record.get("id") or record.get("unc") or _uuid.uuid4())[:64]
    title = str(record.get("title") or "")
    summary = str(record.get("summary") or record.get("description") or "")
    doc = f"{title}\n{summary}".strip() or "no content"
    try:
        col.upsert(documents=[doc], ids=[rec_id],
                   metadatas=[{"stage": stage,
                                "severity": str(record.get("severity") or ""),
                                "created": datetime.now().isoformat(timespec="seconds")}])
    except Exception:
        pass

def pipeline_chroma_count(stage):
    col = _pl_chroma_col(stage)
    if col is None:
        return 0
    try:
        return col.count()
    except Exception:
        return 0

def pipeline_chroma_search(stage, query, n=5):
    col = _pl_chroma_col(stage)
    if col is None or not query.strip():
        return []
    try:
        total = col.count()
        if total == 0:
            return []
        res = col.query(query_texts=[query], n_results=min(n, total))
        return [{"id": res["ids"][0][i], "doc": d,
                 "score": round((1 - res["distances"][0][i]) * 100, 1),
                 "meta": res["metadatas"][0][i]}
                for i, d in enumerate(res["documents"][0])]
    except Exception as e:
        return [{"id": "err", "doc": str(e), "score": 0, "meta": {}}]

def pipeline_chroma_all(stage):
    col = _pl_chroma_col(stage)
    if col is None:
        return []
    try:
        data = col.get(include=["documents", "metadatas"])
        return [{"id": data["ids"][i],
                 "doc": (data["documents"] or [""])[i],
                 "meta": (data["metadatas"] or [{}])[i]}
                for i in range(len(data["ids"]))]
    except Exception:
        return []

def pipeline_insert_full(stage, record):
    """Insert into SQLite and ChromaDB (if connected)."""
    rec_id = pipeline_insert(stage, record)
    pipeline_chroma_insert(stage, record)
    return rec_id

def pipeline_delete(stage, rec_id):
    try:
        with _pl_con() as c:
            c.execute(f"DELETE FROM {stage} WHERE id=?", (rec_id,))
            c.commit()
    except Exception:
        pass
    col = _pl_chroma_col(stage)
    if col:
        try:
            col.delete(ids=[rec_id])
        except Exception:
            pass

pipeline_db_init()

# ──────────────────────────────────────────────────────────────────────────────
# EXPORT HELPERS  — generate downloadable bytes from a pipeline record
# ──────────────────────────────────────────────────────────────────────────────

def _make_csv_bytes(row: dict) -> bytes:
    """
    Build a CSV file from a pipeline SQLite row.
    Flattens the raw_json payload if present so the sheet has all fields.
    Returns UTF-8 bytes suitable for st.download_button.
    """
    import csv, io
    # Start with the fixed SQLite columns
    base = {
        "id":          row.get("id", ""),
        "incident_id": row.get("incident_id", ""),
        "title":       row.get("title", ""),
        "severity":    row.get("severity", ""),
        "stage":       row.get("stage", ""),
        "created_at":  row.get("created_at", ""),
        "summary":     row.get("summary", ""),
    }
    # Merge in any extra keys from raw_json (flat scalars only)
    try:
        extra = _json.loads(row.get("raw_json") or "{}")
        if isinstance(extra, dict):
            for k, v in extra.items():
                if k not in base and not isinstance(v, (dict, list)):
                    base[k] = str(v)
    except Exception:
        pass
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=list(base.keys()))
    writer.writeheader()
    writer.writerow(base)
    return buf.getvalue().encode("utf-8")


def _make_docx_bytes(row: dict) -> bytes:
    """
    Build a professional SOC Initial Ticket .docx from a pipeline record.
    Uses python-docx (pip install python-docx). Falls back to a plain-text
    .txt wrapped as bytes if the library is unavailable.
    """
    import io
    title      = row.get("title", "Untitled Ticket")
    inc_id     = row.get("incident_id", "—")
    severity   = row.get("severity", "—")
    stage      = row.get("stage", "initial_ticket")
    created_at = row.get("created_at", "—")
    summary    = row.get("summary", "No summary available.")

    # Try to parse raw_json for richer data
    extra = {}
    try:
        extra = _json.loads(row.get("raw_json") or "{}")
        if not isinstance(extra, dict):
            extra = {}
    except Exception:
        pass

    try:
        from docx import Document as _DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = _DocxDocument()

        # ── Page margins ─────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # ── Title ────────────────────────────────────────────
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)
            run.font.size = Pt(16)

        doc.add_paragraph()

        # ── Metadata table ───────────────────────────────────
        meta_fields = [
            ("Ticket ID",     row.get("id", "—")),
            ("Incident ID",   inc_id),
            ("Severity",      severity),
            ("Stage",         stage),
            ("Generated At",  created_at),
            ("Assignee",      str(extra.get("assignee") or "Unassigned")),
            ("Status",        str(extra.get("status") or "NEW")),
        ]
        tbl = doc.add_table(rows=len(meta_fields), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(meta_fields):
            tbl.cell(i, 0).text = label
            tbl.cell(i, 1).text = str(value)
            # Bold the label
            for run in tbl.cell(i, 0).paragraphs[0].runs:
                run.bold = True

        doc.add_paragraph()

        # ── Summary section ───────────────────────────────────
        doc.add_heading("Triage Summary", level=2)
        doc.add_paragraph(summary)

        # ── IOC / Alert section if present ───────────────────
        alert_count = extra.get("alertCount") or extra.get("numAlerts") or extra.get("alert_count")
        if alert_count:
            doc.add_paragraph()
            doc.add_heading("Alert Count", level=2)
            doc.add_paragraph(str(alert_count))

        description = extra.get("description") or extra.get("raw_log") or ""
        if description and description != summary:
            doc.add_paragraph()
            doc.add_heading("Description / Raw Log", level=2)
            # Cap at 3000 chars to keep document manageable
            doc.add_paragraph(str(description)[:3000])

        # ── Footer note ───────────────────────────────────────
        doc.add_paragraph()
        footer_para = doc.add_paragraph(
            f"Auto-generated by SOC Platform · {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x3A, 0x60, 0x7A)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # python-docx not installed — return UTF-8 text as fallback
        lines = [
            f"SOC INITIAL TICKET",
            f"==================",
            f"Title:      {title}",
            f"Ticket ID:  {row.get('id', '—')}",
            f"Incident:   {inc_id}",
            f"Severity:   {severity}",
            f"Stage:      {stage}",
            f"Created:    {created_at}",
            f"",
            f"SUMMARY",
            f"-------",
            summary,
            f"",
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        return "\n".join(lines).encode("utf-8")


# Step 1 — load persisted incidents from SQLite immediately
# so the UI is never empty even before the live fetch completes
if not st.session_state.incidents:
    _db_rows = db_load_incidents(limit=500)
    if _db_rows:
        st.session_state.incidents = [
            _json.loads(r["raw_json"]) for r in _db_rows if r.get("raw_json")
        ]

# Step 2 — silently auto-verify & auto-fetch if .env has credentials.
# Runs once per session (tracked by _startup_done flag) in a BACKGROUND
# thread with a soft wait: a slow or wedged NetWitness (mid-response
# stalls evade per-request timeouts) previously hung every page load
# forever, before a single element rendered. The UI now renders from the
# SQLite cache immediately; live data is adopted whenever the fetch lands.
if not st.session_state._startup_done and _env["host"] and _env["username"] and _env["password"]:
    st.session_state._startup_done = True
    st.session_state.nw_host     = _env["host"]
    st.session_state.nw_username = _env["username"]
    st.session_state.nw_password = _env["password"]

    st.session_state._startup_fetching = True

    def _startup_login_and_fetch() -> None:
        try:
            ok, msg, token = nw_login(_env["host"], _env["username"],
                                      _env["password"])
            if not ok:
                st.session_state.nw_msg = msg
                return
            st.session_state.nw_token    = token
            st.session_state.nw_verified = True
            st.session_state.nw_msg      = msg
            # Incremental startup: when the SQLite cache is fresh (<24h since
            # the last fetch), only pull incidents changed since then instead
            # of re-downloading all 53k+ every session. A stale/empty cache
            # still gets the full fetch.
            since = None
            last_dt = None
            try:
                with db_connect() as _c:
                    _last = _c.execute(
                        "SELECT MAX(fetched_at) FROM fetch_log").fetchone()[0]
                if _last and st.session_state.incidents:
                    last_dt = datetime.fromisoformat(str(_last)[:19])
                    if (datetime.now() - last_dt) < timedelta(hours=24):
                        since = incremental_since(last_dt)
            except Exception:
                since = None
            ok2, items, _diag = nw_fetch_incidents(
                since=since,
                deadline_seconds=int(os.environ.get("NW_FETCH_DEADLINE", "300")))
            if ok2 and items:
                st.session_state.incidents = (
                    items if since is None
                    else _merge_incidents(st.session_state.incidents, items))
                st.session_state.last_fetch      = datetime.now()
                # Incremental startup: schedule the next full resync off the
                # last KNOWN fetch, not off now — ground truth stays bounded.
                st.session_state.last_full_fetch = (
                    datetime.now() if since is None else last_dt)
                st.session_state.last_fetch_mode = (
                    "full" if since is None else "incremental")
                db_upsert_incidents(items)
        except Exception as _exc:
            try:
                st.session_state.nw_msg = f"Startup fetch failed: {_exc}"
            except Exception:
                pass
        finally:
            try:
                st.session_state._startup_fetching = False
            except Exception:
                pass

    import threading as _threading
    from streamlit.runtime.scriptrunner import add_script_run_ctx
    _t = _threading.Thread(target=_startup_login_and_fetch, daemon=True)
    add_script_run_ctx(_t)   # lets the thread read/write this session's state
    _t.start()
    # Soft wait: if NetWitness answers quickly, first render has live data;
    # if it's slow or wedged, render proceeds on cached incidents and the
    # thread adopts results whenever (if ever) the fetch completes.
    _t.join(timeout=int(os.environ.get("NW_STARTUP_WAIT", "20")))
    if _t.is_alive():
        st.session_state.nw_msg = ("NetWitness is responding slowly — "
                                   "showing cached incidents; live data "
                                   "will appear when the fetch completes.")

def incremental_since(last_fetch: datetime) -> str:
    """
    ISO8601 cutoff for an incremental refresh — everything since the last
    successful fetch, minus a 5-minute overlap to tolerate clock skew and
    NetWitness indexing lag. Incidents re-fetched inside that overlap just
    get upserted again by id — harmless, not duplicated.
    """
    cutoff = last_fetch - INCREMENTAL_OVERLAP
    return cutoff.strftime("%Y-%m-%dT%H:%M:%S.000Z")


def _merge_incidents(cached: list, fresh: list) -> list:
    """
    Upsert `fresh` incidents into `cached` by id, keeping everything else
    untouched. Used for incremental refreshes, where `fresh` is only the
    new/updated subset NetWitness returned for the `since` window — not
    the full incident set, so a plain replace would drop everything older.
    """
    by_id = {
        str(inc.get("id") or inc.get("incidentId") or f"_unkeyed_{i}"): inc
        for i, inc in enumerate(cached)
    }
    for inc in fresh:
        key = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if key:
            by_id[key] = inc
    return list(by_id.values())


def maybe_auto_fetch():
    if not st.session_state.nw_verified:
        return
    # Don't stack a second fetch on top of the startup background fetch or a
    # previous auto-refresh that is still running in its worker thread.
    if st.session_state.get("_startup_fetching") or st.session_state.get("_bg_fetching"):
        return
    now  = datetime.now()
    last = st.session_state.last_fetch
    if last is None or (now - last).total_seconds() >= REFRESH_INTERVAL:
        # We don't actually know whether NetWitness's `since` filter matches
        # on incident creation time or last-updated time — the API isn't
        # documented clearly enough to bet on it. If it's creation-time only,
        # a purely incremental refresh would silently miss status/alert
        # changes on older incidents. So: incremental fetches most cycles
        # (cheap — only new/updated incidents come back), but force a full
        # ground-truth resync every FULL_RESYNC_INTERVAL regardless, which
        # caps how stale the cache can ever get at a known worst case.
        last_full  = st.session_state.last_full_fetch
        needs_full = last_full is None or (now - last_full) >= FULL_RESYNC_INTERVAL
        since = None if needs_full else incremental_since(last)

        # Budgets kept from the synchronous era as the worker's wall-clock cap.
        _budget = 45 if since is not None else 120

        # The fetch itself runs OFF the render thread (same proven pattern as
        # the startup fetch: daemon thread + add_script_run_ctx + guard flag).
        # It used to run inline here, so every REFRESH_INTERVAL the next rerun
        # (any click / chat message) blocked the whole UI for up to _budget
        # seconds. Now render returns immediately; results are adopted into
        # session state when the fetch lands and show on the next rerun —
        # exactly how the startup fetch already behaves.
        st.session_state._bg_fetching = True

        def _bg_auto_fetch():
            try:
                ok, items, _diag = nw_fetch_incidents(since=since,
                                                      deadline_seconds=_budget)
                if ok:
                    st.session_state.incidents = (
                        items if since is None
                        else _merge_incidents(st.session_state.incidents, items)
                    )
                    st.session_state.last_fetch      = datetime.now()
                    st.session_state.last_fetch_mode = "full" if since is None else "incremental"
                    if since is None:
                        st.session_state.last_full_fetch = st.session_state.last_fetch
                    db_upsert_incidents(items)   # ← persist every fetch to SQLite
            except Exception:
                pass   # failed auto-refresh = stale data; next due cycle retries
            finally:
                try:
                    st.session_state._bg_fetching = False
                except Exception:
                    pass

        import threading as _threading
        from streamlit.runtime.scriptrunner import add_script_run_ctx
        _t = _threading.Thread(target=_bg_auto_fetch, daemon=True)
        add_script_run_ctx(_t)   # thread may read/write this session's state
        _t.start()

def chroma_connect(path: str = "./chroma_db") -> tuple[bool, str]:
    if not CHROMA_OK:
        return False, "chromadb not installed — run: pip install chromadb"
    try:
        client = chromadb.PersistentClient(path=path)
        col    = client.get_or_create_collection(
            name="soc_incidents", metadata={"hnsw:space": "cosine"},
            embedding_function=_openai_ef(),
        )
        st.session_state.chroma_client = client
        st.session_state.chroma_col    = col
        return True, f"Ready — {col.count()} vectors"
    except Exception as e:
        return False, str(e)

def chroma_sync(incidents: list) -> tuple[int, str]:
    col = st.session_state.chroma_col
    if col is None:
        return 0, "Connect ChromaDB first."
    docs, ids, metas = [], [], []
    for inc in incidents:
        inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if not inc_id:
            continue
        title   = inc.get("title") or inc.get("name") or ""
        summary = inc.get("summary") or inc.get("description") or ""
        docs.append(f"{title}\n{summary}".strip() or "no content")
        ids.append(inc_id)
        metas.append({
            "severity": str(inc.get("riskScore") or inc.get("severity") or ""),
            "status":   str(inc.get("status") or ""),
            "created":  str(inc.get("created") or inc.get("createdDate") or "")[:19],
        })
    if not docs:
        return 0, "No valid incidents to store."
    col.upsert(documents=docs, ids=ids, metadatas=metas)
    return len(docs), f"Synced {len(docs)} incidents."

def chroma_search(query: str, n: int = 5) -> list:
    col = st.session_state.chroma_col
    if col is None or not query.strip() or col.count() == 0:
        return []
    try:
        res = col.query(query_texts=[query], n_results=min(n, col.count()))
        return [{"id": res["ids"][0][i], "text": doc,
                 "score": round((1 - res["distances"][0][i]) * 100, 1),
                 "meta": res["metadatas"][0][i]}
                for i, doc in enumerate(res["documents"][0])]
    except Exception as e:
        return [{"id":"err","text":str(e),"score":0,"meta":{}}]

def chat_respond(user_msg: str, incident: Optional[dict] = None,
                 parsed_context: Optional[dict] = None) -> str:
    """
    ── SOC TRIAGE AGENT (LangChain) ─────────────────────────
    Powered by soc_triage_agent.py.
    Trigger words: triage, analyse, analyze, ioc, classify, ticket, investigate
    When triggered with an incident selected, runs the full pipeline:
      Phase 1 → IOC Checklists (Availability / Confidentiality / Integrity)
      Phase 2 → Risk Rating Methodology
      Phase 3 → SOC Classification Template
    Outputs: metakeys_payload + UNC ticket (both queued for downstream agents).
    All other messages fall back to plain SOC analyst Q&A via LangChain.
    ─────────────────────────────────────────────────────────
    """
    return soc_triage_chat_respond(user_msg, incident, llm_config=get_cisco_cfg(),
                                   parsed_context=parsed_context)


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-FETCH on every render
# ══════════════════════════════════════════════════════════════════════════════
maybe_auto_fetch()


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
# ── RBAC role — determined HERE (before the sidebar) so the sidebar can render a
# slimmed view for guests. Driven by the "Analyst mode" toggle rendered further
# down in the sidebar (key="analyst_mode_on") — Streamlit keeps widget-backed
# session_state values across reruns, so reading it here (before the toggle is
# instantiated below) still reflects the click that triggered this rerun.
# NOTE: there is no password/auth behind this — it's a single-operator desktop
# tool, so the point of the toggle is to make "guest" mean something (hide/
# disable the destructive Data Pipeline tools) rather than to gate a login.
st.session_state.setdefault("analyst_mode_on", False)
st.session_state.user_role = "developer" if st.session_state.analyst_mode_on else "guest"
_is_dev = st.session_state.user_role == "developer"

# Redirect set by the "Triage" button flow (jump straight to AI Agents
# after selecting/uploading an incident) -- previously lived inside the
# sidebar block; kept as top-level code so it still runs before routing.
if st.session_state.get("jump_to_ask_tab", False):
    st.session_state.nav_page = "Ask a Question"
    st.session_state.jump_to_ask_tab = False



# ══════════════════════════════════════════════════════════════════════════════
# AUTO-RERUN scheduler
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.nw_verified and st.session_state.last_fetch:
    elapsed   = (datetime.now() - st.session_state.last_fetch).total_seconds()
    remaining = max(REFRESH_INTERVAL - elapsed, 0)
    if remaining <= 1:
        time.sleep(1)
        st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# HEADER METRICS
# ══════════════════════════════════════════════════════════════════════════════
incidents = st.session_state.incidents
total     = len(incidents)
active    = len([i for i in incidents
                 if str(i.get("status","")).upper()
                 not in ("CLOSED","RESOLVED","REMEDIATED")])
by_sev    = Counter(normalise_sev(i) for i in incidents)
vectors   = st.session_state.chroma_col.count() if st.session_state.chroma_col else 0
last_sync = st.session_state.last_fetch.strftime("%H:%M:%S") if st.session_state.last_fetch else "—"
db_total  = db_stats()["total"]

# Aegis design-system components (ui_components.py) — page title + stat cards
import importlib
import ui_components as _ui
importlib.reload(_ui)
st.markdown(_ui.COMPONENT_CSS, unsafe_allow_html=True)


def _build_case_findings(inc: dict):
    """Real 'key findings' for an incident: the distilled alert behaviours
    (alertMeta.AlertTitles) + the elevated unified-verdict signals. Returns
    (findings, verdict). Deterministic, guarded — never raises."""
    findings: list[dict] = []
    try:
        from triage_verdict import aggregate_verdict
        v = aggregate_verdict(inc)
    except Exception:
        v = {"available": False, "signals": []}
    _conf = {3: "high", 2: "elevated", 1: "moderate", 0: "low"}
    _kw = [("hta", ""), ("c2", ""), ("command", ""), ("exfil", ""),
           ("autorun", ""), ("credential", ""), ("powershell", "⌘"),
           ("lateral", "↔"), ("ransom", ""), ("phish", ""), ("beacon", "")]
    am = inc.get("alertMeta") or {}
    for t in list(dict.fromkeys(am.get("AlertTitles") or []))[:6]:
        tl = str(t).lower()
        icon = next((e for k, e in _kw if k in tl), "")
        findings.append({"icon": icon, "title": str(t)[:72],
                         "desc": "Observed alert behaviour", "confidence": ""})
    if v.get("available"):
        _si = {"base severity": "", "asset criticality": "",
               "internal IOC correlation": "", "external threat intel": ""}
        for s in sorted(v.get("signals", []), key=lambda s: -s.get("level", 0)):
            if s.get("error") or s.get("absent") or s.get("level", 0) == 0:
                continue
            findings.append({"icon": _si.get(s["name"], "•"),
                             "title": f"{s['name'].title()} — {s['label']}",
                             "desc": s.get("detail", ""),
                             "confidence": _conf.get(s["level"], "")})
    return findings, v


def _build_case_context(inc: dict, sev: str, status: str, alerts, verdict: dict):
    """Aegis 'case context' grid facts from the incident + unified verdict."""
    am = inc.get("alertMeta") or {}
    host = (am.get("Hostname") or [None])[0] or inc.get("hostname") or "—"
    user = (am.get("User") or am.get("AdUser") or [None])[0] or "—"
    ips = len(am.get("SourceIp") or []) + len(am.get("DestinationIp") or [])
    _tone = {"CRITICAL": "crit", "HIGH": "warn", "MEDIUM": "warn", "LOW": "ok"}
    vlvl = verdict.get("level", "—") if verdict.get("available") else "—"
    return [
        ("Severity", sev.title(), _tone.get(sev, ""),
         "How serious NetWitness rated this incident on its own."),
        ("Unified verdict", str(vlvl), _tone.get(vlvl, ""),
         "This app's combined risk read, blending severity, asset "
         "importance, and known-bad-indicator matches into one rating."),
        ("Host", str(host)),
        ("User", str(user)),
        ("Status", str(status)),
        ("IOC IPs", str(ips),
         "", "IOC = Indicator of Compromise — a technical signal (like an "
             "IP address) known to be associated with malicious activity. "
             "This counts how many were seen in this incident."),
    ]


def _pipeline_worked_ids() -> set:
    """Incident ids already present anywhere in the SOC pipeline DB —
    used to steer the 'next move' pick toward untouched cases. Read-only,
    guarded; a missing/locked DB just means 'nothing worked yet'."""
    ids: set = set()
    try:
        with _pl_con() as c:
            for _s in PIPELINE_STAGES:
                if _s == "workflow_runs":
                    continue
                try:
                    for (i,) in c.execute(
                            f"SELECT DISTINCT incident_id FROM {_s}").fetchall():
                        if i:
                            ids.add(str(i))
                except Exception:
                    continue
    except Exception:
        pass
    return ids


# Most-advanced-first order for _pipeline_stage_map() below.
_STAGE_ORDER_ADVANCING = [
    "alerts_to_triage", "initial_ticket",
    "post_triage_no_investigate", "post_triage_investigate",
    "post_investigation", "pending_ticket_report", "finalized_report",
]

_STAGE_LABEL_SHORT = {
    "alerts_to_triage":           "Triaged",
    "initial_ticket":             "Triaged",
    "post_triage_no_investigate": "Triaged",
    "post_triage_investigate":    "Investigating",
    "post_investigation":         "Findings ready",
    "pending_ticket_report":      "Reporting",
    "finalized_report":           "Reported",
}

# stage key -> (next action, why) — the Incidents table's Next Action column.
_NEXT_ACTION = {
    "alerts_to_triage":           ("Review triage result", "Triaged"),
    "initial_ticket":             ("Review triage result", "Ticket created"),
    "post_triage_no_investigate": ("Review triage result",
                                    "Triaged — no investigation needed"),
    "post_triage_investigate":    ("Awaiting investigation",
                                    "Queued for the investigation agent"),
    "post_investigation":         ("Review findings",
                                    "Investigation complete — analyst review needed"),
    "pending_ticket_report":      ("Awaiting report",
                                    "Reporting agent is generating the report"),
    "finalized_report":           ("Review final report",
                                    "Report ready — Pipeline DB → Finalized Report"),
}


def _pipeline_stage_map() -> dict:
    """incident_id -> most-advanced pipeline stage it's reached, computed
    once per page render (one query per stage table) — NOT once per row,
    which is what a naive per-incident lookup would cost at 200+ rows."""
    stage_map: dict = {}
    try:
        with _pl_con() as c:
            for _s in _STAGE_ORDER_ADVANCING:  # least-advanced first —
                try:                            # a later write below wins.
                    for (i,) in c.execute(
                            f"SELECT DISTINCT incident_id FROM {_s}").fetchall():
                        if i:
                            stage_map[str(i)] = _s
                except Exception:
                    continue
    except Exception:
        pass
    return stage_map


# The 6-stage display order for My Workspace's per-case circular stepper —
# matches the names already used by the Overview page's aggregate pipeline
# (_render_circular_pipeline_section) so both views read as one system.
_CASE_DISPLAY_STAGES = [
    "Parsing", "Triage", "Threat Intelligence Enrichment",
    "Investigation", "Reporting",
]
# internal pipeline stage key -> its position in _CASE_DISPLAY_STAGES
_CASE_STAGE_POSITION = {
    "alerts_to_triage": 1, "post_triage_no_investigate": 1,
    "initial_ticket": 1,
    "post_triage_investigate": 2,
    "post_investigation": 3,
    "pending_ticket_report": 4,
    "finalized_report": 4,
}


def _case_stage_states(inc_id: str, status: str, stage_map: dict,
                       parsing_status: str | None = None) -> list:
    """[{'name', 'state': 'done'|'current'|'queued'}] x6 for one specific
    case's circular workflow stepper (My Workspace) — reuses the same
    stage-position data _pipeline_stage_map() already computed, just reads
    a single case's position out of it instead of aggregating counts."""
    stage = stage_map.get(str(inc_id))
    if str(status or "").upper() in _CLOSED_STATUSES and stage == "finalized_report":
        current_pos = 4
    elif stage in _CASE_STAGE_POSITION:
        current_pos = _CASE_STAGE_POSITION[stage]
    elif str(parsing_status or "").lower() == "complete":
        # Parsing done, nothing triaged yet — Triage is current.
        current_pos = 1
    else:
        # Parsing not yet run/complete — Parsing is the current stage, not
        # an impossible state where every stage is queued.
        current_pos = 0
    out = []
    for i, name in enumerate(_CASE_DISPLAY_STAGES):
        state = "done" if i < current_pos else "current" if i == current_pos else "queued"
        out.append({"name": name, "state": state})
    return out


def _render_case_stage_selector(
    stages: list, selected_stage: str, case_id: str
) -> str:
    """Render the case workflow as native controls.

    The previous workflow used HTML anchors with query parameters.  That made
    selecting a stage a browser navigation, which briefly left/reloaded the
    workspace.  Native buttons keep the selection in Streamlit session state
    and only rerender the detail content below this control.
    """
    _selected_index = next(
        (i for i, stage in enumerate(stages)
         if str(stage.get("name") or "") == selected_stage),
        0,
    )
    _state_rules = []
    for _i, _stage in enumerate(stages, start=1):
        _state = str(_stage.get("state") or "queued").lower()
        if _state == "done":
            _state_rules.append(f"""
            div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_i})
            div.stButton > button {{
                background:#303b83 !important;border-color:#7680ff !important;
                color:#fff !important;box-shadow:0 0 0 6px rgba(111,124,255,.10) !important;
            }}
            div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_i})
            .case-stage-status {{ color:#aeb4ff !important; }}
            div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_i})::after {{
                background:#7778f6 !important;
            }}
            """)
        elif _state == "current":
            _state_rules.append(f"""
            div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_i})
            div.stButton > button {{
                background:#332812 !important;border-color:#9c722c !important;
                color:#ffd36b !important;box-shadow:0 0 0 7px rgba(244,188,95,.09) !important;
            }}
            div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_i})
            .case-stage-status {{ color:#ffd36b !important; }}
            """)
    _selected_child = _selected_index + 1
    st.html(
        f"""
        <style>
        .case-workflow-title {{
            color:#f7f9fd;font-size:1.05rem;font-weight:800;
            letter-spacing:-.01em;margin:0 0 5px;
        }}
        .case-workflow-subtitle {{
            color:#91acd5;font-size:.76rem;margin:0 0 18px;
        }}
        div.st-key-case_stage_selector {{
            margin:0 0 12px;padding:0 4px 4px;
        }}
        div.st-key-case_stage_selector [data-testid="stHorizontalBlock"] {{
            gap:0 !important;
            align-items: stretch !important;
            overflow-x:auto !important;
            padding:0 0 2px !important;
        }}
        div.st-key-case_stage_selector [data-testid="stColumn"] {{
            position:relative !important;z-index:1;
            min-width:130px !important;min-height:112px !important;
            padding:7px 10px 10px !important;
            border:1px solid transparent !important;border-radius:12px !important;
            text-align:center !important;
        }}
        div.st-key-case_stage_selector [data-testid="stColumn"]:not(:last-child)::after {{
            content:"";position:absolute;z-index:0;top:27px;
            left:calc(50% + 20px);width:calc(100% - 40px);
            height:2px;background:#2a354a;
        }}
        div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_selected_child}) {{
            background:#1c263b !important;border-color:#485a79 !important;
        }}
        div.st-key-case_stage_selector [data-testid="stColumn"]:nth-child({_selected_child})::after {{
            left:100% !important;width:calc(50% - 20px) !important;
        }}
        div.st-key-case_stage_selector [data-testid="stColumn"]
        [data-testid="stElementContainer"]:has(div.stButton) {{
            width:100% !important;
        }}
        div.st-key-case_stage_selector div.stButton {{
            position:relative;z-index:2;
            width:100% !important;
            display:flex !important;justify-content:center !important;
        }}
        div.st-key-case_stage_selector div.stButton > button {{
            width:40px !important;min-width:40px !important;max-width:40px !important;
            height:40px !important;min-height:40px !important;padding:0 !important;
            border-radius:50% !important;background:#172137 !important;
            border:1px solid #32415b !important;color:#8793a8 !important;
            font-size:.78rem !important;font-weight:800 !important;line-height:1 !important;
            box-shadow:none !important;
        }}
        div.st-key-case_stage_selector div.stButton > button:hover {{
            border-color:#7182ff !important;color:#fff !important;
            transform:translateY(-1px);
        }}
        div.st-key-case_stage_selector .case-stage-name {{
            color:#f2f4fa !important;font-size:.76rem !important;
            font-weight:750 !important;line-height:1.22 !important;
            text-align:center !important;margin:11px 0 0 !important;
        }}
        div.st-key-case_stage_selector .case-stage-status {{
            color:#8e99ad !important;font-size:.69rem !important;
            line-height:1.2 !important;text-align:center !important;
            margin:5px 0 0 !important;
        }}
        {"".join(_state_rules)}
        </style>
        """,
    )
    st.markdown(
        '<div class="case-workflow-title">Case workflow</div>'
        '<div class="case-workflow-subtitle">'
        'Select a stage to inspect its work, outputs and actions</div>',
        unsafe_allow_html=True,
    )
    with st.container(key="case_stage_selector"):
        columns = st.columns(len(stages))
        for index, (column, stage) in enumerate(zip(columns, stages)):
            name = str(stage.get("name") or "")
            state = str(stage.get("state") or "queued").lower()
            if state == "done":
                marker, status_label = "✓", "Complete"
            elif state == "current":
                marker, status_label = str(index + 1), "Current stage"
            else:
                marker, status_label = str(index + 1), "Queued"
            with column:
                if st.button(
                    marker,
                    key=f"case_stage_{case_id}_{index}",
                    type="secondary",
                    help=f"Show {name} details below",
                ):
                    st.query_params.clear()
                    st.session_state.case_selected_stage = name
                    st.session_state.case_selected_stage_id = str(case_id)
                    st.rerun()
                st.markdown(
                    f'<div class="case-stage-name">{_esc_html(name)}</div>'
                    f'<div class="case-stage-status">{_esc_html(status_label)}</div>',
                    unsafe_allow_html=True,
                )
    return selected_stage


def _next_action_for(inc_id: str, status: str, stage_map: dict) -> tuple:
    """(action, detail) for the Incidents table's Next Action column.
    Deliberately never invents an action with no matching control on the
    row — e.g. NOT "Assign analyst", since this app has no assign button."""
    if str(status or "").upper() in _CLOSED_STATUSES:
        return "No action needed", "Closed"
    stage = stage_map.get(str(inc_id))
    if stage:
        return _NEXT_ACTION.get(stage, ("Review", "In pipeline"))
    return "Triage this case", "Not yet triaged"


def _esc_html(s) -> str:
    return (str(s if s is not None else "")
            .replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


_SEV_BORDER = {"CRITICAL": "#ff6e7c", "HIGH": "#f4bc5f",
               "MEDIUM": "#c9a6f7", "LOW": "#43d28c"}


def _render_case_table(rows: list, *, key_prefix: str, heading: str = "Open cases",
                       subheading: str = "") -> None:
    """Shared dense case table — used by both Overview (a capped, ranked
    subset) and All Cases (the full filtered archive), so the row-rendering
    logic exists exactly once. Clicking a case name navigates to My
    Workspace with that case loaded — this replaces the old per-row Triage
    button; triage is now triggered from inside My Workspace itself via its
    primary action button (same _next_action_for() label), not from the
    list."""
    if not rows:
        st.markdown(
            '<div style="text-align:center;padding:50px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '● NO CASES<br>'
            '<span style="font-size:0.62rem">'
            'Incidents are saved automatically on every fetch</span></div>',
            unsafe_allow_html=True,
        )
        return

    # Computed once for the whole table — NOT per row (see
    # _pipeline_stage_map's own docstring for why that matters at scale).
    stage_map = _pipeline_stage_map()
    _col_ratios = [1.0, 3.25, 0.8, 1.35, 0.95, 1.65, 0.55]
    _table_key = f"{key_prefix}_table"

    # Style-only st.html calls do not create a visible layout element. Using
    # st.markdown here leaves an empty element in the table's vertical stack,
    # which Streamlit spaces like real content.
    st.html(f"""
    <style>
    div.st-key-{_table_key} {{
        background: #0b1422 !important;
        border: 1px solid #1d2d42 !important;
        border-radius: 10px !important;
        padding: 16px 18px 8px !important;
        overflow: hidden !important;
    }}
    div.st-key-{_table_key} [data-testid="stVerticalBlock"] {{
        gap: 0 !important;
    }}
    div.st-key-{_table_key} [data-testid="stHorizontalBlock"] {{
        gap: 0.8rem !important;
    }}
    div.st-key-{_table_key} [data-testid="stMarkdownContainer"] p {{
        margin: 0 !important;
    }}
    div.st-key-{_table_key} [data-testid="stColumn"] {{
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        text-align: center !important;
    }}
    div.st-key-{_table_key} [data-testid="stColumn"]
    [data-testid="stMarkdownContainer"] {{
        width: 100% !important;
        text-align: center !important;
    }}
    div.st-key-{_table_key} [data-testid="stColumn"]
    [data-testid="stMarkdownContainer"] > * {{
        width: 100% !important;
        margin-left: auto !important;
        margin-right: auto !important;
        font-size: 14px !important;
        text-align: center !important;
    }}
    </style>
    """)

    with st.container(key=_table_key):
        st.markdown(
            f'<div style="font-size:1.1rem;font-weight:750;color:var(--text);'
            f'letter-spacing:-.01em">'
            f'{_esc_html(heading)}</div>'
            + (f'<div style="font-size:0.8rem;color:var(--sub);margin:2px 0 14px">'
               f'{_esc_html(subheading)}</div>' if subheading else
               '<div style="margin-bottom:12px"></div>'),
            unsafe_allow_html=True)
        _hdr = st.columns(_col_ratios)
        for _hc, _htxt in zip(_hdr, ["CASE ID", "CASE", "SEVERITY",
                                     "CURRENT STAGE", "OWNER",
                                     "NEXT ACTION", "ACTIONS"]):
            _hc.markdown(
                f'<div style="font-family:var(--mono);font-size:0.68rem;'
                f'color:#7895bd;text-transform:uppercase;'
                f'letter-spacing:.08em;font-weight:750;white-space:nowrap">'
                f'{_htxt}</div>',
                unsafe_allow_html=True)
        st.markdown('<hr style="border:0;border-top:1px solid #21334a;'
                   'margin:7px 0 0">', unsafe_allow_html=True)

        for row in rows:
            try:
                _raw = _json.loads(row.get("raw_json") or "{}")
            except Exception:
                _raw = {}
            inc = _raw if isinstance(_raw, dict) and _raw.get("id") else row

            sev      = normalise_sev(inc) if ("severity" in inc or "riskScore" in inc) else row.get("severity", "LOW")
            inc_id   = str(inc.get("id") or inc.get("incidentId") or row.get("id") or "—")
            title    = inc.get("title") or inc.get("name") or row.get("title") or "Untitled"
            status   = str(inc.get("status") or row.get("status") or "—")
            assignee = inc.get("assignee") or row.get("assignee") or "Unassigned"

            _border_col  = _SEV_BORDER.get(sev.upper(), "#43d28c")
            _stage_key   = stage_map.get(inc_id)
            _stage_label = _STAGE_LABEL_SHORT.get(_stage_key, "Not triaged")
            _action, _action_detail = _next_action_for(inc_id, status, stage_map)
            _unassigned = assignee in ("Unassigned", "", None)
            _row_key = f"{key_prefix}_{inc_id}"

            # Keep the per-row style out of the layout flow. A style-only
            # st.markdown call still occupies an stElementContainer and was
            # the source of the large blank band above every incident.
            st.html(f"""
            <style>
            div.st-key-{_row_key} {{
                border: none !important;
                border-bottom: 1px solid #1b2b3f !important;
                border-radius: 0 !important;
                background: transparent !important;
                padding: 0 !important;
                /* Cancel Streamlit's reserved 1rem slot before each nested
                   container so the hover surface starts at the separator. */
                margin: -1rem 0 0 !important;
                min-height: 49px !important;
                transition: background-color 160ms ease,
                            box-shadow 160ms ease !important;
            }}
            div.st-key-{_row_key}:hover {{
                background: #101c2f !important;
                box-shadow: 0 8px 22px rgba(0, 0, 0, .28),
                            0 0 0 1px rgba(117, 165, 255, .14) !important;
                position: relative;
                z-index: 1;
            }}
            div.st-key-{_row_key} [data-testid="stColumn"],
            div.st-key-{_row_key} [data-testid="stVerticalBlock"],
            div.st-key-{_row_key} [data-testid="stHorizontalBlock"] {{
                background: transparent !important;
            }}
            div.st-key-{_row_key} > [data-testid="stVerticalBlock"],
            div.st-key-{_row_key} [data-testid="stHorizontalBlock"] {{
                min-height: 48px !important;
                height: 100% !important;
                gap: 0 !important;
            }}
            div.st-key-{_row_key} [data-testid="stHorizontalBlock"] {{
                align-items: stretch !important;
                column-gap: 0.8rem !important;
            }}
            div.st-key-{_row_key} [data-testid="stColumn"] {{
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                min-height: 48px !important;
            }}
            div.st-key-{_row_key} [data-testid="stColumn"] > [data-testid="stVerticalBlock"] {{
                min-height: 48px !important;
                height: 48px !important;
                width: 100% !important;
                justify-content: center !important;
            }}
            div.st-key-{_row_key} [data-testid="stElementContainer"] {{
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                width: 100% !important;
                margin-bottom: 0 !important;
            }}
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"],
            div.st-key-{_row_key} div.stButton,
            div.st-key-{_row_key} [data-testid="stPopover"] {{
                width: 100% !important;
            }}
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"] {{
                min-height: 48px !important;
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                line-height: 1.25 !important;
            }}
            div.st-key-{_row_key} div.stButton {{
                height: 48px !important;
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
            }}
            div.st-key-{_row_key} div.stButton > button p {{
                width: 100% !important;
                margin: 0 !important;
                text-align: center !important;
                line-height: 1.25 !important;
            }}
            div.st-key-{_row_key} div.stButton > button,
            div.st-key-{_row_key} button[data-testid="stBaseButton-secondary"] {{
                background: transparent !important;
                border: 0 !important;
                border-radius: 0 !important;
                text-align: center !important;
                justify-content: center !important;
                padding: 3px 11px !important;
                color: var(--text) !important; font-weight: 600 !important;
                font-size: 14px !important; box-shadow: none !important;
                min-height: 48px !important;
                height: 48px !important;
                position: relative !important;
                line-height: 1.25 !important;
                cursor: pointer !important;
                transition: color 140ms ease, background-color 140ms ease,
                            transform 140ms ease, text-shadow 140ms ease !important;
            }}
            div.st-key-{_row_key} div.stButton > button::before {{
                content: "" !important;
                position: absolute !important;
                left: 0 !important;
                top: 3px !important;
                bottom: 3px !important;
                width: 3px !important;
                border-radius: 3px !important;
                background: {_border_col} !important;
            }}
            div.st-key-{_row_key} div.stButton > button:hover {{
                color: var(--accent) !important;
                background: #122036 !important;
                text-shadow: 0 0 14px rgba(117, 165, 255, .32) !important;
                transform: none !important;
                box-shadow: none !important;
            }}
            div.st-key-{_row_key} [data-testid="stPopover"] button {{
                min-height: 28px !important;
                height: 28px !important;
                padding: 0 8px !important;
                border: 1px solid #2a3d57 !important;
                border-radius: 6px !important;
                color: #91a8c7 !important;
                background: #0d1828 !important;
                justify-content: center !important;
            }}
            div.st-key-{_row_key} [data-testid="stColumn"] {{
                text-align: center !important;
            }}
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"],
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"] > *,
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"] span,
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"] div {{
                text-align: center !important;
                font-size: 14px !important;
            }}
            div.st-key-{_row_key} [data-testid="stMarkdownContainer"] > p {{
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                width: 100% !important;
                margin: 0 !important;
            }}
            div.st-key-{_row_key} div.stButton > button {{
                width: 100% !important;
            }}
            div.st-key-{_row_key} div.stButton > button > div,
            div.st-key-{_row_key} div.stButton > button p {{
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                width: 100% !important;
                text-align: center !important;
            }}
            div.st-key-{_row_key} [data-testid="stPopover"] button::before {{
                display: none !important;
            }}
            div.st-key-{_row_key} .case-id-cell {{
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                width: 100% !important;
                min-height: 48px !important;
                margin: 0 !important;
                text-align: center !important;
            }}
            </style>
            """)

            with st.container(key=_row_key):
                _c = st.columns(_col_ratios, vertical_alignment="center")
                _c[0].markdown(
                    f'<div class="case-id-cell" style="font-family:var(--mono);'
                    f'font-size:0.8rem;color:#7ea6d8">'
                    f'{_esc_html(inc_id[:16])}</div>',
                    unsafe_allow_html=True)
                with _c[1]:
                    if st.button(title[:56], key=f"{_row_key}_open",
                                use_container_width=True,
                                help="Open this case in My Workspace"):
                        st.session_state.selected_case_id = inc_id
                        st.session_state.nav_page = "My Workspace"
                        st.rerun()
                _c[2].markdown(_ui.pill(sev, _ui.sev_class(sev)), unsafe_allow_html=True)
                _c[3].markdown(_ui.pill(_stage_label, "stage"), unsafe_allow_html=True)
                _c[4].markdown(
                    f'<span style="font-size:0.85rem;'
                    f'color:{"var(--warn)" if _unassigned else "var(--text)"}">'
                    f'{_esc_html(assignee)}</span>', unsafe_allow_html=True)
                _c[5].markdown(
                    f'<div style="font-size:0.85rem;color:'
                    f'{"var(--warn)" if _action not in ("Awaiting report", "No action needed") else "var(--text)"};'
                    f'font-weight:700;line-height:1.3">{_esc_html(_action)}</div>',
                    unsafe_allow_html=True)
                with _c[6]:
                    with st.popover("⋯", use_container_width=True):
                        do_json = st.button("View Raw JSON", key=f"{_row_key}_json",
                                            use_container_width=True)

            if do_json:
                with st.expander(f"JSON — {inc_id}", expanded=True):
                    st.json(inc)


def _pick_next_move(incs: list):
    """Deterministic 'most urgent case to pick up next': unworked incidents
    first, then severity band, then riskScore, then distilled-behaviour
    richness, then newest. Cheap (no skills, no network) — safe to run every
    rerun. Returns (incident, meta) or (None, {})."""
    if not incs:
        return None, {}
    worked = _pipeline_worked_ids()
    _rank = {"CRITICAL": 3, "HIGH": 2, "MEDIUM": 1, "LOW": 0}

    def _risk(i):
        try:
            return int(i.get("riskScore") or 0)
        except (TypeError, ValueError):
            return 0

    def _behaviours(i):
        return len((i.get("alertMeta") or {}).get("AlertTitles") or [])

    def _id(i):
        return str(i.get("id") or i.get("incidentId") or "")

    # Deterministic ordering: id asc → created desc → main urgency key.
    # The stable sorts make created/id pure tie-breaks for the urgency key.
    ordered = sorted(incs, key=_id)
    ordered = sorted(ordered, key=lambda i: str(i.get("created") or ""), reverse=True)
    ordered = sorted(ordered, key=lambda i: (
        1 if _id(i) in worked else 0,
        -_rank.get(normalise_sev(i), 0),
        -_risk(i),
        -_behaviours(i),
    ))
    top = ordered[0]
    return top, {
        "already_worked": _id(top) in worked,
        "sev": normalise_sev(top),
        "risk": _risk(top),
        "behaviours": _behaviours(top),
    }

def _render_overview_header():
    st.markdown(_ui.page_title(
        "Operations overview",
        f"{active:,} active · {total:,} in session · last sync {last_sync}"),
        unsafe_allow_html=True)
    return

    try:
        _nm, _nm_meta = _pick_next_move(incidents)
        if _nm is not None:
            _nm_id = str(_nm.get("id") or _nm.get("incidentId") or "?")
            _nm_title = str(_nm.get("title") or _nm.get("name") or "Untitled incident")[:90]
            _vc = st.session_state.get("_next_move_verdict") or {}
            if _vc.get("id") != _nm_id:
                try:
                    from triage_verdict import aggregate_verdict as _agg
                    _vc = {"id": _nm_id, "v": _agg(_nm)}
                except Exception:
                    _vc = {"id": _nm_id, "v": {"available": False}}
                st.session_state._next_move_verdict = _vc
            _v = _vc.get("v") or {}
            _why = []
            if _v.get("available"):
                _why.append(f"Unified verdict {_v.get('level')} — {_v.get('action')}")
            else:
                _why.append(f"Severity {_nm_meta.get('sev', '—').title()}")
            if _nm_meta.get("risk"):
                _why.append(f"risk {_nm_meta['risk']}")
            if _nm_meta.get("behaviours"):
                _why.append(f"{_nm_meta['behaviours']} observed behaviour(s)")
            _why.append("already in pipeline" if _nm_meta.get("already_worked")
                        else "not yet worked")
            _hot = (_v.get("level") in ("CRITICAL", "HIGH")
                    or _nm_meta.get("sev") in ("CRITICAL", "HIGH"))
            _hero_border = "#633645" if _hot else "#33407a"
            _hero_bg = "linear-gradient(105deg, #351d2acc, #111b2c 58%)" if _hot else "linear-gradient(105deg, #1a2350cc, #111b2c 58%)"
            _hero_eyebrow = "#ff939d" if _hot else "#aeb7ff"

            st.markdown(f"""
            <style>
            div.st-key-hero_container,
            div[data-testid="stVerticalBlockBorderWrapper"].st-key-hero_container {{
                border: 2px solid {_hero_border} !important;
                background: {_hero_bg} !important;
                box-shadow: 0 24px 60px #0009 !important;
                border-radius: 20px !important;
                padding: 36px 44px 40px 44px !important;
                margin-bottom: 28px !important;
            }}
            div.st-key-hero_container [data-testid="stColumn"],
            div.st-key-hero_container [data-testid="stVerticalBlock"],
            div.st-key-hero_container [data-testid="stHorizontalBlock"] {{
                background: transparent !important;
            }}
            div.st-key-hero_container button {{
                font-size: 1.12rem !important;
                font-weight: 800 !important;
                min-height: 58px !important;
                border-radius: 12px !important;
                letter-spacing: 0.03em !important;
            }}
            </style>
            """, unsafe_allow_html=True)

            with st.container(key="hero_container", border=True):
                h_col1, h_col2 = st.columns([3.4, 1.6], vertical_alignment="center")
                with h_col1:
                    st.markdown(f'''
                    <div class="ag-hero-body" style="padding-bottom:6px;">
                        <div class="e" style="color:{_hero_eyebrow};font-size:0.75rem;font-weight:900;letter-spacing:.14em;text-transform:uppercase;margin-bottom:4px;">YOUR NEXT MOVE</div>
                        <h2 style="margin:6px 0 10px;font-size:2.3rem !important;font-weight:900 !important;color:#ffffff !important;line-height:1.25 !important;letter-spacing:-0.01em;"><strong>{_nm_id} — {_nm_title}</strong></h2>
                        <p style="margin:0;color:#b4c3d6;font-size:0.85rem;line-height:1.55;">Why this case: {" · ".join(_why)}</p>
                    </div>
                    ''', unsafe_allow_html=True)
                with h_col2:
                    if st.button(f"Triage {_nm_id} now", key="hero_triage", use_container_width=True):
                        st.session_state.chat_incident       = _nm
                        st.session_state.pending_auto_triage = True
                        st.session_state.nav_page            = "Ask a Question"
                        st.rerun()
    except Exception:
        pass


def _render_circular_pipeline_section():
    try:
        _pstages = [
            ("Triage", "alerts_to_triage"),
            ("Investigation", "post_triage_investigate"),
            ("Findings", "post_investigation"),
            ("Ticketing", "initial_ticket"),
            ("Reporting", "pending_ticket_report"),
            ("Finalized", "finalized_report"),
        ]
        _steps = []
        for _nm, _tbl in _pstages:
            _c = pipeline_count(_tbl)
            _steps.append({"name": _nm, "count": _c,
                           "label": (f"{_c} in stage" if _c else "empty")})

        # Real average, not a placeholder — computed from the same
        # workflow_runs audit rows _workflow_worker() already writes
        # (duration_seconds is embedded in each row's raw_json).
        _avg_cycle = "—"
        try:
            _recent_runs = pipeline_load("workflow_runs", limit=20)
            _durations = []
            for _rr in _recent_runs:
                try:
                    _rrj = _json.loads(_rr.get("raw_json") or "{}")
                    _d = _rrj.get("duration_seconds")
                    if _d is not None:
                        _durations.append(float(_d))
                except Exception:
                    continue
            if _durations:
                _avg_s = int(round(sum(_durations) / len(_durations)))
                _avg_cycle = f"{_avg_s // 60}m {_avg_s % 60}s"
        except Exception:
            pass

        st.markdown(f'''
        <div style="margin:26px 0 16px;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:16px;">
            <div style="font-size:1.3rem;font-weight:800;color:#ffffff;letter-spacing:-0.01em;">System cases pipeline</div>
            <div style="display:flex;align-items:center;gap:12px;flex-wrap:wrap;">
                <div style="background:#0c1626;border:1px solid #1e2d42;padding:6px 14px;border-radius:8px;font-family:var(--mono);font-size:0.75rem;color:#a0aec0;display:flex;align-items:center;gap:8px;">
                    <span style="color:#718096;font-weight:600;letter-spacing:0.5px;">AVG CYCLE TIME</span>
                    <strong style="color:#43d28c;font-size:0.85rem;">{_avg_cycle}</strong>
                </div>
            </div>
        </div>
        ''', unsafe_allow_html=True)
        st.markdown(_ui.circular_pipeline(_steps), unsafe_allow_html=True)
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════════════
# PAGE ROUTING
# ══════════════════════════════════════════════════════════════════════════════
# Workflow steps are normal links so they remain accessible and can be opened in
# a new tab. Rehydrate both pieces of workspace context from their URL instead
# of allowing a link-triggered page reload to fall back to Overview.
_linked_case_id = st.query_params.get("case_id")
_linked_case_stage = st.query_params.get("case_stage")
if _linked_case_id and _linked_case_stage in _CASE_DISPLAY_STAGES:
    st.session_state.selected_case_id = str(_linked_case_id)
    st.session_state.nav_page = "My Workspace"

active_page = st.session_state.get("nav_page", "Overview")

def _render_top_nav(active_page: str) -> None:
    """Top navigation bar -- replaces the old left sidebar entirely. Real
    st.button widgets drive page routing (raw HTML can't trigger reruns
    on click)."""
    _tabs = [
        ("Overview", "Overview"),
        ("My Workspace", "My Workspace"),
        ("All Cases", "All Cases"),
    ]
    _page_heading = {
        "Overview": "Operations overview",
        "My Workspace": "My Workspace",
        "All Cases": "All Cases",
    }.get(active_page, active_page)
    _logo_path = (
        Path(__file__).parent
        / "soc_reporting_agent"
        / "dashboard"
        / "assets"
        / "aegis-logo-sidebar.png"
    )
    try:
        _logo_src = (
            "data:image/png;base64,"
            + base64.b64encode(_logo_path.read_bytes()).decode("ascii")
        )
    except OSError:
        _logo_src = ""
    st.markdown("""
    <style>
    div.st-key-top_nav, div[data-testid="stVerticalBlockBorderWrapper"].st-key-top_nav {
        border: 1px solid #17243a !important;
        border-radius: 10px !important; padding: 8px 10px 7px !important;
        background: #050b14 !important;
        box-shadow: 0 10px 28px rgba(0, 0, 0, .22) !important;
        margin-bottom: 2px !important;
    }
    div.st-key-top_nav [data-testid="stColumn"] { background: transparent !important; }
    div.st-key-top_nav [data-testid="stColumn"]:first-child,
    div.st-key-top_nav [data-testid="stColumn"]:first-child [data-testid="stElementContainer"],
    div.st-key-top_nav [data-testid="stColumn"]:first-child .stMarkdown,
    div.st-key-top_nav [data-testid="stColumn"]:first-child [data-testid="stMarkdownContainer"] {
        display: flex !important;
        align-items: center !important;
    }
    div.st-key-top_nav [data-testid="stColumn"]:first-child [data-testid="stElementContainer"],
    div.st-key-top_nav [data-testid="stColumn"]:first-child .stMarkdown,
    div.st-key-top_nav [data-testid="stColumn"]:first-child [data-testid="stMarkdownContainer"] {
        width: 100%;
        height: 48px !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    .top-nav-identity {
        width: 100%; height: 48px; display: flex; align-items: center; gap: 12px;
        white-space: nowrap; margin: 0 !important;
    }
    .top-nav-identity img {
        display: block; width: 94px; height: auto; object-fit: contain;
    }
    .top-nav-page {
        min-width: 0; height: 48px; border-left: 1px solid #26344a; padding-left: 12px;
        display: flex; flex-direction: column; justify-content: center;
    }
    .top-nav-page-kicker {
        color: #8290a7; font-size: 9px; font-weight: 500;
        line-height: 1; text-transform: uppercase; margin-bottom: 7px;
    }
    .top-nav-page-heading {
        color: #f5f7fb; font-size: 14px; font-weight: 800;
        line-height: 1; letter-spacing: -.01em;
    }
    div.st-key-top_nav div.stButton > button {
        background: #08111f !important; border: 1px solid #1d2b42 !important;
        color: var(--muted) !important; font-weight: 600 !important;
    }
    div.st-key-top_nav div.stButton > button:hover {
        background: #0d1828 !important; border-color: #2a3b56 !important;
        color: #f3f6fb !important;
    }
    div.st-key-top_nav div.stButton > button[kind="primary"] {
        background: #3478f6 !important;
        border: 1px solid #4385ff !important; color: #ffffff !important;
        box-shadow: 0 4px 12px rgba(52,120,246,.22) !important;
    }
    div.st-key-top_nav [data-testid="stTextInput"] input {
        background: #08111f !important; border-color: #1d2b42 !important;
        color: #f3f6fb !important;
    }
    </style>
    """, unsafe_allow_html=True)
    with st.container(key="top_nav"):
        _ratios = [2.35] + [0.95] * len(_tabs) + [2.0, 1.1]
        _cols = st.columns(_ratios, vertical_alignment="center")
        with _cols[0]:
            st.markdown(
                f"""
                <div class="top-nav-identity">
                    <img src="{_logo_src}" alt="Aegis">
                    <div class="top-nav-page">
                        <div class="top-nav-page-kicker">Dashboard</div>
                        <div class="top-nav-page-heading">{_esc_html(_page_heading)}</div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        for _i, (_pid, _label) in enumerate(_tabs):
            with _cols[_i + 1]:
                if st.button(_label, key=f"topnav_{_pid}", use_container_width=True,
                             type="primary" if active_page == _pid else "secondary"):
                    st.query_params.clear()
                    st.session_state.nav_page = _pid
                    st.rerun()
        _search_col, _profile_col = _cols[-2], _cols[-1]
        with _search_col:
            _q = st.text_input("Search cases", key="topnav_search",
                               placeholder="Search cases by title or ID...",
                               label_visibility="collapsed")
            if _q.strip() and _q.strip() != st.session_state.get("_topnav_search_prev", ""):
                st.session_state._topnav_search_prev = _q.strip()
                st.session_state.hist_search_override = _q.strip()
                st.session_state.nav_page = "All Cases"
                st.rerun()
        with _profile_col:
            _who = (st.session_state.get("analyst_display_name") or "").strip() or "Guest"
            if st.button(f"⚙ {_who[:10]}", key="topnav_profile", use_container_width=True,
                        type="primary" if active_page == "Settings" else "secondary"):
                st.session_state.nav_page = "Settings"
                st.rerun()


_render_top_nav(active_page)


SEV_COLORS = {
    "CRITICAL": "#FF3B3B",
    "HIGH":     "#FF7700",
    "MEDIUM":   "#FFB700",
    "LOW":      "#0AF0A0",
}
STATUS_COLORS = {
    "NEW":         "#00D4FF",
    "ASSIGNED":    "#FFB700",
    "IN_PROGRESS": "#FF7700",
    "CLOSED":      "#3A607A",
    "RESOLVED":    "#0AF0A0",
    "REMEDIATED":  "#0AF0A0",
}


# ─────────────────────────────────────────────────────────────
# PAGE 1 — OVERVIEW
# ─────────────────────────────────────────────────────────────
if active_page == "Overview":
    _render_overview_header()
    _render_circular_pipeline_section()

    # ── Open cases — moved here from the old standalone Incidents page.
    # Capped to the 15 highest-priority matching cases (severity, then
    # unassigned, then newest). The compact controls intentionally mirror
    # the three fields analysts scan in the table below.
    st.markdown("<div style='margin-top:10px'></div>", unsafe_allow_html=True)
    _sev_rank = {"CRITICAL": 3, "HIGH": 2, "MEDIUM": 1, "LOW": 0}
    _ov_pool = db_load_incidents(limit=500)
    _ov_pool.sort(key=lambda r: str(r.get("created") or ""), reverse=True)
    _ov_pool.sort(key=lambda r: (
        -_sev_rank.get(str(r.get("severity", "")).upper(), 0),
        0 if (r.get("assignee") in (None, "", "Unassigned")) else 1,
    ))

    _ov_stage_map = _pipeline_stage_map()
    _ov_stage_groups = {
        "Parsing": {None},
        "Triage": {
            "alerts_to_triage", "initial_ticket",
            "post_triage_no_investigate",
        },
        "Threat Intelligence Enrichment": {"post_triage_investigate"},
        "Investigation": {"post_investigation"},
        "Reporting": {"pending_ticket_report", "finalized_report"},
    }
    _ov_owners = sorted({
        str(r.get("assignee") or "Unassigned")
        for r in _ov_pool
    }, key=lambda owner: (owner != "Unassigned", owner.casefold()))

    st.html("""
    <style>
    div.st-key-ovw_filters [data-testid="stHorizontalBlock"] {
        gap: 0.6rem !important;
        align-items: center !important;
    }
    div.st-key-ovw_filters [data-testid="stSelectbox"] > div > div {
        min-height: 36px !important;
        background: #0b1422 !important;
        border-color: #21334a !important;
        border-radius: 7px !important;
        font-size: 0.78rem !important;
    }
    </style>
    """)
    with st.container(key="ovw_filters"):
        _ov_spacer, _ov_f1, _ov_f2, _ov_f3 = st.columns(
            [3.2, 1.15, 1.75, 1.15], vertical_alignment="center")
        _ov_severity = _ov_f1.selectbox(
            "Severity",
            ["All severities", "Critical", "High", "Medium", "Low"],
            key="ovw_severity_filter", label_visibility="collapsed")
        _ov_stage = _ov_f2.selectbox(
            "Stage",
            ["All stages", "Parsing", "Triage",
             "Threat Intelligence Enrichment", "Investigation", "Reporting"],
            key="ovw_stage_filter", label_visibility="collapsed")
        _ov_owner = _ov_f3.selectbox(
            "Owner", ["All owners"] + _ov_owners,
            key="ovw_owner_filter", label_visibility="collapsed")

    if _ov_severity != "All severities":
        _ov_pool = [
            r for r in _ov_pool
            if str(r.get("severity") or "").upper() == _ov_severity.upper()
        ]
    if _ov_stage != "All stages":
        _ov_allowed_stages = _ov_stage_groups[_ov_stage]
        _ov_pool = [
            r for r in _ov_pool
            if _ov_stage_map.get(str(r.get("id"))) in _ov_allowed_stages
        ]
    if _ov_owner != "All owners":
        _ov_pool = [
            r for r in _ov_pool
            if str(r.get("assignee") or "Unassigned") == _ov_owner
        ]

    _render_case_table(
        _ov_pool[:15], key_prefix="ovwcase", heading="Open cases",
        subheading="Your highest-priority cases right now, ranked by "
                   "severity and ownership.")


# ─────────────────────────────────────────────────────────────
# PAGE — MY WORKSPACE (case queue + case detail workspace)
# ─────────────────────────────────────────────────────────────
elif active_page == "My Workspace":
    _sel_id = st.session_state.get("selected_case_id")

    if not _sel_id:
        # ── Queue view — landing on My Workspace directly shows "what's on
        # my plate": cases actively moving through the pipeline right now.
        # No real per-analyst login/assignment exists in this app, so this
        # is an honest proxy for "in progress," not a fabricated count.
        st.markdown(_ui.page_title(
            "My Workspace", "Cases currently in progress across the pipeline.",
            "Your queue"), unsafe_allow_html=True)
        _stage_map_q = _pipeline_stage_map()
        _in_progress = {"post_triage_investigate", "post_investigation",
                        "pending_ticket_report"}
        _wip_ids = {i for i, s in _stage_map_q.items() if s in _in_progress}
        _wip_rows = [r for r in db_load_incidents(limit=2000)
                    if str(r.get("id")) in _wip_ids]
        _render_case_table(
            _wip_rows, key_prefix="wscase", heading="In progress",
            subheading=f"{len(_wip_rows)} case(s) currently being worked.")
    else:
        _inc_row = db_get_incident(_sel_id)
        if not _inc_row:
            st.error(f"Case {_sel_id} not found — it may have been removed.")
            if st.button("← Back to My Workspace", key="ws_back_missing"):
                st.query_params.clear()
                st.session_state.selected_case_id = None
                st.rerun()
        else:
            try:
                _raw = _json.loads(_inc_row.get("raw_json") or "{}")
            except Exception:
                _raw = {}
            inc = _raw if isinstance(_raw, dict) and _raw.get("id") else _inc_row

            sev = (normalise_sev(inc) if ("severity" in inc or "riskScore" in inc)
                  else _inc_row.get("severity", "LOW"))
            title    = inc.get("title") or inc.get("name") or _inc_row.get("title") or "Untitled"
            status   = str(inc.get("status") or _inc_row.get("status") or "—")
            assignee = inc.get("assignee") or _inc_row.get("assignee") or "Unassigned"
            alerts   = inc.get("alertCount") or inc.get("numAlerts") or _inc_row.get("alert_count") or "—"
            created  = str(inc.get("created") or inc.get("createdDate")
                          or _inc_row.get("created") or "—")[:16]
            updated  = str(inc.get("updated") or inc.get("lastUpdated")
                          or _inc_row.get("updated") or created)[:16]
            _alert_label = (
                f"{alerts} {'alert' if str(alerts) == '1' else 'alerts'}"
                if str(alerts) != "—" else "NetWitness incident"
            )

            _stage_map   = _pipeline_stage_map()
            _stage_key   = _stage_map.get(_sel_id)
            _stage_label = _STAGE_LABEL_SHORT.get(_stage_key, "Not triaged")
            _action, _action_detail = _next_action_for(_sel_id, status, _stage_map)
            _stage_states = _case_stage_states(_sel_id, status, _stage_map,
                                               parsing_status=_inc_row.get("parsing_status"))
            _current_stage = next(
                (s["name"] for s in _stage_states if s["state"] == "current"),
                _CASE_DISPLAY_STAGES[0],
            )
            _requested_stage = st.query_params.get("case_stage")
            _remembered_stage = st.session_state.get("case_selected_stage")
            _remembered_case = st.session_state.get("case_selected_stage_id")
            if str(_remembered_case) == str(_sel_id) and (
                _remembered_stage in _CASE_DISPLAY_STAGES
            ):
                _selected_stage = _remembered_stage
            elif _requested_stage in _CASE_DISPLAY_STAGES:
                _selected_stage = _requested_stage
            else:
                _selected_stage = _current_stage
            st.session_state.case_selected_stage = _selected_stage
            st.session_state.case_selected_stage_id = str(_sel_id)

            _bc1, _bc2 = st.columns([6, 1])
            with _bc1:
                st.markdown(
                    f'<div style="font-size:0.75rem;color:var(--muted)">'
                    f'My Workspace &nbsp;›&nbsp; {_esc_html(_sel_id)}</div>',
                    unsafe_allow_html=True)
            with _bc2:
                if st.button("← Back", key="ws_back", use_container_width=True):
                    st.query_params.clear()
                    st.session_state.selected_case_id = None
                    st.rerun()

            _main_col, _chat_col = st.columns([2.5, 1.0])

            with _main_col:
                # Case-level actions live in the header so they remain visible
                # before scrolling and are not mistaken for stage controls.
                # This reuses the existing full workflow trigger.
                _active_run = _workflow_store().get("run") or {}
                _parsing_status = str(_inc_row.get("parsing_status") or "Pending")
                _process_running = (
                    _parsing_status == "Processing"
                    or (str(_active_run.get("incident_id") or "") == str(_sel_id)
                        and not _active_run.get("done"))
                )
                _process_closed = status.upper() in _CLOSED_STATUSES
                _process_label = (
                    "Process running" if _process_running else
                    "Process complete" if _process_closed else
                    "▶ Start Process"
                )
                st.html("""
                <style>
                div.st-key-case_header_action {
                    margin: 8px 0 12px;
                    padding: 14px 16px;
                    border: 1px solid #293b57;
                    border-left: 3px solid var(--case-action-border, #7182ff);
                    border-radius: 12px;
                    background: #0d1929;
                    box-shadow: 0 10px 28px #0002;
                }
                div.st-key-case_header_action .ag-casehdr {
                    margin: 0 !important;
                    padding: 0 !important;
                    border: 0 !important;
                    border-left: 0 !important;
                    border-radius: 0 !important;
                    background: transparent !important;
                    box-shadow: none !important;
                }
                div.st-key-case_header_action div.stButton {
                    width: 100%;
                }
                div.st-key-case_header_action div.stButton > button {
                    width: 100% !important;
                    min-height: 42px !important;
                    border-radius: 9px !important;
                    font-size: .74rem !important;
                    font-weight: 800 !important;
                    white-space: nowrap !important;
                    box-shadow: 0 8px 20px rgba(92, 105, 255, .22) !important;
                }
                @media(max-width: 1100px) {
                    div.st-key-case_header_action [data-testid="stHorizontalBlock"] {
                        flex-wrap: wrap !important;
                    }
                    div.st-key-case_header_action [data-testid="stColumn"] {
                        min-width: 100% !important;
                    }
                }
                </style>
                """)
                with st.container(key="case_header_action"):
                    _header_content, _header_action = st.columns(
                        [6.5, 1.15], vertical_alignment="center"
                    )
                    with _header_content:
                        st.markdown(_ui.case_header(
                            _sel_id, title, sev=sev, status=status,
                            subtitle=_alert_label,
                            metas=[("Owner", assignee), ("Current stage", _stage_label),
                                  ("Incidents", _alert_label), ("Last updated", updated)],
                            icon={"CRITICAL": "▲", "HIGH": "◆", "MEDIUM": "●",
                                 "LOW": "○"}.get(sev, "○"),
                        ), unsafe_allow_html=True)
                    with _header_action:
                        if st.button(
                            _process_label,
                            key=f"case_start_process_{_sel_id}",
                            type="primary",
                            use_container_width=True,
                            disabled=_process_running or _process_closed,
                            help=(
                                "This case's workflow is already running."
                                if _process_running else
                                "This case is closed."
                                if _process_closed else
                                "Start triage, investigation and reporting for this case."
                            ),
                        ):
                            _full_inc, _is_full = _resolve_full_incident(_sel_id, inc)
                            if not _is_full:
                                st.warning(
                                    "Full alert data isn't cached in this session yet "
                                    "— Parsing will run on the cached incident summary "
                                    "only. Refresh the Incidents list first for "
                                    "complete per-alert data.")
                            result = _run_triage_workflow_with_ui(_full_inc)

                            # Keep the analyst on the case detail page after
                            # starting the process.  This action used to arm the
                            # Ask-a-Question auto-triage flow and navigate there
                            # on rerun, which unexpectedly replaced the case
                            # view with the agent board.
                            if result is not None:
                                st.session_state.chat_incident = _full_inc
                                st.session_state.pending_auto_triage = False
                                st.session_state.nav_page = "My Workspace"
                                st.rerun()
                            # else: already running, or the workflow failed —
                            # the message was already shown above; no rerun,
                            # so the analyst can read it.

                _render_case_stage_selector(
                    _stage_states, selected_stage=_selected_stage,
                    case_id=_sel_id,
                )

                _selected_state = next(
                    (s["state"] for s in _stage_states
                     if s["name"] == _selected_stage),
                    "queued",
                )
                st.markdown(
                    '<div style="display:flex;align-items:center;gap:10px;'
                    'margin:12px 0 10px">'
                    f'<div style="font-size:1rem;font-weight:800;color:var(--text)">'
                    f'{_esc_html(_selected_stage)}</div>'
                    '<div style="border:1px solid #334966;border-radius:999px;'
                    'padding:3px 9px;color:#9eb0c8;font-size:.65rem;'
                    'text-transform:uppercase;letter-spacing:.06em">'
                    f'{_esc_html(_selected_state)}</div></div>',
                    unsafe_allow_html=True,
                )

                _findings, _ctx = [], []
                try:
                    _findings, _verdict = _build_case_findings(inc)
                    _ctx = _build_case_context(inc, sev, status, alerts, _verdict)
                    _sum_txt = _verdict.get("action") if _verdict.get("available") else ""
                    _titles = (inc.get("alertMeta") or {}).get("AlertTitles") or []
                    _summary = str(inc.get("summary") or "")[:280].strip()
                    _thinking_parts = []
                    if _verdict.get("available"):
                        _thinking_parts.append(
                            f"Unified triage verdict: {_verdict.get('level')}."
                        )
                    if _sum_txt:
                        _thinking_parts.append(f"Recommended action: {_sum_txt}.")
                    if _titles:
                        _thinking_parts.append(
                            "Observed behaviours: "
                            f"{', '.join(list(dict.fromkeys(_titles))[:4])}."
                        )
                    _thinking = " ".join(_thinking_parts)
                    _nar = str(inc.get("summary") or "") + " " + " ".join(str(t) for t in _titles)
                    if _selected_stage == "Parsing":
                        try:
                            _parsing_ai = _json.loads(_inc_row.get("parsing_result_json") or "{}")
                        except Exception:
                            _parsing_ai = {}
                        _parsing_summary = str(_parsing_ai.get("ai_summary") or "").strip()
                        st.markdown(_ui.ai_summary(
                            _parsing_summary or
                            "No AI summary yet — run Start Process to populate.",
                            _ui.detect_fallback(_parsing_summary)), unsafe_allow_html=True)
                    elif _selected_stage == "Triage":
                        # Real LLM-generated summary/thinking from
                        # soc_workflow.generate_triage_ai_summary(), stored
                        # alongside the ticket by run_until_triage_approval() —
                        # same pattern as the Parsing stage above, not the
                        # ad hoc triage_verdict text used for other stages.
                        try:
                            _triage_saved = _json.loads(_inc_row.get("triage_result_json") or "{}")
                        except Exception:
                            _triage_saved = {}
                        # Backfill: cases triaged before this feature existed
                        # have a saved ticket but no ai_summary yet — generate
                        # it once, on demand, and persist it so the LLM call
                        # doesn't repeat on every render (a saved failure
                        # message still counts as "generated" and is left
                        # alone, same as a successful one).
                        if (WORKFLOW_OK and _triage_saved.get("ticket")
                                and not _triage_saved.get("ai_summary")):
                            try:
                                _triage_saved.update(
                                    wf_generate_triage_ai_summary(_triage_saved))
                                _run_id_bf = _inc_row.get("run_id")
                                if _run_id_bf:
                                    wss_save_triage_result(
                                        _sel_id, _run_id_bf, _triage_saved)
                            except Exception:
                                pass
                        _triage_ai_summary = str(_triage_saved.get("ai_summary") or "").strip()
                        # Thinking Process is always computed live from the
                        # agent's own trace (never cached/backfilled) — it's
                        # deterministic and free, so every case shows the
                        # real reasoning immediately, including ones triaged
                        # before this rendering existed at all.
                        try:
                            _triage_ai_thinking = (
                                wf_render_triage_thinking_plain(_triage_saved).strip()
                                if WORKFLOW_OK and _triage_saved.get("trace") else "")
                        except Exception:
                            _triage_ai_thinking = ""
                        _summary_col, _thinking_col = st.columns(2)
                        with _summary_col:
                            st.markdown(_ui.ai_summary(
                                _triage_ai_summary or
                                "No AI summary yet — run Triage to populate.",
                                _ui.detect_fallback(_triage_ai_summary)),
                                unsafe_allow_html=True)
                        with _thinking_col:
                            st.markdown(_ui.ai_summary(
                                _triage_ai_thinking or
                                "No thinking process yet — run Triage to populate.",
                                _ui.detect_fallback(_triage_ai_thinking),
                                title="Thinking Process"),
                                unsafe_allow_html=True)
                    else:
                        _summary_col, _thinking_col = st.columns(2)
                        with _summary_col:
                            st.markdown(_ui.ai_summary(
                                _summary or
                                "No AI summary yet — run Triage to populate.",
                                _ui.detect_fallback(_nar)),
                                unsafe_allow_html=True)
                        with _thinking_col:
                            st.markdown(_ui.ai_summary(
                                _thinking or
                                "No thinking process yet — run Triage to populate.",
                                _ui.detect_fallback(_nar),
                                title="Thinking Process"),
                                unsafe_allow_html=True)
                except Exception:
                    pass

                # Generated Files — only meaningful once the reporting agent
                # has actually run for this case (finalized_report stage).
                if (_selected_stage != "Parsing"
                        and _stage_key == "finalized_report"):
                    _fin_row = next(
                        (r for r in pipeline_load("finalized_report", limit=300)
                         if str(r.get("incident_id")) == _sel_id), None)
                    if _fin_row:
                        try:
                            _fj = _json.loads(_fin_row.get("raw_json") or "{}")
                            _exports = ((_fj.get("report") or {})
                                       .get("document_exports") or {})
                        except Exception:
                            _exports = {}
                        st.markdown(
                            '<div style="font-size:0.85rem;font-weight:700;'
                            'color:var(--text);margin:14px 0 6px">Generated Files</div>',
                            unsafe_allow_html=True)
                        _gc1, _gc2 = st.columns(2)
                        _docx_p, _pdf_p = _exports.get("docx"), _exports.get("pdf")
                        if _docx_p and Path(str(_docx_p)).exists():
                            _gc1.download_button(
                                "Export Word", data=Path(str(_docx_p)).read_bytes(),
                                file_name=f"report_{_sel_id}.docx", key="ws_docx",
                                mime=("application/vnd.openxmlformats-officedocument"
                                     ".wordprocessingml.document"),
                                use_container_width=True)
                        else:
                            _gc1.button("Export Word", disabled=True, key="ws_docx_dis",
                                       use_container_width=True)
                        if _pdf_p and Path(str(_pdf_p)).exists():
                            _gc2.download_button(
                                "Export PDF", data=Path(str(_pdf_p)).read_bytes(),
                                file_name=f"report_{_sel_id}.pdf", key="ws_pdf",
                                mime="application/pdf", use_container_width=True)
                        else:
                            _gc2.button("Export PDF", disabled=True, key="ws_pdf_dis",
                                       use_container_width=True)

                if _selected_stage == "Parsing":
                    # Parsing produces structured data rather than an analytical
                    # case view. Keep the downstream Overview/Timeline/etc. tabs
                    # out of this stage and expose the parsed record directly.
                    _parsed_normalised_output = (
                        inc if isinstance(inc, dict) else dict(_inc_row)
                    )
                    _parsed_json = _json.dumps(
                        _parsed_normalised_output, indent=2, default=str
                    ).encode("utf-8")
                    _parsed_suffix = re.sub(
                        r"[^A-Za-z0-9_\-]", "_", str(_sel_id)
                    )[:40]
                    _parsed_saved = updated if updated and updated != "â€”" else created

                    st.markdown(
                        '<div class="parsing-stage-view"></div>'
                        '<style>'
                        'div[data-testid="stColumn"]:has(.parsing-stage-view) '
                        'div[data-testid="stTabs"]{'
                        'display:none !important;}'
                        '</style>'
                        '<div style="font-size:.85rem;font-weight:700;'
                        'color:var(--text);margin:16px 0 2px">Generated Files</div>'
                        '<div style="font-size:.7rem;color:var(--muted);'
                        'margin-bottom:10px">Download the structured parsed and '
                        'normalised case data.</div>',
                        unsafe_allow_html=True,
                    )
                    with st.container(border=True):
                        _gf_doc, _gf_saved, _gf_action = st.columns(
                            [3.8, 1.6, 1.7], vertical_alignment="center"
                        )
                        with _gf_doc:
                            st.markdown(
                                '<div style="font-size:.58rem;color:var(--faint);'
                                'letter-spacing:.08em;margin-bottom:10px">DOCUMENT</div>'
                                '<div style="display:flex;gap:10px;align-items:center">'
                                '<div style="width:34px;height:34px;border-radius:8px;'
                                'background:#41516d;color:#e8eef8;display:flex;'
                                'align-items:center;justify-content:center;'
                                'font-family:var(--mono);font-size:.8rem">&lt;&gt;</div>'
                                '<div><div style="font-size:.78rem;font-weight:700;'
                                'color:var(--text)">Parsed &amp; Normalised output</div>'
                                '<div style="font-size:.65rem;color:var(--muted);'
                                'margin-top:2px">Structured parsed case data ready '
                                'for downstream agents.</div></div></div>',
                                unsafe_allow_html=True,
                            )
                        with _gf_saved:
                            st.markdown(
                                '<div style="font-size:.58rem;color:var(--faint);'
                                'letter-spacing:.08em;margin-bottom:10px">LAST SAVED</div>'
                                f'<div style="font-size:.66rem;color:var(--muted)">'
                                f'{_esc_html(_parsed_saved)}</div>',
                                unsafe_allow_html=True,
                            )
                        with _gf_action:
                            st.markdown(
                                '<div style="font-size:.58rem;color:var(--faint);'
                                'letter-spacing:.08em;margin-bottom:4px">ACTIONS</div>',
                                unsafe_allow_html=True,
                            )
                            st.download_button(
                                "Download JSON",
                                data=_parsed_json,
                                file_name=(
                                    f"parsed_normalised_output_{_parsed_suffix}.json"
                                ),
                                mime="application/json",
                                key=f"ws_parsed_json_{_parsed_suffix}",
                                use_container_width=True,
                            )

                st.markdown("<div style='margin:10px 0'></div>", unsafe_allow_html=True)
                if _selected_stage == "Triage":
                    (_t_output,) = st.tabs(["Output"])
                    with _t_output:
                        try:
                            _triage_out = _json.loads(
                                _inc_row.get("triage_result_json") or "{}")
                        except Exception:
                            _triage_out = {}
                        _out_ticket = _triage_out.get("ticket")
                        if _out_ticket:
                            st.markdown(
                                render_triage_trace(_triage_out.get("trace") or [])
                                + "\n" + format_ticket_display(_out_ticket))
                        else:
                            st.caption("No triage output yet — run Triage to populate.")
                elif _selected_stage == "Threat Intelligence Enrichment":
                    (_t_output,) = st.tabs(["Output"])
                    with _t_output:
                        try:
                            from threat_intel import enrich_iocs, format_enrichment
                            try:
                                _ti_triage = _json.loads(
                                    _inc_row.get("triage_result_json") or "{}") or None
                            except Exception:
                                _ti_triage = None
                            _enr = enrich_iocs(inc, _ti_triage)
                            if _enr.get("results"):
                                st.code(format_enrichment(_enr), language=None)
                            else:
                                st.caption("No IOCs found to enrich for this case.")
                        except Exception as _ti_err:
                            st.caption(
                                f"Threat intelligence enrichment unavailable: {_ti_err}")
                else:
                    (_t_ovw, _t_time, _t_mitre, _t_graph, _t_evid, _t_act) = st.tabs(
                        ["Overview", "Timeline", "MITRE ATT&CK", "Entity Graph",
                         "Evidence", "Activity"])

                    with _t_ovw:
                        oc1, oc2 = st.columns([1.4, 0.9])
                        with oc1:
                            _fh = (_ui.key_findings(_findings) if _findings else
                                  '<div style="color:var(--sub);font-size:.8rem">'
                                  'No findings distilled yet — run Triage.</div>')
                            st.markdown(_ui.panel_open(
                                "Key findings", "Behaviours &amp; analytic signals")
                                + _fh + _ui.panel_close(), unsafe_allow_html=True)
                        with oc2:
                            st.markdown(_ui.panel_open(
                                "Case context", "Key facts for the current decision")
                                + _ui.context_grid(_ctx) + _ui.panel_close(),
                                unsafe_allow_html=True)

                    with _t_time:
                        try:
                            from incident_map import build_incident_map
                            _imap_t = build_incident_map(inc)
                            if _imap_t.get("timeline"):
                                st.markdown("\n".join(
                                    f"- `{t['time'][:19]}` — {t['event']}"
                                    for t in _imap_t["timeline"][:20]))
                            else:
                                st.caption("No timeline events recorded for this case.")
                        except Exception as _tl_err:
                            st.caption(f"Timeline unavailable: {_tl_err}")

                    with _t_mitre:
                        try:
                            from tactic_inference import infer_tactics
                            _ti = infer_tactics(inc)
                            _mitre_maps = []
                            # Prefer the investigation agent's evidence-rich mappings.
                            for _candidate in (
                                inc.get("mitre_mappings"),
                                (_raw or {}).get("mitre_mappings"),
                                ((_raw or {}).get("final_report") or {}).get("mitre_mappings"),
                                ((_raw or {}).get("investigation_result") or {}).get("mitre_mappings"),
                            ):
                                if isinstance(_candidate, list) and _candidate:
                                    _mitre_maps = _candidate
                                    break
                            if not _mitre_maps:
                                _tactics = (_ti.get("tactics") if _ti.get("available")
                                            else inc.get("tactics")) or []
                                _techniques = (_ti.get("techniques") if _ti.get("available")
                                               else inc.get("techniques")) or []
                                _fallback_tactic = (_ti.get("tactic") if _ti.get("available")
                                                    else inc.get("mitre_tactic"))
                                _fallback_tech = (_ti.get("technique") if _ti.get("available")
                                                  else inc.get("mitre_technique"))
                                if not _tactics and _fallback_tactic:
                                    _tactics = [_fallback_tactic]
                                if not _techniques and _fallback_tech:
                                    _techniques = [_fallback_tech]
                                _map_count = max(len(_tactics), len(_techniques))
                                for _mi in range(_map_count):
                                    _mitre_maps.append({
                                        "tactic": (_tactics[_mi] if _mi < len(_tactics)
                                                   else _tactics[-1] if _tactics else "Unclassified"),
                                        "technique_id": (_techniques[_mi] if _mi < len(_techniques)
                                                         else _techniques[-1] if _techniques else ""),
                                        "technique_name": (_ti.get("technique_name") if _mi == 0
                                                           else "MITRE ATT&CK technique"),
                                        "confidence": _ti.get("confidence") or "high",
                                        "evidence": _ti.get("evidence") or [],
                                        "source": _ti.get("source") or "Aegis Investigation Agent",
                                    })
                            st.markdown(_ui.mitre_mapping_workspace(_mitre_maps),
                                        unsafe_allow_html=True)
                        except Exception as _mt_err:
                            st.caption(f"MITRE mapping unavailable: {_mt_err}")

                    with _t_graph:
                        # This IS the app's existing "Map"/entity-relationship
                        # feature, relocated here rather than rebuilt.
                        try:
                            from incident_map import build_incident_map, to_dot, map_caption
                            _imap_g = build_incident_map(inc)
                            st.graphviz_chart(to_dot(_imap_g), width="stretch")
                            st.caption(map_caption(_imap_g))
                        except Exception as _mg_err:
                            st.caption(f"Entity graph unavailable: {_mg_err}")

                    with _t_evid:
                        _alerts_list = (_raw or {}).get("alerts")
                        if _alerts_list:
                            for _al in _alerts_list[:20]:
                                st.markdown(
                                    f'<div style="background:#091624;padding:8px 12px;'
                                    f'border-radius:4px;margin-bottom:6px;'
                                    f'border-left:3px solid var(--accent)">'
                                    f'{_esc_html(_al.get("title") or _al.get("name") or "Untitled Alert")}'
                                    f'</div>', unsafe_allow_html=True)
                        else:
                            st.caption("No associated alerts recorded for this case.")
                        try:
                            from ioc_correlation import correlate_iocs, format_correlation
                            _corr = correlate_iocs(inc)
                            if _corr.get("available") and _corr.get("results"):
                                st.markdown("**Internal IOC correlation**")
                                st.code(format_correlation(_corr), language=None)
                        except Exception:
                            pass

                    with _t_act:
                        _case_runs = [r for r in pipeline_load("workflow_runs", limit=300)
                                      if str(r.get("incident_id")) == _sel_id]
                        if _case_runs:
                            for _r in _case_runs:
                                st.markdown(f"- `{_r.get('created_at', '')}` — "
                                           f"{_esc_html(_r.get('summary', ''))}")
                        else:
                            st.caption("No workflow runs recorded for this case yet.")

            with _chat_col:
                # ── Ask Aegis — reuses the existing chat_respond() Q&A path,
                # scoped to this case. No new LLM/agent logic; new UI surface.
                st.markdown("""
                <style>
                /* The inline style element is the first item in this column.
                   Remove Streamlit's default inter-element gap so it does not
                   push the Ask Aegis card below the case summary beside it. */
                div[data-testid="stColumn"]:has(div.st-key-ws_ask_aegis)
                    > div[data-testid="stVerticalBlock"] {
                    gap: 0 !important;
                }
                /* Keep Ask Aegis visible while the case details beside it scroll.
                   The column is the sticky element so its full Streamlit-rendered
                   card (header, messages, actions and composer) moves together. */
                div[data-testid="stColumn"]:has(div.st-key-ws_ask_aegis) {
                    position: sticky !important;
                    top: 4.75rem !important;
                    align-self: flex-start !important;
                    height: fit-content !important;
                    max-height: calc(100vh - 5.75rem) !important;
                    z-index: 5;
                }
                div.st-key-ws_ask_aegis {
                    --ws-aegis-inline-padding: 20px;
                    background: #0b1524;
                    border: 1px solid #263750;
                    border-radius: 14px;
                    margin-top: 8px;
                    overflow: hidden;
                    box-shadow: 0 15px 45px rgba(0, 0, 0, .2);
                }
                div.st-key-ws_ask_aegis > div[data-testid="stVerticalBlock"] {
                    gap: 0 !important;
                }
                div.st-key-ws_ask_header {
                    border-bottom: 1px solid #263750;
                    padding: 18px var(--ws-aegis-inline-padding) 16px;
                }
                .ws-aegis-head {
                    display: grid;
                    grid-template-columns: 42px minmax(0, 1fr) auto;
                    align-items: center;
                    gap: 13px;
                }
                .ws-aegis-icon {
                    display: grid;
                    place-items: center;
                    width: 42px;
                    height: 42px;
                    border-radius: 11px;
                    color: #fff;
                    font-size: 21px;
                    background: linear-gradient(145deg, #7889ff, #4b5bd4);
                    box-shadow: inset 0 1px 0 rgba(255,255,255,.2);
                }
                .ws-aegis-title {
                    color: #eef2f8;
                    font-size: 1rem;
                    font-weight: 750;
                    line-height: 1.2;
                }
                .ws-aegis-sub {
                    color: #8d9bb0;
                    font-size: .76rem;
                    line-height: 1.4;
                    margin-top: 3px;
                }
                .ws-aegis-case {
                    border: 1px solid #34455e;
                    border-radius: 999px;
                    color: #8d9bb0;
                    font: 600 .7rem/1 var(--mono);
                    letter-spacing: .02em;
                    padding: 7px 12px;
                    white-space: nowrap;
                }
                div.st-key-ws_ask_messages {
                    background: #0b1524;
                    padding: 12px 14px 8px;
                }
                div.st-key-ws_ask_messages .bubble-user,
                div.st-key-ws_ask_messages .bubble-agent {
                    background: #111e32;
                    border: 1px solid #293b56;
                    border-left-width: 1px;
                    border-radius: 10px;
                    line-height: 1.55;
                    margin: 0 0 9px;
                    padding: 11px 12px;
                }
                div.st-key-ws_ask_actions {
                    padding: 4px var(--ws-aegis-inline-padding) 10px;
                }
                div.st-key-ws_ask_actions [data-testid="stHorizontalBlock"] {
                    display: flex !important;
                    flex-wrap: wrap !important;
                    gap: 6px !important;
                }
                div.st-key-ws_ask_actions [data-testid="stColumn"] {
                    flex: 1 1 calc(50% - 3px) !important;
                    width: calc(50% - 3px) !important;
                    max-width: calc(50% - 3px) !important;
                    min-width: 0 !important;
                }
                div.st-key-ws_ask_actions [data-testid="stColumn"]:nth-child(3) {
                    flex: 0 0 100% !important;
                    width: 100% !important;
                    max-width: 100% !important;
                }
                div.st-key-ws_ask_actions div.stButton > button {
                    width: 100% !important;
                    min-height: 27px !important;
                    height: 27px !important;
                    border-radius: 999px !important;
                    border: 1px solid #2b405d !important;
                    background: #0d192a !important;
                    color: #aab7ca !important;
                    font-size: .62rem !important;
                    padding: 4px 10px !important;
                    box-shadow: none !important;
                    white-space: nowrap;
                }
                div.st-key-ws_ask_actions div.stButton > button:hover {
                    border-color: #6076e8 !important;
                    color: #e8ecff !important;
                    transform: none !important;
                }
                div.st-key-ws_ask_composer {
                    border-top: 1px solid #263750;
                    padding: 11px var(--ws-aegis-inline-padding) 12px;
                }
                div.st-key-ws_ask_composer [data-testid="stHorizontalBlock"] {
                    align-items: center;
                    gap: 8px !important;
                }
                div.st-key-ws_ask_composer [data-testid="stColumn"]:first-child {
                    flex: 1 1 auto !important;
                    width: auto !important;
                    min-width: 0 !important;
                }
                div.st-key-ws_ask_composer [data-testid="stColumn"]:last-child {
                    flex: 0 0 74px !important;
                    width: 74px !important;
                    min-width: 74px !important;
                }
                div.st-key-ws_ask_composer [data-testid="stElementContainer"] {
                    margin-bottom: 0 !important;
                }
                div.st-key-ws_ask_composer .stTextInput input {
                    min-height: 35px !important;
                    height: 35px !important;
                    border-color: #293b56 !important;
                    border-radius: 8px !important;
                    font-size: .68rem !important;
                }
                div.st-key-ws_ask_composer div.stButton > button {
                    min-height: 35px !important;
                    height: 35px !important;
                    border: 0 !important;
                    border-radius: 8px !important;
                    background: #6473e8 !important;
                    color: #fff !important;
                    font-size: .68rem !important;
                    font-weight: 700 !important;
                    padding: 0 14px !important;
                    box-shadow: none !important;
                    white-space: nowrap !important;
                    word-break: normal !important;
                }
                div.st-key-ws_ask_composer div.stButton > button p {
                    white-space: nowrap !important;
                    word-break: normal !important;
                    margin: 0 !important;
                }
                </style>
                """, unsafe_allow_html=True)

                _hist = st.session_state.workspace_chat.setdefault(_sel_id, [])
                _prompt = None
                with st.container(key="ws_ask_aegis"):
                    with st.container(key="ws_ask_header"):
                        st.markdown(
                            '<div class="ws-aegis-head">'
                            '<div class="ws-aegis-icon">✦</div>'
                            '<div><div class="ws-aegis-title">Ask Aegis</div>'
                            '<div class="ws-aegis-sub">Shared across every agent stage &amp; '
                            'Threat Intelligence Enrichment</div></div>'
                            f'<div class="ws-aegis-case">{_esc_html(_sel_id)}</div>'
                            '</div>',
                            unsafe_allow_html=True,
                        )

                    with st.container(height=520, border=False, key="ws_ask_messages"):
                        for _m in _hist[-20:]:
                            if _m["role"] == "user":
                                st.markdown(
                                    '<div class="bubble-user" style="font-size:0.72rem">'
                                    f'{_esc_html(_m["content"])}</div>',
                                    unsafe_allow_html=True,
                                )
                            else:
                                st.markdown(
                                    '<div class="bubble-agent" style="font-size:0.72rem">'
                                    f'{_m["content"]}</div>',
                                    unsafe_allow_html=True,
                                )

                    with st.container(key="ws_ask_actions"):
                        _p1, _p2, _p3 = st.columns([1, 1, 1])
                        if _p1.button("Summarise investigation", key="ws_chip_sum",
                                     help="Summarise investigation"):
                            _prompt = "Summarise the investigation on this case."
                        if _p2.button("Explain key findings", key="ws_chip_find",
                                     help="Explain key findings"):
                            _prompt = "Explain the key findings for this case."
                        if _p3.button("What requires analyst attention?",
                                     key="ws_chip_next",
                                     help="What requires analyst attention?"):
                            _prompt = "What requires analyst attention on this case?"

                    with st.container(key="ws_ask_composer"):
                        _wc1, _wc2 = st.columns([5, 1])
                        _typed = _wc1.text_input(
                            "Ask about this case…",
                            key="ws_chat_typed",
                            label_visibility="collapsed",
                            placeholder="Ask about this case…",
                        )
                        if _wc2.button("Send", key="ws_chat_send",
                                       use_container_width=True):
                            _prompt = _typed.strip() or None

                    if _prompt:
                        _hist.append({"role": "user", "content": _prompt})
                        with st.spinner("Aegis is thinking…"):
                            try:
                                _reply = chat_respond(
                                    _prompt, incident=inc,
                                    parsed_context=db_load_parsed_context(
                                        str(inc.get("id") or inc.get("incidentId") or "")))
                            except Exception as _ce:
                                _reply = f"Error: {_ce}"
                        _hist.append({"role": "assistant", "content": _reply})
                        st.rerun()


# ─────────────────────────────────────────────────────────────
# PAGE 2 — ASK A QUESTION
# ─────────────────────────────────────────────────────────────
elif active_page == "Ask a Question":

    st.markdown(
        '<div class="info-box"><div class="title">AI Agents</div>'
        'You can ask plain-language questions about your security alerts here. '
        'For example: <em>"What are the most critical incidents today?"</em> or '
        '<em>"Summarise the latest high-priority alerts."</em> '
        'You do not need any technical knowledge to use this.</div>',
        unsafe_allow_html=True,
    )

    # ── Global CSS for the spinning phase icon (injected once, never reset) ──
    st.markdown("""
    <style>
    @keyframes soc-phase-spin { to { transform: rotate(360deg); } }
    .soc-phase-spinner {
        display: inline-block;
        animation: soc-phase-spin 1.2s linear infinite;
        color: #E8623A;
        font-size: 0.95rem;
        line-height: 1;
        margin-right: 8px;
        vertical-align: middle;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Helper: parse uploaded file into an incident dict ──────
    def _parse_uploaded_file(uploaded_file) -> tuple[dict, str]:
        """
        Parse a Streamlit UploadedFile into an incident dict.
        Supports: .json, .csv, .txt, .log
        Returns (incident_dict, error_message).  error_message is "" on success.
        """
        import io
        name = uploaded_file.name
        ext  = name.rsplit(".", 1)[-1].lower()
        raw  = uploaded_file.read()

        try:
            if ext == "json":
                data = _json.loads(raw.decode("utf-8", errors="replace"))
                # Accept a list → take first item; accept a dict → use directly
                if isinstance(data, list):
                    if not data:
                        return {}, "JSON array is empty."
                    incident = data[0] if isinstance(data[0], dict) else {"raw": data[0]}
                elif isinstance(data, dict):
                    # NetWitness envelope: {"items": [...]}
                    if "items" in data and isinstance(data["items"], list) and data["items"]:
                        incident = data["items"][0]
                    else:
                        incident = data
                else:
                    return {}, "JSON root must be an object or array."

            elif ext == "csv":
                import csv, io as _io
                text    = raw.decode("utf-8", errors="replace")
                reader  = list(csv.DictReader(_io.StringIO(text)))
                if not reader:
                    return {}, "CSV has no data rows."
                # Use the first row as the incident dict
                incident = dict(reader[0])
                # Attach all rows as an alert list for richer context
                incident["_csv_alerts"] = reader

            elif ext in ("txt", "log"):
                text = raw.decode("utf-8", errors="replace")
                # Build a minimal incident dict wrapping the raw text
                incident = {
                    "id":          name,
                    "title":       f"Uploaded log — {name}",
                    "description": text[:4000],   # cap to avoid token overflow
                    "raw_log":     text,
                    "source":      "file_upload",
                }

            else:
                return {}, f"Unsupported file type '.{ext}'. Upload a .json, .csv, .txt, or .log file."

        except Exception as exc:
            return {}, f"Failed to parse file: {exc}"

        # Multi-source alert normalization (defensive-security skill's
        # analyze_alert): a SIEM/EDR/NDR/syslog alert that isn't already
        # NetWitness-shaped (no alertMeta) gets normalized into the incident
        # schema — extracted IPs into alertMeta, a triage verdict, MITRE —
        # so the whole pipeline (map/TI/detection/triage) runs on it. Purely
        # additive: NetWitness-shaped uploads keep their alertMeta untouched.
        try:
            if not isinstance(incident.get("alertMeta"), dict) or not incident.get("alertMeta"):
                from alert_triage import normalize_to_incident, validate_alert
                # for a raw txt/log upload, feed the text as the message
                if ext in ("txt", "log") and "message" not in incident:
                    incident["message"] = incident.get("raw_log", "")[:8000]
                    incident.setdefault("timestamp", datetime.now().isoformat())
                    incident.setdefault("source", "file_upload")
                if validate_alert(incident)["ok"]:
                    ctx = ("edr" if ext in ("txt", "log") else "siem")
                    incident = normalize_to_incident(incident, ctx)
        except Exception:
            pass  # normalization is best-effort; never block an upload

        # Ensure there's always an id and title
        if "id" not in incident:
            incident["id"] = name
        if "title" not in incident:
            incident["title"] = name

        return incident, ""

    # ── Resolve which incident the agent will use ──────────────
    # Priority: NW incident (from Incidents tab) > uploaded file
    nw_inc       = st.session_state.chat_incident
    up_inc       = st.session_state.uploaded_incident
    active_inc   = nw_inc if nw_inc else up_inc   # what the agent receives

    # ── Context banner ─────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:10px">■ INCIDENT CONTEXT</div>',
        unsafe_allow_html=True,
    )

    ctx_col, clr_col = st.columns([5, 1])

    with ctx_col:
        if nw_inc:
            sev = normalise_sev(nw_inc)
            st.markdown(
                f'<div style="background:#050F1A;border:1px solid var(--accent);'
                f'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                f'font-size:0.72rem">'
                f'<span class="badge badge-{sev.lower()}">{sev}</span>'
                f'&nbsp;'
                f'<strong>{nw_inc.get("id","?")}</strong> — {nw_inc.get("title","?")}'
                f'<span style="color:var(--muted);font-size:0.6rem;margin-left:10px">'
                f'● from NetWitness</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if nw_inc.get("alerts_fetch_error"):
                st.warning(_alerts_fetch_warning(nw_inc))
        elif up_inc:
            st.markdown(
                f'<div style="background:#050F1A;border:1px solid var(--warn);'
                f'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                f'font-size:0.72rem">'
                f'<strong>{st.session_state.uploaded_filename}</strong>'
                f'<span style="color:var(--muted);font-size:0.6rem;margin-left:10px">'
                f'● from uploaded file</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#07080F;border:1px solid var(--border);'
                'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                'font-size:0.68rem;color:var(--muted)">'
                'No incident context — select one from the Incidents tab '
                'or upload a file below.'
                '</div>',
                unsafe_allow_html=True,
            )

    with clr_col:
        if nw_inc or up_inc:
            if st.button("✕ Clear", use_container_width=True, key="clear_ctx"):
                st.session_state.chat_incident    = None
                st.session_state.uploaded_incident = None
                st.session_state.uploaded_filename = ""
                st.rerun()

    st.markdown("---")

    # ── AGENT BOARD — live thinking + outputs for all 3 agents ─────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:8px">■ AGENT BOARD — CLICK AN AGENT '
        'TO SEE ITS THINKING &amp; OUTPUT</div>',
        unsafe_allow_html=True,
    )

    _AGENTS = [
        ("triage",        "◎", "Triage Agent",        "#00D4FF"),
        ("investigation", "⌕", "Investigation Agent", "#FF7700"),
        ("reporting",     "▤", "Reporting Agent",     "#A78BFA"),
    ]
    _BOARD_BADGES = {
        "idle":    ("○ IDLE",    "#3A607A"),
        "queued":  ("QUEUED",  "#8B9DC3"),
        "running": ("RUNNING", "#FFB700"),
        "done":    ("DONE",    "#0AF0A0"),
        "cached":  ("CACHED",  "#0AF0A0"),
        "skipped": ("SKIPPED", "#3A607A"),
        "failed":  ("FAILED",  "#FF3B3B"),
        "rejected": ("REJECTED", "#FF6E7C"),
    }

    # ── Background workflow sync: worker state → board, finalize once ──────
    _wf_active = _workflow_store().get("run")
    if _wf_active:
        # Adopt the worker's live panel dicts (same objects — subsequent
        # worker writes are visible on every poll rerun).
        for _ag in ("investigation", "reporting"):
            if _ag in _wf_active.get("panels", {}):
                st.session_state.agent_board[_ag] = _wf_active["panels"][_ag]
        if _wf_active.get("done") and not _wf_active.get("finalized"):
            _wf_active["finalized"] = True
            for _stage, _rec in _wf_active.get("chroma_queue", []):
                try:
                    pipeline_chroma_insert(_stage, _rec)
                except Exception:
                    pass
            st.session_state.chat_history.append(
                {"role": "assistant",
                 "content": "\n".join(_wf_active.get("wf_md") or
                                      ["Workflow finished."]),
                 "ts": datetime.now().strftime("%H:%M:%S")})
            st.session_state.setdefault("_surfaced_runs", []).append(
                _wf_active.get("run_id"))
            _workflow_store()["run"] = None
    else:
        # No active worker — sweep zombie states from interrupted sessions.
        for _ag in ("investigation", "reporting"):
            if st.session_state.agent_board[_ag]["status"] in ("running", "queued"):
                st.session_state.agent_board[_ag]["status"] = "failed"
                st.session_state.agent_board[_ag]["thinking"].append(
                    f"[{datetime.now().strftime('%H:%M:%S')}] previous run "
                    "was interrupted before completion")

    # ── Disk fallback: surface a completed run this session never saw ──────
    # The in-memory finalize above only works while the session's poll loop
    # stays alive for the whole run. The worker also persists its finished
    # results to disk; if a recent run (< 30 min) hasn't been surfaced in
    # THIS session, deliver its board panels + chat summary now.
    try:
        _lwr_path = SOC_DB_DIR / "last_workflow_result.json"
        if _lwr_path.exists():
            _lwr = _json.loads(_lwr_path.read_text(encoding="utf-8"))
            _surf = st.session_state.setdefault("_surfaced_runs", [])
            _fresh = (time.time() - float(_lwr.get("finished_at") or 0)) < 1800
            if (_lwr.get("done") and _fresh
                    and _lwr.get("run_id") not in _surf
                    and not (_workflow_store().get("run") or {})):
                for _ag in ("investigation", "reporting"):
                    _p = (_lwr.get("panels") or {}).get(_ag)
                    if _p:
                        st.session_state.agent_board[_ag] = _p
                st.session_state.chat_history.append(
                    {"role": "assistant",
                     "content": "\n".join(_lwr.get("wf_md") or
                                          ["Workflow finished."]),
                     "ts": datetime.now().strftime("%H:%M:%S")})
                _surf.append(_lwr.get("run_id"))
    except Exception:
        pass

    _ANSI_RE = re.compile(r"\x1b\[[0-9;]*m")

    def _board_set(agent: str, status: str | None = None,
                   think: str | None = None, output: str | None = None,
                   progress: int | None = None) -> None:
        """Update an agent's board state (+ live card/detail refresh mid-run)."""
        panel = st.session_state.agent_board[agent]
        if status is not None:
            panel["status"] = status
            if status in ("done", "cached", "skipped"):
                panel["progress"] = 100
        if progress is not None:
            panel["progress"] = max(0, min(100, int(progress)))
        if think is not None:
            clean = _ANSI_RE.sub("", str(think))   # subprocess logs carry colours
            panel["thinking"].append(f"[{datetime.now().strftime('%H:%M:%S')}] {clean}")
            panel["thinking"] = panel["thinking"][-60:]   # keep the tail
        if output is not None:
            panel["output"] = output
        panel["updated"] = datetime.now().strftime("%H:%M:%S")
        # Live refresh of this agent's card + open detail panel during a run.
        slot = _board_live.get(agent)
        if slot is not None:
            try:
                _render_board_card(slot, agent)
            except Exception:
                pass
        try:
            _render_board_detail(agent)
        except Exception:
            pass

    class _BoardTee:
        """Duck-typed st.empty() that mirrors writes to several containers —
        used to send the triage LLM token stream to both the in-status panel
        and the agent board's detail view."""
        def __init__(self, *targets):
            self.targets = [t for t in targets if t is not None]

        def markdown(self, *a, **k):
            for t in self.targets:
                try:
                    t.markdown(*a, **k)
                except Exception:
                    pass

        def empty(self):
            for t in self.targets:
                try:
                    t.empty()
                except Exception:
                    pass

    def _render_board_card(container, agent: str) -> None:
        icon, name, color = {a: (i, n, c) for a, i, n, c in _AGENTS}[agent]
        panel  = st.session_state.agent_board[agent]
        status = panel["status"]
        badge, bcolor = _BOARD_BADGES.get(status, _BOARD_BADGES["idle"])
        upd = panel.get("updated") or "—"
        tail = ""
        if status == "running" and panel["thinking"]:
            last = panel["thinking"][-1]
            tail = (f'<div style="font-family:var(--mono);font-size:var(--fs-3xs);'
                    f'color:var(--muted);margin-top:5px;white-space:nowrap;'
                    f'overflow:hidden;text-overflow:ellipsis">{last[-80:]}</div>')
        pct = int(panel.get("progress", 0) or 0)
        _bar = ""
        if status in ("running", "done", "cached") or pct:
            _bar = (f'<div style="height:3px;background:#12202f;border-radius:2px;'
                    f'margin-top:7px;overflow:hidden"><div style="width:{pct}%;'
                    f'height:100%;background:{color};border-radius:2px;'
                    f'transition:width .4s ease"></div></div>')
        container.markdown(
            f'<div style="background:#060C16;border:1px solid {color}44;'
            f'border-left:3px solid {color};border-radius:7px;'
            f'padding:10px 12px;min-height:74px">'
            f'<div style="display:flex;align-items:center;gap:8px">'
            f'<span style="font-size:1.15rem">{icon}</span>'
            f'<strong style="flex:1;font-size:0.8rem;color:{color}">{name}</strong>'
            f'<span style="background:{bcolor}22;color:{bcolor};'
            f'border:1px solid {bcolor}44;padding:1px 7px;border-radius:3px;'
            f'font-family:var(--mono);font-size:0.55rem">{badge}</span></div>'
            f'<div style="font-family:var(--mono);font-size:var(--fs-3xs);'
            f'color:var(--muted);margin-top:5px">last activity: {upd}'
            f'{(" · " + str(pct) + "%") if (status == "running" or pct) else ""}</div>'
            f'{_bar}{tail}</div>',
            unsafe_allow_html=True,
        )

    # HITL manual-review toggle — OFF (default) keeps the pipeline auto-chaining
    # exactly as before; ON pauses after triage and after investigation for the
    # analyst to review and click Approve before the next agent runs.
    st.toggle("Manual review — approve each hand-off",
              key="manual_review",
              help="When on, the pipeline pauses after triage and after "
                   "investigation. Review each agent's output on this board, "
                   "then click Approve to hand off to the next agent.")

    _board_live: dict = {}
    _b_cols = st.columns(3)
    for _bi, (_ag, _icon, _name, _color) in enumerate(_AGENTS):
        with _b_cols[_bi]:
            _slot = st.empty()
            _board_live[_ag] = _slot
            _render_board_card(_slot, _ag)
            if st.button("View", key=f"board_view_{_ag}", use_container_width=True):
                st.session_state.agent_board_sel = (
                    None if st.session_state.agent_board_sel == _ag else _ag)
                st.rerun()

    # ── HITL approval controls (shown only when a run awaits the analyst) ──
    # Previously Approve was the ONLY option at either gate — an analyst who
    # disagreed with triage/investigation had no control to press. Reject now
    # exists at both gates, requires a reason (see _reject_with_reason), and
    # is written to the same workflow_runs audit trail as every other run.
    _hitl_run = _workflow_store().get("run")
    if _hitl_run and _hitl_run.get("manual_review") and not _hitl_run.get("done"):
        _await = _hitl_run.get("awaiting")
        if _await == "investigate" and not _hitl_run.get("_spawned"):
            st.info("**Triage complete.** Review it above, then approve to hand "
                    "off to the Investigation agent, or reject to stop here.")
            _hg1a, _hg1b = st.columns(2)
            with _hg1a:
                if st.button("Approve → Investigate", type="primary",
                             use_container_width=True, key="hitl_go_inv"):
                    import threading as _th2
                    _hitl_run["_spawned"] = True
                    _hitl_run["awaiting"] = None
                    _th2.Thread(target=_workflow_worker,
                                args=(_hitl_run, _hitl_run["_tri"],
                                      _hitl_run["_incident"]),
                                daemon=True).start()
                    st.rerun()
            with _hg1b:
                _rej_reason = _reject_with_reason("hitl_reject_inv")
                if _rej_reason:
                    _analyst = _hitl_run.get("analyst") or "—"
                    pipeline_insert("workflow_runs", {
                        "id": f"run_{_hitl_run['run_id']}_{_hitl_run['incident_id'][:20]}",
                        "incident_id": _hitl_run["incident_id"],
                        "title": f"Run {_hitl_run['started_hms']} — {_hitl_run['title']}",
                        "severity": _hitl_run["cls"],
                        "summary": f"triage: rejected · ticket {_hitl_run['unc']} "
                                  f"· analyst {_analyst} · reason: {_rej_reason}",
                        "stages": {"triage": "rejected"},
                        "ticket_unc": _hitl_run["unc"],
                        "analyst": _analyst,
                        "gate_decision": "rejected", "gate_reason": _rej_reason})
                    st.session_state.agent_board["triage"]["status"] = "rejected"
                    st.session_state.agent_board["triage"]["thinking"].append(
                        f"[{datetime.now().strftime('%H:%M:%S')}] "
                        f"Rejected by analyst — {_rej_reason}")
                    _workflow_store()["run"] = None
                    st.success("Rejected — the run has been stopped; no "
                              "investigation or report will be generated.")
                    st.rerun()
        elif _await == "report":
            st.info("**Investigation complete.** Review it above, then approve "
                    "to generate the final report, or reject to stop here.")
            _hg2a, _hg2b = st.columns(2)
            with _hg2a:
                if st.button("Approve → Generate report", type="primary",
                             use_container_width=True, key="hitl_go_rep"):
                    _hitl_run["gate_decision"] = "approved"
                    _g = _hitl_run.get("gate_report")
                    if _g is not None:
                        _g.set()
                    _hitl_run["awaiting"] = None
                    st.rerun()
            with _hg2b:
                _rej_reason2 = _reject_with_reason("hitl_reject_rep")
                if _rej_reason2:
                    # The worker thread is already running, blocked on
                    # gate.wait() — set the decision BEFORE releasing it so
                    # it knows to skip run_reporting() instead of running it.
                    _hitl_run["gate_decision"] = "rejected"
                    _hitl_run["gate_reason"] = _rej_reason2
                    _g = _hitl_run.get("gate_report")
                    if _g is not None:
                        _g.set()
                    _hitl_run["awaiting"] = None
                    st.rerun()

    # ── Detail panel for the selected agent (live slots) ───────────────────
    # The thinking/output areas are st.empty() slots registered in
    # _board_live_detail; _board_set rewrites them mid-run, so the panel
    # updates in real time while the workflow executes further down the page.
    _board_live_detail: dict = {}

    def _render_board_detail(agent: str) -> None:
        slots = _board_live_detail.get(agent)
        if not slots:
            return
        panel = st.session_state.agent_board[agent]
        _acolor = {a: c for a, _i, _n, c in _AGENTS}.get(agent, "#00D4FF")
        if panel["thinking"]:
            # Chat-style live transcript: one row per thinking line, newest at the
            # bottom. Show the recent tail so the latest activity is always in view
            # (Streamlit strips <script>, so we can't auto-scroll a tall box).
            _rows = []
            for _ln in panel["thinking"][-16:]:
                _ts, _msg = "", _ln
                if _ln.startswith("[") and "]" in _ln:
                    _ts = _ln[1:_ln.index("]")]
                    _msg = _ln[_ln.index("]") + 1:].strip()
                _esc = (_msg.replace("&", "&amp;").replace("<", "&lt;")
                            .replace(">", "&gt;"))
                _rows.append(
                    f'<div style="display:flex;gap:9px;padding:5px 0;'
                    f'border-bottom:1px solid #0e1826">'
                    f'<span style="font-family:var(--mono);font-size:0.55rem;'
                    f'color:var(--faint);flex-shrink:0;min-width:52px">{_ts}</span>'
                    f'<span style="font-size:0.73rem;color:var(--text);'
                    f'line-height:1.45">{_esc}</span></div>')
            slots["think"].markdown(
                f'<div style="max-height:300px;overflow-y:auto;'
                f'padding:4px 12px;background:#070d16;border:1px solid {_acolor}33;'
                f'border-radius:9px;border-left:3px solid {_acolor}">'
                + "".join(_rows)
                + '</div>',
                unsafe_allow_html=True)
        else:
            slots["think"].caption("No activity yet — run a triage to see "
                                   "this agent think.")
        if panel["output"]:
            slots["out"].markdown(panel["output"], unsafe_allow_html=True)
        else:
            slots["out"].caption("No output yet.")

    # ── Generated Files: real downloadable exports for the Reporting Agent ──
    # Replaces the raw filesystem paths that used to be printed in the Output
    # panel with actual st.download_button controls, wired to the paths the
    # reporting agent's export step already returns in document_exports.
    def _saved_label(path: str | None) -> str | None:
        if not path:
            return None
        try:
            p = Path(str(path))
            if not p.exists():
                return None
            ts = datetime.fromtimestamp(p.stat().st_mtime)
        except Exception:
            return None
        hh = ts.strftime("%#I:%M %p") if os.name == "nt" else ts.strftime("%-I:%M %p")
        if ts.date() == datetime.now().date():
            return f"Generated today at {hh}"
        return f"Generated on {ts.strftime('%b %d, %Y')} at {hh}"

    def _gf_badge(text: str, color: str = "#A78BFA") -> str:
        return (f'<span style="background:{color}1A;color:{color};'
                f'border:1px solid {color}44;padding:1px 8px;border-radius:10px;'
                f'font-size:0.6rem;font-weight:600;margin-left:8px">{text}</span>')

    def _render_generated_files(panel: dict) -> None:
        exports = panel.get("exports") or {}
        reporting_data = panel.get("reporting_data")
        if not exports and reporting_data is None:
            return   # reporting hasn't produced anything yet — nothing to show

        _inc_suffix = ""
        if isinstance(reporting_data, dict) and reporting_data.get("incident_id"):
            _inc_suffix = "_" + re.sub(r"[^A-Za-z0-9_\-]", "_",
                                       str(reporting_data["incident_id"]))[:40]

        st.markdown(
            '<style>'
            'div[class*="st-key-gfpdf_"] button{'
            'border-color:var(--danger) !important;color:var(--danger) !important;'
            'background:transparent !important;}'
            'div[class*="st-key-gfpdf_"] button:hover{'
            'background:#ff6e7c1A !important;}'
            '</style>', unsafe_allow_html=True)
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.6rem;color:#A78BFA;'
            'letter-spacing:2px;margin:16px 0 4px"> GENERATED FILES</div>'
            '<div style="font-size:0.72rem;color:var(--muted);margin-bottom:10px">'
            'Download the structured stage data or review and export the '
            'formatted report.</div>',
            unsafe_allow_html=True)

        combined_saved = (_saved_label(exports.get("docx"))
                          or _saved_label(exports.get("pdf")))

        _rows = [
            {"id": "reporting_data", "icon": "◧", "icon_bg": "#1E293B",
             "name": "Reporting Data", "badge": None,
             "desc": "Structured data used to generate the reporting documents.",
             "saved": combined_saved if reporting_data is not None else None,
             "kind": "json"},
            {"id": "executive_summary", "icon": "▤", "icon_bg": "#1D4ED8",
             "name": "Executive Summary", "badge": "Draft ready",
             "desc": "High-level overview of the incident and key findings.",
             "saved": _saved_label(exports.get("executive_summary_docx")
                                   or exports.get("executive_summary_pdf")),
             "kind": "section", "section_key": "executive_summary"},
            {"id": "technical_findings", "icon": "⌕", "icon_bg": "#15803D",
             "name": "Technical Findings", "badge": "Draft ready",
             "desc": "Detailed technical analysis, evidence, and indicators.",
             "saved": _saved_label(exports.get("technical_findings_docx")
                                   or exports.get("technical_findings_pdf")),
             "kind": "section", "section_key": "technical_findings"},
            {"id": "soc_analyst_review", "icon": "✓", "icon_bg": "#7E22CE",
             "name": "SOC Analyst Review", "badge": "Draft ready",
             "desc": "Analyst assessment, decisions and recommendations.",
             "saved": _saved_label(exports.get("soc_analyst_review_docx")
                                   or exports.get("soc_analyst_review_pdf")),
             "kind": "section", "section_key": "soc_analyst_review"},
            {"id": "final_incident_report", "icon": "✦", "icon_bg": "#B45309",
             "name": "Final Incident Report", "badge": "Draft ready",
             "desc": "Comprehensive report combining all approved sections.",
             "saved": combined_saved,
             "kind": "combined"},
        ]

        with st.container(border=True):
            for _i, _row in enumerate(_rows):
                _c_label, _c_saved, _c_actions = st.columns([3.4, 1.6, 3.4])
                with _c_label:
                    _badge_html = _gf_badge(_row["badge"]) if _row["badge"] else ""
                    st.markdown(
                        f'<div style="display:flex;gap:10px;align-items:flex-start">'
                        f'<div style="width:30px;height:30px;border-radius:7px;'
                        f'background:{_row["icon_bg"]}33;color:{_row["icon_bg"]};'
                        f'display:flex;align-items:center;justify-content:center;'
                        f'font-size:0.85rem;flex-shrink:0">{_row["icon"]}</div>'
                        f'<div><div style="font-size:0.82rem;font-weight:600;'
                        f'color:var(--text)">{_row["name"]}{_badge_html}</div>'
                        f'<div style="font-size:0.68rem;color:var(--muted);'
                        f'margin-top:2px;max-width:320px">{_row["desc"]}</div>'
                        f'</div></div>',
                        unsafe_allow_html=True)
                with _c_saved:
                    st.markdown(
                        f'<div style="font-size:0.68rem;color:var(--faint);'
                        f'padding-top:6px">{_row["saved"] or "Not generated yet"}'
                        f'</div>', unsafe_allow_html=True)
                with _c_actions:
                    if _row["kind"] == "json":
                        if reporting_data is not None:
                            st.download_button(
                                "Download JSON",
                                data=_json.dumps(reporting_data, indent=2,
                                                 default=str).encode("utf-8"),
                                file_name=f"reporting_data{_inc_suffix}.json",
                                mime="application/json",
                                key=f"gf_json_{_i}", use_container_width=True)
                        else:
                            st.button("Download JSON", disabled=True,
                                     key=f"gf_json_dis_{_i}",
                                     use_container_width=True)
                    else:
                        _b1, _b2, _b3 = st.columns(3)
                        _b1.button("Open & Edit", type="primary", disabled=True,
                                  key=f"gf_edit_{_i}", use_container_width=True,
                                  help="Section editing lives in the Reporting "
                                       "Agent's own dashboard — not yet wired "
                                       "into this view.")
                        if _row["kind"] == "combined":
                            _docx_path, _pdf_path = exports.get("docx"), exports.get("pdf")
                        else:
                            _sk = _row["section_key"]
                            _docx_path = exports.get(f"{_sk}_docx")
                            _pdf_path = exports.get(f"{_sk}_pdf")
                        if _docx_path and Path(str(_docx_path)).exists():
                            _b2.download_button(
                                "Export Word",
                                data=Path(str(_docx_path)).read_bytes(),
                                file_name=f"{_row['id']}{_inc_suffix}.docx",
                                mime=("application/vnd.openxmlformats-officedocument"
                                      ".wordprocessingml.document"),
                                key=f"gf_docx_{_i}", use_container_width=True)
                        else:
                            _b2.button("Export Word", disabled=True,
                                      key=f"gf_docx_dis_{_i}", use_container_width=True)
                        if _pdf_path and Path(str(_pdf_path)).exists():
                            with _b3:
                                with st.container(key=f"gfpdf_{_i}"):
                                    st.download_button(
                                        "Export PDF",
                                        data=Path(str(_pdf_path)).read_bytes(),
                                        file_name=f"{_row['id']}{_inc_suffix}.pdf",
                                        mime="application/pdf",
                                        key=f"gf_pdf_{_i}", use_container_width=True)
                        else:
                            _b3.button("Export PDF", disabled=True,
                                      key=f"gf_pdf_dis_{_i}", use_container_width=True)
                if _i < len(_rows) - 1:
                    st.markdown('<hr style="border-color:var(--border);opacity:0.4;'
                               'margin:6px 0">', unsafe_allow_html=True)

    # Auto-follow: surface the currently-running agent's live thinking without a
    # manual click (keeps the "live LLM chat" feel as work moves between agents).
    if st.session_state.agent_board_sel is None:
        for _ag2, _i2, _n2, _c2 in _AGENTS:
            if st.session_state.agent_board[_ag2]["status"] == "running":
                st.session_state.agent_board_sel = _ag2
                break

    _sel_ag = st.session_state.agent_board_sel
    if _sel_ag:
        _panel = st.session_state.agent_board[_sel_ag]
        _icon, _name, _color = {a: (i, n, c) for a, i, n, c in _AGENTS}[_sel_ag]
        st.markdown(
            f'<div style="font-family:var(--mono);font-size:0.6rem;color:{_color};'
            f'letter-spacing:2px;margin:10px 0 4px">{_icon} {_name.upper()} — DETAIL</div>',
            unsafe_allow_html=True,
        )
        with st.expander("Thinking process", expanded=True):
            _think_slot = st.empty()
            _token_slot = st.empty()   # live LLM token stream (triage phases)
        with st.expander("Output", expanded=bool(_panel["output"])):
            _out_slot = st.empty()
        _board_live_detail[_sel_ag] = {"think": _think_slot,
                                       "token": _token_slot, "out": _out_slot}
        _render_board_detail(_sel_ag)
        if _sel_ag == "reporting":
            _render_generated_files(_panel)

    # ── File uploader (kept — tucked below the board) ──────────────────────
    with st.expander("Upload incident file (JSON · CSV · TXT · LOG)"):
        uploaded_file = st.file_uploader(
            "Upload incident file",
            type=["json", "csv", "txt", "log"],
            label_visibility="collapsed",
            help="Upload a JSON export, CSV log, or plain-text log file to use as incident context.",
        )

        if uploaded_file is not None:
            # Only re-parse if it's a new file
            if uploaded_file.name != st.session_state.uploaded_filename:
                parsed_inc, err = _parse_uploaded_file(uploaded_file)
                if err:
                    st.error(f"{err}")
                else:
                    st.session_state.uploaded_incident = parsed_inc
                    st.session_state.uploaded_filename  = uploaded_file.name
                    st.session_state.chat_incident      = None   # clear NW incident
                    st.rerun()

        # Preview of parsed file fields
        if st.session_state.uploaded_incident and not nw_inc:
            preview_inc = st.session_state.uploaded_incident
            preview_keys = [k for k in preview_inc if not k.startswith("_")][:12]
            preview_html = " &nbsp;·&nbsp; ".join(
                f'<span style="color:var(--accent)">{k}</span>'
                f':<span style="color:var(--text)"> '
                f'{str(preview_inc[k])[:40]}</span>'
                for k in preview_keys
            )
            st.markdown(
                f'<div style="background:#060C16;border:1px solid var(--border);'
                f'border-radius:5px;padding:8px 12px;font-family:var(--mono);'
                f'font-size:0.62rem;margin:6px 0;line-height:1.8">'
                f'Parsed fields: {preview_html}'
                f'</div>',
                unsafe_allow_html=True,
            )
            # Multi-source triage verdict when the upload was normalized from a
            # SIEM/EDR/NDR/log alert (defensive-security skill's analyze_alert)
            _av = preview_inc.get("_analyze_alert")
            if isinstance(_av, dict) and _av.get("classification") not in (None, "invalid"):
                _tp = "true positive" if _av["is_true_positive"] else "needs review"
                st.markdown(
                    f'<div style="background:#0A1A10;border:1px solid #1E5A38;'
                    f'border-radius:5px;padding:8px 12px;font-family:var(--mono);'
                    f'font-size:0.62rem;margin:6px 0;line-height:1.7">'
                    f'<b>Alert triage</b> ({preview_inc.get("_source_format","?").upper()} '
                    f'source): {_av["classification"]} · '
                    f'severity <b>{_av["severity"]}</b> · {_tp}'
                    + (" · MITRE " + ", ".join(m["technique"] for m in _av.get("mitre", [])[:4])
                       if _av.get("mitre") else "")
                    + '</div>',
                    unsafe_allow_html=True,
                )
            # (checkbox, not an expander — Streamlit forbids nesting expanders)
            if st.checkbox("View full parsed incident JSON", key="upl_json_view"):
                st.json({k: v for k, v in preview_inc.items() if k != "raw_log"})

    st.markdown("---")

    # ── Trigger hint ───────────────────────────────────────────
    # Leads with the plain-language instruction — the panel above promises
    # "no technical knowledge needed", so the primary call to action
    # shouldn't be a list of CLI-style keywords. The keywords still work
    # (power-user shortcut, matched by _TRIAGE_TRIGGER below) but are now a
    # smaller secondary line instead of the headline instruction.
    if active_inc:
        st.markdown(
            '<div style="font-size:0.78rem;color:var(--text);margin-bottom:2px">'
            'Select an incident above, or upload a file below, then press '
            '<strong style="color:var(--accent)">Triage</strong> — or just describe '
            'what you want in your own words in the box below.</div>'
            '<div style="font-family:var(--mono);font-size:0.58rem;color:var(--muted);'
            'margin-bottom:6px">Shortcut: typing '
            '<strong style="color:var(--accent)">triage</strong>, '
            '<strong style="color:var(--accent)">ioc</strong>, '
            '<strong style="color:var(--accent)">classify</strong> or '
            '<strong style="color:var(--accent)">ticket</strong> also runs the full pipeline.'
            '</div>',
            unsafe_allow_html=True,
        )

    # ── Chat input  (called first — always fixed at page bottom by Streamlit)
    user_input = st.chat_input("Ask the SOC agent…")

    # The "Triage" button sets this flag instead of the user typing a
    # message — synthesize the trigger word so it flows through the exact
    # same pipeline below, with no manual typing required.
    _triage_restart = False
    if not user_input and st.session_state.pending_auto_triage and active_inc:
        st.session_state.pending_auto_triage = False
        user_input = "Triage this incident"

    # Interaction-interrupt recovery: triage runs INLINE (kept there for live
    # token streaming), so ANY click mid-run — e.g. opening the board's View
    # panel to watch the thinking — kills the script run silently and the
    # board freezes at "Triage started" forever (observed live 2026-07-17
    # 20:30). The in-flight marker set when triage starts survives into this
    # rerun; if it's still here, the previous run died mid-triage → restart.
    elif not user_input and isinstance(st.session_state.get("triage_in_flight"), dict) and active_inc:
        _tif = st.session_state.triage_in_flight
        if str(active_inc.get("id") or "") != _tif.get("incident_id"):
            st.session_state.triage_in_flight = None   # different incident — stale marker
        elif _tif.get("attempts", 1) < 3:
            _tif["attempts"] = _tif.get("attempts", 1) + 1
            user_input = "Triage this incident"
            _triage_restart = True
        else:
            st.session_state.triage_in_flight = None
            _board_set("triage", status="failed",
                       think="triage was interrupted repeatedly — click Triage to run it again",
                       output="Triage did not complete: the run was interrupted "
                              "repeatedly by page interactions. Click Triage to rerun.")

    # Append user message immediately so it shows in the history render below
    # (not on auto-restart — the original message is already in the history)
    if user_input and not _triage_restart:
        now = datetime.now().strftime("%H:%M:%S")
        st.session_state.chat_history.append(
            {"role": "user", "content": user_input, "ts": now}
        )

    # ── Chat history ───────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:8px">■ CONVERSATION</div>',
        unsafe_allow_html=True,
    )

    if not st.session_state.chat_history:
        st.markdown(
            '<div style="text-align:center;padding:40px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            'SOC TRIAGE AGENT READY<br>'
            '<span style="font-size:0.62rem">'
            'Upload a file or select an incident, then type a message below</span></div>',
            unsafe_allow_html=True,
        )

    for msg in st.session_state.chat_history:
        ts = msg.get("ts", "")
        if msg["role"] == "user":
            st.markdown(
                f'<div class="bubble-user">'
                f'<div class="bubble-label" style="color:var(--muted)">YOU · {ts}</div>'
                f'{msg["content"]}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div class="bubble-agent">'
                f'<div class="bubble-label">SOC AGENT · {ts}</div>'
                f'{msg["content"]}</div>',
                unsafe_allow_html=True,
            )

    # ── Thinking containers — in conversation flow, above fixed input ──────────
    # ── Agent execution ────────────────────────────────────────────────────────
    if user_input:

        reply = "No response was generated."   # safe default

        # ── Triage trigger ─────────────────────────────────────────────────────
        # Calls the SAME shared helper as the Start Process button — there is
        # exactly one implementation of "run Parsing+Triage", used by both
        # triggers. Does not start Investigation; the analyst must approve
        # the Triage result first (see the mandatory approval gate in
        # soc_workflow.run_until_triage_approval()).
        if active_inc and _TRIAGE_TRIGGER.search(user_input):
            _inc_id = str(active_inc.get("id") or active_inc.get("incidentId") or "")
            result = _run_triage_workflow_with_ui(active_inc)

            if result is None:
                existing = wss_get_state(_inc_id)
                _tri = (_json.loads(existing["triage_result_json"])
                       if existing and existing.get("triage_result_json") else None)
                if _tri:
                    reply = (format_ticket_display(_tri["ticket"])
                            + f"\n\n---\n\n**Status: {existing.get('workflow_status')}** "
                              f"(approval stage: {existing.get('approval_stage') or '—'})")
                else:
                    reply = (f"Triage is already "
                            f"{existing.get('workflow_status') if existing else 'in progress'} "
                            f"for this incident.")
            else:
                _err = result["errors"].get("parsing") or result["errors"].get("triage")
                if _err:
                    reply = f"Workflow error: {_err}"
                else:
                    _tri = result["triage"]
                    reply = (render_triage_trace(_tri.get("trace") or [])
                            + "\n" + format_ticket_display(_tri["ticket"])
                            + "\n\n---\n\n**Awaiting SOC Analyst approval.** Review "
                              "the ticket above, then approve it from the case view "
                              "before Investigation can run.")

        # ── Plain Q&A fallback ─────────────────────────────────────────────────
        else:
            try:
                with st.spinner("Agent thinking…"):
                    reply = chat_respond(user_input, incident=active_inc)
                if not reply:
                    reply = "Agent returned an empty response."
            except Exception as exc:
                reply = f"Error: {exc}"

        st.session_state.chat_history.append(
            {"role": "assistant", "content": reply,
             "ts": datetime.now().strftime("%H:%M:%S")}
        )
        st.rerun()

    if st.session_state.chat_history:
        if st.button("Clear Chat", key="clear_chat"):
            st.session_state.chat_history = []
            st.rerun()



# ─────────────────────────────────────────────────────────────
# PAGE 3 — INCIDENTS (permanent SQLite store)
# ─────────────────────────────────────────────────────────────
elif active_page == "All Cases":
    st.markdown(
        '<div class="info-box"><div class="title">All Cases</div>'
        'Every case ever loaded, kept permanently, even after a restart. '
        'Search, filter, export to CSV, or open any case in My Workspace.</div>',
        unsafe_allow_html=True,
    )
    stats = db_stats()

    # ── Summary ── operational, not cumulative ────────────────────────────────────────────
    # These used to be all-time counters ("High (ever): 48,718") with no
    # bearing on what needs attention today. Swapped for three questions an
    # analyst actually asks: what's critical right now, what has no owner,
    # and is anything sitting in front of me waiting on a decision.
    _awaiting = (_workflow_store().get("run") or {}).get("awaiting")
    st.markdown(_ui.attention_row([
        {"label": "Critical cases",
         "value": stats["by_sev_active"].get("CRITICAL", 0),
         "sub": "Requires immediate attention", "tone": "red"},
        {"label": "Unassigned cases",
         "value": stats["unassigned_active"],
         "sub": "Waiting for an owner", "tone": "amber"},
        {"label": "Awaiting your approval",
         "value": 1 if _awaiting else 0,
         "sub": "Awaiting SOC analyst approval" if _awaiting else "Nothing pending",
         "tone": "blue" if _awaiting else "green"},
    ]), unsafe_allow_html=True)
    st.caption(f"{stats['total']:,} cases logged all-time · "
              f"{stats['fetches']:,} fetches")

    # ── Filters ───────────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:10px">■ FILTER & SEARCH</div>',
        unsafe_allow_html=True,
    )
    hf1, hf2, hf3, hf4 = st.columns([1.2, 1.2, 2, 1])

    hist_sev    = hf1.selectbox("Severity", ["ALL","CRITICAL","HIGH","MEDIUM","LOW"],
                                 key="hist_sev", label_visibility="collapsed")
    # Collect unique statuses from DB for filter
    with db_connect() as _con:
        _statuses = ["ALL"] + [
            r[0] for r in _con.execute(
                "SELECT DISTINCT status FROM incidents ORDER BY status"
            ).fetchall() if r[0]
        ]
    hist_status = hf2.selectbox("Status", _statuses,
                                 key="hist_status", label_visibility="collapsed")
    # Pre-filled once from the top-nav search box (see _render_top_nav),
    # then consumed -- doesn't stick around overriding manual edits below.
    hist_search = hf3.text_input(
        "Search title / ID / assignee",
        value=st.session_state.pop("hist_search_override", ""),
        placeholder="Type to filter…",
        label_visibility="collapsed")
    hist_limit  = hf4.number_input("Limit", 10, 5000, 200,
                                    label_visibility="collapsed")

    # ── Load rows ────────────────────────────────────────────
    log_rows = db_load_incidents(
        severity=hist_sev,
        status=hist_status,
        search=hist_search,
        limit=hist_limit,
    )

    # ── Export button ──────────────────────────────────────
    csv_data = db_export_csv()
    if csv_data:
        st.download_button(
            label="Export all as CSV",
            data=csv_data,
            file_name=f"soc_incidents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
        )

    st.markdown("<div style='margin-top:8px'></div>", unsafe_allow_html=True)

    _render_case_table(
        log_rows, key_prefix="allcase", heading="All cases",
        subheading=f"{len(log_rows)} records shown.")

    # ── Endpoint Diagnostics & Manual Endpoint Config ────────────────────
    # Moved below the case list (was between the filters and the incidents,
    # interrupting the primary "scan the list" path) and gated behind
    # Analyst mode — these are connection/API diagnostics, not something a
    # read-only guest needs to see at all, let alone step over.
    if _is_dev:
        st.markdown("---")
        with st.expander("Endpoint Diagnostics", expanded=False):
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);margin-bottom:8px">'
                'Tests all known NW endpoints with all auth styles to find the working combination. '
                'Click "Use this" on a hit to wire it into the app automatically.</div>',
                unsafe_allow_html=True,
            )

            # Map scanner's auth-style labels → nw_headers() style names
            _AUTH_STYLE_MAP = {
                "NW-Token": "NetWitness-Token",
                "Bearer":   "Bearer",
                "Cookie":   "Cookie",
                "Both":     "Both",
            }

            if st.button("Run Endpoint Scan", use_container_width=False, key="hist_ep_scan"):
                if not st.session_state.nw_token:
                    st.error("Login first.")
                else:
                    host  = st.session_state.nw_host.rstrip("/")
                    token = st.session_state.nw_token
                    eps   = [
                        "/rest/api/incidents",
                        "/rest/api/respond/incidents",
                        "/rest/api/v1/incidents",
                        "/rest/api/v2/incidents",
                        "/rest/api/respond/incidents/list",
                        "/rest/api/incidents/list",
                        "/rest/api/investigation/incidents",
                        "/respond/api/incidents",
                        "/respond/api/v1/incidents",
                        "/respond/api/v2/incidents",
                        "/rsa/investigation/incidents",
                        "/rsa/respond/incidents",
                        "/sa/api/incidents",
                        "/esa/api/incidents",
                        "/api/respond/incidents",
                        "/rest/api/alerting/incidents",
                        "/rest/api/incidents?start=0",
                    ]
                    auth_styles = {
                        "Bearer":     {"Authorization": f"Bearer {token}"},
                        "Cookie":     {"Cookie": f"access_token={token}"},
                        "NW-Token":   {"NetWitness-Token": token},
                        "Both":       {"Authorization": f"Bearer {token}", "Cookie": f"access_token={token}"},
                    }
                    results = []
                    for ep in eps:
                        for style, ah in auth_styles.items():
                            try:
                                ah = dict(ah)  # avoid mutating shared dict across iterations
                                ah["Accept"] = "application/json"
                                r = requests.get(f"{host}{ep}?limit=1", headers=ah,
                                                 timeout=10, verify=False)
                                ct   = r.headers.get("Content-Type","")
                                is_j = "json" in ct or r.text.strip().startswith(("{","["))
                                results.append({
                                    "endpoint": ep, "auth": style,
                                    "status": r.status_code,
                                    "json": "JSON" if is_j else "HTML",
                                    "is_hit": is_j and r.status_code == 200,
                                    "preview": r.text[:80] if is_j else "",
                                })
                            except Exception as e:
                                results.append({"endpoint": ep, "auth": style,
                                                "status": "ERR", "json": str(e)[:50],
                                                "is_hit": False, "preview": ""})
                    # Persist across reruns (so "Use this" buttons survive)
                    st.session_state.endpoint_scan_results = results

            results = st.session_state.get("endpoint_scan_results", [])
            if results:
                for i, res in enumerate(results):
                    color   = "#00E676" if res.get("is_hit") else "#3A607A"
                    preview = res["preview"]
                    preview_html = (
                        f'<br><span style="color:#00E676">{preview}</span>'
                        if preview else ""
                    )
                    rc1, rc2 = st.columns([6, 1])
                    with rc1:
                        st.markdown(
                            f'<div style="font-family:var(--mono);font-size:0.65rem;'
                            f'padding:3px 8px;border-left:2px solid {color};margin:2px 0">'
                            f'<span style="color:{color}">{res["json"]}</span> '
                            f'HTTP {res["status"]} | {res["auth"]} | {res["endpoint"]}'
                            f'{preview_html}'
                            f'</div>',
                            unsafe_allow_html=True,
                        )
                    with rc2:
                        if res.get("is_hit"):
                            if st.button("Use this", key=f"hist_use_ep_{i}", use_container_width=True):
                                # Wire the found endpoint + auth style into the app
                                clean_path = res["endpoint"].split("?")[0]
                                st.session_state.nw_incidents_path = clean_path
                                st.session_state.nw_auth_style     = _AUTH_STYLE_MAP.get(res["auth"], "NetWitness-Token")
                                st.session_state.nw_working_ep      = res["endpoint"]
                                st.session_state.nw_working_auth    = {"style": res["auth"]}

                                ok_v, msg_v = nw_verify_token()
                                st.session_state.nw_verified = ok_v
                                st.session_state.nw_msg      = msg_v
                                if ok_v:
                                    ok_f, items_f, diag_f = nw_fetch_incidents()
                                    if ok_f:
                                        st.session_state.incidents  = items_f
                                        st.session_state.last_fetch = datetime.now()
                                        db_upsert_incidents(items_f)
                                    st.success(f"Applied {clean_path} ({res['auth']}) — {msg_v}")
                                else:
                                    st.error(f"Applied but verify failed: {msg_v}")
                                st.rerun()

            if st.session_state.nw_working_ep:
                st.success(
                    f"Active endpoint: `{st.session_state.nw_incidents_path}` "
                    f"· auth style: `{st.session_state.nw_auth_style}`"
                )

        with st.expander("Manual Endpoint Config"):
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);margin-bottom:6px">'
                'Set these directly if you already know your NW instance\'s working values.</div>',
                unsafe_allow_html=True,
            )
            mc1, mc2 = st.columns(2)
            new_path = mc1.text_input(
                "Incidents path", value=st.session_state.nw_incidents_path, key="hist_manual_inc_path"
            )
            new_style = mc2.selectbox(
                "Auth header style",
                ["NetWitness-Token", "Bearer", "Cookie", "Both"],
                index=["NetWitness-Token", "Bearer", "Cookie", "Both"].index(
                    st.session_state.nw_auth_style
                    if st.session_state.nw_auth_style in ["NetWitness-Token","Bearer","Cookie","Both"]
                    else "NetWitness-Token"
                ),
                key="hist_manual_auth_style",
            )
            if st.button("Apply & Re-verify", key="hist_manual_apply"):
                st.session_state.nw_incidents_path = new_path.strip() or "/rest/api/incidents"
                st.session_state.nw_auth_style     = new_style
                ok_v, msg_v = nw_verify_token()
                st.session_state.nw_verified = ok_v
                st.session_state.nw_msg      = msg_v
                if ok_v:
                    ok_f, items_f, diag_f = nw_fetch_incidents()
                    if ok_f:
                        st.session_state.incidents  = items_f
                        st.session_state.last_fetch = datetime.now()
                        db_upsert_incidents(items_f)
                    st.success(f"{msg_v}")
                else:
                    st.error(f"{msg_v}")
                st.rerun()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — DATA PIPELINE
# Fully inline — single streamlit run app.py, no second process.
# ═══════════════════════════════════════════════════════════════════════════════
elif active_page == "Data Pipeline":
    st.markdown(
        '<div class="info-box"><div class="title"> Data Pipeline</div>'
        'This tool helps IT staff manage how security data flows through the system — '
        'from collection, to review, to archiving. If you are not sure what this is for, '
        'you likely do not need to use it.</div>',
        unsafe_allow_html=True,
    )

    # ── session defaults ──────────────────────────────────────────────────────
    if "pl_stage"      not in st.session_state: st.session_state.pl_stage      = None
    if "pl_cv_results" not in st.session_state: st.session_state.pl_cv_results = []

    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:14px">'
        '■ SOC PIPELINE — CLICK SELECT UNDER ANY STAGE TO OPEN ITS DB VIEWER</div>',
        unsafe_allow_html=True,
    )

    # ══════════════════════════════════════════════════════════════════════════
    # STAGE CARDS  (always visible — one column per stage)
    # ══════════════════════════════════════════════════════════════════════════
    _pl_cols = st.columns(len(PIPELINE_STAGES))
    for _idx, _stage in enumerate(PIPELINE_STAGES):
        _cnt    = pipeline_count(_stage)
        _icon   = PIPELINE_ICONS[_stage]
        _label  = PIPELINE_LABELS[_stage]
        _color  = PIPELINE_COLORS[_stage]
        _active = st.session_state.pl_stage == _stage
        _border = f"3px solid {_color}" if _active else f"1px solid {_color}44"
        _bg     = f"{_color}18"          if _active else "#060C16"
        _last = pipeline_last_write(_stage)
        with _pl_cols[_idx]:
            st.markdown(
                f'<div style="background:{_bg};border:{_border};border-radius:7px;'
                f'padding:10px 8px;text-align:center;transition:all 0.15s">'
                f'<div style="font-size:1.3rem">{_icon}</div>'
                f'<div style="font-family:var(--mono);font-size:1.5rem;'
                f'color:{_color};margin:4px 0">{_cnt}</div>'
                f'<div style="font-family:var(--mono);font-size:var(--fs-3xs);'
                f'color:var(--muted);letter-spacing:1px;line-height:1.4">'
                f'{_label.upper()}</div>'
                f'<div style="font-family:var(--mono);font-size:var(--fs-3xs);'
                f'color:{_color}99;margin-top:3px">last: {_last}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if st.button("Select", key=f"pl_sel_{_stage}", use_container_width=True):
                st.session_state.pl_stage      = None if _active else _stage
                st.session_state.pl_cv_results = []
                st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # INLINE VIEWER  (expands below when a stage is selected)
    # ══════════════════════════════════════════════════════════════════════════
    _sel = st.session_state.pl_stage

    if not _sel:
        st.markdown(
            '<div style="text-align:center;padding:60px 20px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '↑ Click <strong style="color:var(--accent)">Select</strong> '
            'under any stage above to open its database viewer<br>'
            '<span style="font-size:0.62rem;margin-top:8px;display:block">'
            'Records auto-insert when the triage agent runs in the Chat tab.'
            '</span></div>',
            unsafe_allow_html=True,
        )
    else:
        _color = PIPELINE_COLORS[_sel]
        _icon  = PIPELINE_ICONS[_sel]
        _label = PIPELINE_LABELS[_sel]

        st.markdown("<br>", unsafe_allow_html=True)

        # ── header banner ─────────────────────────────────────────────────────
        st.markdown(
            f'<div style="background:#060C16;border:2px solid {_color};'
            f'border-left:8px solid {_color};border-radius:8px;'
            f'padding:18px 24px;margin-bottom:16px">'
            f'<div style="font-family:var(--mono);font-size:1.05rem;'
            f'color:{_color};letter-spacing:3px">{_icon} {_label.upper()}</div>'
            f'<div style="font-family:var(--mono);font-size:0.58rem;'
            f'color:var(--muted);margin-top:6px">'
            f'ChromaDB collection: <code style="color:var(--accent)">pipeline_{_sel}</code>'
            f'&nbsp;·&nbsp;SQLite table: <code style="color:var(--accent)">{_sel}</code>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

        # ── metrics ───────────────────────────────────────────────────────────
        _sql_cnt = pipeline_count(_sel)
        _vec_cnt = pipeline_chroma_count(_sel)
        _vm1, _vm2, _vm3, _vm4 = st.columns(4)
        _vm1.metric("SQLite Records",   _sql_cnt)
        _vm2.metric("ChromaDB Vectors", _vec_cnt)
        _vm3.metric("Stage",            _label[:22])
        _vm4.metric("Collection",       f"pipeline_{_sel}"[:24])

        # ── action strip ──────────────────────────────────────────────────────
        _ac1, _ac2, _ac3, _ac4, _ = st.columns([1, 1.4, 1, 1.1, 2.5])

        if _ac1.button("Test Insert", key=f"pl_add_{_sel}", disabled=not _is_dev,
                       help=None if _is_dev else
                       "Analyst mode required — toggle it on in Settings."):
            import uuid as _uuid3
            pipeline_insert_full(_sel, {
                "id":          f"test_{str(_uuid3.uuid4())[:8]}",
                "incident_id": "DEMO-001",
                "title":       f"[Demo] Test record · {_label}",
                "severity":    "MEDIUM",
                "summary":     f"Auto-inserted test record into: {_label}",
            })
            st.success("Test record inserted!")
            st.rerun()

        if _sel == "pending_ticket_report":
            if _ac2.button("Finalize All", key="pl_finalize_btn"):
                _pending = pipeline_load("pending_ticket_report")
                for _prow in _pending:
                    _fin = dict(_prow)
                    _fin["id"]      = f"final_{_prow['id']}"
                    _fin["title"]   = (_fin.get("title") or "").replace("[PENDING]","[FINAL]")
                    _fin["summary"] = "Finalized. " + (_prow.get("summary") or "")
                    pipeline_insert_full("finalized_report", _fin)
                st.success(f"Finalized {len(_pending)} records → Finalized Report.")
                st.rerun()

        with _ac3:
            if not _is_dev:
                st.button("Clear Stage", key=f"pl_clear_{_sel}", disabled=True,
                         use_container_width=True,
                         help="Analyst mode required — toggle it on in Settings.")
            elif _confirm_action(
                    f"pl_clear_{_sel}", "Clear Stage",
                    f"This permanently deletes **all** records in "
                    f"**{_label}** (SQLite + its ChromaDB vectors). "
                    f"This cannot be undone.",
                    confirm_label="Yes, clear stage"):
                with _pl_con() as _c:
                    _c.execute(f"DELETE FROM {_sel}")
                    _c.commit()
                _wipe_col = _pl_chroma_col(_sel)
                if _wipe_col:
                    try:
                        st.session_state.chroma_client.delete_collection(f"pipeline_{_sel}")
                    except Exception:
                        pass
                st.success(f"Cleared {_label}.")
                st.rerun()

        if _ac4.button("✕ Close Viewer", key="pl_close_viewer"):
            st.session_state.pl_stage      = None
            st.session_state.pl_cv_results = []
            st.rerun()

        st.markdown("---")

        # ── sub-tabs ──────────────────────────────────────────────────────────
        _vtab_sql, _vtab_chroma = st.tabs(["SQLITE RECORDS", "CHROMADB SEARCH"])

        # ════════════════════ SQLite Records ════════════════════
        with _vtab_sql:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:10px">'
                '■ ALL RECORDS IN THIS STAGE</div>',
                unsafe_allow_html=True,
            )
            _fc1, _fc2 = st.columns([5, 1])
            _sql_srch = _fc1.text_input(
                "filter", placeholder="Filter by title / ID / summary…",
                key=f"sql_srch_{_sel}", label_visibility="collapsed")
            _sql_lim = int(_fc2.number_input(
                "limit", 10, 2000, 200,
                key=f"sql_lim_{_sel}", label_visibility="collapsed"))

            _rows = pipeline_load(_sel, limit=_sql_lim)
            if _sql_srch.strip():
                _s = _sql_srch.lower()
                _rows = [r for r in _rows if (
                    _s in (r.get("title")   or "").lower() or
                    _s in (r.get("id")      or "").lower() or
                    _s in (r.get("summary") or "").lower())]

            if not _rows:
                st.markdown(
                    '<div style="text-align:center;padding:50px;font-family:var(--mono);'
                    'font-size:0.78rem;color:var(--muted)">● NO RECORDS FOUND<br>'
                    '<span style="font-size:0.62rem">'
                    'Run a triage in the Chat tab, or click Test Insert above.'
                    '</span></div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-family:var(--mono);font-size:0.6rem;'
                    f'color:var(--muted);margin-bottom:8px">{len(_rows)} records</div>',
                    unsafe_allow_html=True,
                )
                # ── Aegis queue table (My Queue mockup) — compact default view.
                # Detail cards (JSON / CSV / DOCX / delete per record) stay fully
                # available behind the checkbox below; nothing was removed.
                try:
                    _qrows = [[
                        {"mono": str(_r.get("id", "—"))[:28]},
                        str(_r.get("title") or "Untitled")[:70],
                        {"mono": str(_r.get("incident_id") or "—")},
                        {"pill": str(_r.get("severity") or "—").upper(),
                         "kind": _ui.sev_class(_r.get("severity") or "")},
                        {"pill": _label, "kind": "stage"},
                        str(_r.get("created_at") or "")[:16],
                    ] for _r in _rows[:100]]
                    st.markdown(_ui.queue_table(
                        ["Record", "Case", "Incident", "Severity", "Stage", "Created"],
                        _qrows), unsafe_allow_html=True)
                    if len(_rows) > 100:
                        st.caption(f"Table shows the first 100 of {len(_rows)} records — "
                                   "use the filter above or the detail cards below.")
                except Exception:
                    pass
                _show_detail = st.checkbox(
                    "Show detail cards (JSON · exports · delete per record)",
                    value=False, key=f"pl_detail_{_sel}")
                _rows_detail = _rows if _show_detail else []
                for _row in _rows_detail:
                    _r_id  = _row.get("id", "—")
                    _r_ttl = _row.get("title", "Untitled")
                    _r_sev = str(_row.get("severity", "—")).upper()
                    _r_inc = _row.get("incident_id", "—")
                    _r_sum = (_row.get("summary") or "")[:200]
                    _r_ts  = str(_row.get("created_at", ""))[:16]
                    _sev_c = SEV_COLORS.get(_r_sev, "#3A607A")

                    # ── Stage-specific export badge ─────────────────────────
                    _is_csv_stage  = _sel in ("post_triage_investigate", "post_triage_no_investigate")
                    _is_docx_stage = _sel == "initial_ticket"
                    _is_report_stage = _sel == "finalized_report"
                    _is_postinv_stage = _sel == "post_investigation"
                    # Post-investigation records carry the analysis narrative
                    # (markdown) produced by the investigation agent.
                    _postinv_md = ""
                    if _is_postinv_stage:
                        try:
                            _rj_inv = _json.loads(_row.get("raw_json") or "{}")
                            _inv_j  = _rj_inv.get("investigation") or {}
                            _postinv_md = _inv_j.get("narrative_report") or ""
                            if not _postinv_md:
                                _md_path = (_inv_j.get("artifacts") or {}).get("report_markdown")
                                if _md_path and Path(str(_md_path)).exists():
                                    _postinv_md = Path(str(_md_path)).read_text(encoding="utf-8")
                        except Exception:
                            _postinv_md = ""
                    # Finalized reports carry real Word/PDF files generated by
                    # the reporting agent — resolve their paths from raw_json.
                    _report_exports = {}
                    if _is_report_stage:
                        try:
                            _rj_exp = _json.loads(_row.get("raw_json") or "{}")
                            _report_exports = ((_rj_exp.get("report") or {})
                                               .get("document_exports") or {})
                        except Exception:
                            _report_exports = {}
                        _report_exports = {
                            k: v for k, v in _report_exports.items()
                            if k in ("docx", "pdf") and v and Path(str(v)).exists()
                        }
                    _export_badge  = ""
                    if _is_csv_stage:
                        _export_badge = (
                            f'<span style="background:#00403A;color:#0AF0A0;'
                            f'border:1px solid #0AF0A044;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> CSV</span>'
                        )
                    elif _is_docx_stage:
                        _export_badge = (
                            f'<span style="background:#1A0040;color:#A78BFA;'
                            f'border:1px solid #A78BFA44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> DOCX</span>'
                        )
                    elif _is_report_stage and _report_exports:
                        _fmt_txt = " · ".join(k.upper() for k in ("docx", "pdf")
                                              if k in _report_exports)
                        _export_badge = (
                            f'<span style="background:#1A0040;color:#A78BFA;'
                            f'border:1px solid #A78BFA44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> {_fmt_txt}</span>'
                        )
                    elif _is_postinv_stage and _postinv_md:
                        _export_badge = (
                            f'<span style="background:#04342C;color:#2DD4BF;'
                            f'border:1px solid #2DD4BF44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> REPORT</span>'
                        )

                    st.markdown(
                        f'<div style="background:#060C16;border:1px solid {_color}44;'
                        f'border-left:3px solid {_color};border-radius:6px;'
                        f'padding:10px 14px;margin:4px 0">'
                        f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap">'
                        f'<code style="font-family:var(--mono);font-size:0.65rem;color:{_color}">'
                        f'{_r_id}</code>'
                        f'<strong style="flex:1;font-size:0.85rem">{_r_ttl}</strong>'
                        f'{_export_badge}'
                        f'<span style="background:{_sev_c}22;color:{_sev_c};'
                        f'border:1px solid {_sev_c}44;padding:2px 8px;border-radius:3px;'
                        f'font-family:var(--mono);font-size:0.6rem">{_r_sev}</span>'
                        f'<span style="font-family:var(--mono);font-size:0.56rem;'
                        f'color:var(--muted)">{_r_ts}</span>'
                        f'</div>'
                        f'<div style="font-size:0.7rem;color:var(--muted);margin-top:6px">'
                        f'Incident: <code style="color:var(--accent)">{_r_inc}</code>'
                        f'{"&nbsp;·&nbsp;" + _r_sum if _r_sum else ""}'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )

                    # ── Button row: JSON | export | delete ──────────────────
                    _bj2b = None
                    if _is_csv_stage:
                        _bj1, _bj2, _bj3, _ = st.columns([0.6, 1.2, 0.28, 7])
                    elif _is_docx_stage:
                        _bj1, _bj2, _bj3, _ = st.columns([0.6, 1.3, 0.28, 7])
                    elif _is_report_stage and _report_exports:
                        _bj1, _bj2, _bj2b, _bj3, _ = st.columns([0.6, 1.1, 1.1, 0.28, 6])
                    elif _is_postinv_stage and _postinv_md:
                        _bj1, _bj2, _bj2b, _bj3, _ = st.columns([0.6, 1.1, 1.3, 0.28, 6])
                    else:
                        _bj1, _bj3, _ = st.columns([0.6, 0.28, 9])
                        _bj2 = None

                    if _bj1.button("{ } JSON", key=f"pl_json_{_sel}_{_r_id}"):
                        try:
                            _raw_j = _json.loads(_row.get("raw_json") or "{}")
                        except Exception:
                            _raw_j = _row
                        with st.expander(f"JSON — {_r_id}", expanded=True):
                            st.json(_raw_j)

                    # ── CSV download for post-triage stages ─────────────────
                    if _is_csv_stage and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _csv_bytes = _make_csv_bytes(_row)
                        _bj2.download_button(
                            label="View as Sheet",
                            data=_csv_bytes,
                            file_name=f"triage_{_safe_id}.csv",
                            mime="text/csv",
                            key=f"pl_csv_{_sel}_{_r_id}",
                        )

                    # ── DOCX download for initial_ticket stage ──────────────
                    elif _is_docx_stage and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _docx_bytes = _make_docx_bytes(_row)
                        # Detect if python-docx is available (bytes start with PK zip magic)
                        _is_real_docx = _docx_bytes[:2] == b'PK'
                        _bj2.download_button(
                            label="View as Ticket",
                            data=_docx_bytes,
                            file_name=f"ticket_{_safe_id}.{'docx' if _is_real_docx else 'txt'}",
                            mime=(
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document"
                                if _is_real_docx else "text/plain"
                            ),
                            key=f"pl_docx_{_sel}_{_r_id}",
                        )

                    # ── Word/PDF downloads for finalized_report stage ───────
                    # Serves the actual documents generated by the reporting
                    # agent (outputs/<incident>/reports/exports/), not a
                    # reconstruction from the pipeline record.
                    elif _is_report_stage and _report_exports and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        if _report_exports.get("docx"):
                            try:
                                _bj2.download_button(
                                    label="Word Report",
                                    data=Path(str(_report_exports["docx"])).read_bytes(),
                                    file_name=f"incident_report_{_safe_id}.docx",
                                    mime=("application/vnd.openxmlformats-officedocument"
                                          ".wordprocessingml.document"),
                                    key=f"pl_repdocx_{_sel}_{_r_id}",
                                )
                            except Exception:
                                pass
                        if _report_exports.get("pdf") and _bj2b is not None:
                            try:
                                _bj2b.download_button(
                                    label="PDF Report",
                                    data=Path(str(_report_exports["pdf"])).read_bytes(),
                                    file_name=f"incident_report_{_safe_id}.pdf",
                                    mime="application/pdf",
                                    key=f"pl_reppdf_{_sel}_{_r_id}",
                                )
                            except Exception:
                                pass

                    # ── Investigation findings for post_investigation stage ──
                    elif _is_postinv_stage and _postinv_md and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _bj2.download_button(
                            label="Report (MD)",
                            data=_postinv_md.encode("utf-8"),
                            file_name=f"investigation_{_safe_id}.md",
                            mime="text/markdown",
                            key=f"pl_invmd_{_sel}_{_r_id}",
                        )
                        if _bj2b is not None and _bj2b.button(
                                "View Findings", key=f"pl_invview_{_sel}_{_r_id}"):
                            with st.expander(f"Investigation findings — {_r_id}",
                                             expanded=True):
                                # ✦ AI-summary card (Phase 4a) — flags an
                                # error/fallback report with the amber tag so a
                                # failed LLM run is visible at a glance.
                                try:
                                    _fb = _ui.detect_fallback(_postinv_md)
                                    _first = next((ln.strip() for ln in _postinv_md.splitlines()
                                                   if ln.strip() and not ln.lstrip().startswith("#")),
                                                  "Investigation report")
                                    st.markdown(_ui.ai_summary(_first[:400], _fb,
                                                title="Investigation report"),
                                                unsafe_allow_html=True)
                                except Exception:
                                    pass
                                st.markdown(_postinv_md)

                    with _bj3:
                        if not _is_dev:
                            st.button("✕", key=f"pl_del_{_sel}_{_r_id}", disabled=True,
                                     help="Analyst mode required — toggle it "
                                          "on in Settings.")
                        elif _confirm_action(
                                f"pl_del_{_sel}_{_r_id}", "✕",
                                f"Permanently delete record **{_r_id}** from "
                                f"**{_label}**? This cannot be undone.",
                                confirm_label="Delete", use_container_width=False):
                            pipeline_delete(_sel, _r_id)
                            st.rerun()

        # ════════════════════ ChromaDB Search ════════════════════
        with _vtab_chroma:
            _chroma_ok = CHROMA_OK and bool(st.session_state.get("chroma_client"))

            if not _chroma_ok:
                st.markdown(
                    '<div style="background:#0A0608;border:1px solid #2A1010;'
                    'border-radius:6px;padding:12px 16px;font-family:var(--mono);'
                    'font-size:0.68rem;margin-bottom:12px">'
                    '<span style="color:var(--warn)"> ChromaDB not connected</span><br>'
                    '<span style="color:var(--muted);font-size:0.6rem">'
                    'Connect ChromaDB from Settings. SQLite records persist regardless.'
                    '</span></div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-family:var(--mono);font-size:0.65rem;'
                    f'color:var(--muted);margin-bottom:10px">'
                    f'{pipeline_chroma_count(_sel)} vectors in '
                    f'<code style="color:var(--accent)">pipeline_{_sel}</code></div>',
                    unsafe_allow_html=True,
                )

            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ SEMANTIC SEARCH</div>',
                unsafe_allow_html=True,
            )
            _sq2, _sn2 = st.columns([5, 1])
            _cv_q = _sq2.text_input(
                "query", key=f"cv_q_{_sel}",
                placeholder="e.g. ransomware lateral movement C2",
                label_visibility="collapsed")
            _cv_n = int(_sn2.number_input(
                "topn", 1, 20, 5,
                key=f"cv_n_{_sel}", label_visibility="collapsed"))

            if st.button("Search", key=f"cv_srch_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from Settings first.")
                elif not _cv_q.strip():
                    st.warning("Enter a query.")
                else:
                    st.session_state.pl_cv_results = pipeline_chroma_search(
                        _sel, _cv_q, n=_cv_n)

            for _r in st.session_state.pl_cv_results:
                _rm   = _r["meta"]
                _scol = PIPELINE_COLORS.get(_rm.get("stage", _sel), _color)
                st.markdown(
                    f'<div style="background:#060C16;border:1px solid {_scol}55;'
                    f'border-left:3px solid {_scol};border-radius:6px;'
                    f'padding:10px 14px;margin:4px 0">'
                    f'<div style="display:flex;align-items:center;gap:10px">'
                    f'<span style="background:{_scol}22;color:{_scol};'
                    f'border:1px solid {_scol}44;padding:2px 8px;border-radius:3px;'
                    f'font-family:var(--mono);font-size:0.6rem">{_r["score"]}%</span>'
                    f'<strong style="flex:1">{_r["id"]}</strong>'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:var(--muted)">'
                    f'sev:{_rm.get("severity","?")} · {_rm.get("created","")[:10]}'
                    f'</span></div>'
                    f'<div style="font-size:0.72rem;color:var(--text);margin-top:6px">'
                    f'{str(_r["doc"])[:300]}</div></div>',
                    unsafe_allow_html=True,
                )

            st.markdown("---")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ BROWSE ALL VECTORS</div>',
                unsafe_allow_html=True,
            )
            if st.button("Load All Vectors", key=f"cv_all_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from Settings first.")
                else:
                    _all_v = pipeline_chroma_all(_sel)
                    if not _all_v:
                        st.info("No vectors in this collection yet.")
                    for _v in _all_v:
                        _vm  = _v["meta"]
                        _vc  = PIPELINE_COLORS.get(_vm.get("stage", _sel), _color)
                        st.markdown(
                            f'<div style="background:#060C16;border:1px solid {_vc}44;'
                            f'border-left:3px solid {_vc};border-radius:6px;'
                            f'padding:8px 14px;margin:3px 0">'
                            f'<div style="display:flex;gap:10px;align-items:center">'
                            f'<code style="font-family:var(--mono);font-size:0.62rem;'
                            f'color:{_vc}">{_v["id"]}</code>'
                            f'<span style="font-family:var(--mono);font-size:0.56rem;'
                            f'color:var(--muted)">'
                            f'sev:{_vm.get("severity","?")} · {_vm.get("created","")[:10]}'
                            f'</span></div>'
                            f'<div style="font-size:0.7rem;color:var(--text);margin-top:5px">'
                            f'{str(_v["doc"])[:250]}</div></div>',
                            unsafe_allow_html=True,
                        )

            st.markdown("---")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ COLLECTION ACTIONS</div>',
                unsafe_allow_html=True,
            )
            if st.button(f"Wipe Chroma: pipeline_{_sel}", key=f"cv_wipe_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from Settings first.")
                else:
                    try:
                        st.session_state.chroma_client.delete_collection(f"pipeline_{_sel}")
                        _pl_chroma_col(_sel)
                        st.success(f"Wiped pipeline_{_sel}.")
                        st.rerun()
                    except Exception as _we:
                        st.error(str(_we))

elif active_page == "Settings":
    st.markdown(
        '<div class="info-box"><div class="title">Settings</div>'
        'Connection, security certificate, and AI settings -- previously in '
        'the sidebar, moved here when the app switched to a top navigation '
        'bar.</div>',
        unsafe_allow_html=True,
    )
    # Aegis workspace card (mockup .workspace) — analyst identity at the top
    # of the sidebar. Guarded: the sidebar renders before the main
    # `import ui_components` below, so import locally here.
    #
    # analyst_display_name is separate from nw_username: nothing downstream
    # previously recorded WHO approved/rejected a hand-off (see the Agent
    # Board below) — this is the minimum viable fix without building a full
    # login system. Seeded from nw_username once, then editable.
    st.session_state.setdefault("analyst_display_name", "")
    if not st.session_state.analyst_display_name and st.session_state.get("nw_username"):
        st.session_state.analyst_display_name = st.session_state.nw_username
    try:
        import ui_components as _uisb
        _ws_user = (st.session_state.analyst_display_name or "").strip() or "not signed in"
        st.markdown(_uisb.workspace_card(
            "SOC Workspace", f"Analyst · {_ws_user}", ""),
            unsafe_allow_html=True)
    except Exception:
        pass
    st.session_state.analyst_display_name = st.text_input(
        "Your name — attributed on every approve/reject",
        value=st.session_state.analyst_display_name,
        placeholder="e.g. J. Tan",
        key="analyst_name_input",
    ).strip()

    # Previously this label was hard-coded to "Guest · read-only view" with
    # no control anywhere that could ever flip it to "Developer mode" — the
    # role never changed, and worse, several destructive tools (Clear Stage,
    # per-record delete, the endpoint scanner) weren't actually gated by it.
    # This toggle makes the label true: off = genuinely read-only.
    st.toggle("Analyst mode", key="analyst_mode_on",
              help="Off: read-only — you can view everything but can't "
                   "delete pipeline data, run the endpoint scanner, or "
                   "change connection/LLM settings. On: unlocks those "
                   "tools. No password — this is a single-operator tool, "
                   "the toggle just makes the role label accurate.")
    st.caption("Analyst mode — connection settings and destructive tools "
               "unlocked" if _is_dev else "Guest · read-only view")
    st.markdown("---")

    if _is_dev:
        # ── Connection status card ─────────────────────────────────
        if st.session_state.nw_verified:
            last      = st.session_state.last_fetch
            last_str  = last.strftime("%H:%M:%S") if last else "—"
            elapsed   = int((datetime.now() - last).total_seconds()) if last else 0
            remaining = max(REFRESH_INTERVAL - elapsed, 0)
            pct       = min(elapsed / REFRESH_INTERVAL, 1.0)
            bar_color = "var(--accent)" if pct < 0.8 else "var(--warn)"

            # Check GP status via the existing connection state
            _gp_title = ('title="GlobalProtect — the corporate VPN required to '
                        'reach the on-prem NetWitness server"')
            _gp_status = (
                f'<span class="dot dot-green"></span>'
                f'<span {_gp_title} style="color:var(--green);font-size:0.58rem;'
                f'text-decoration:underline dotted;text-decoration-color:var(--green);'
                f'cursor:help">GP VPN ACTIVE</span>'
                if st.session_state.nw_verified else
                f'<span class="dot dot-yellow"></span>'
                f'<span {_gp_title} style="color:var(--warn);font-size:0.58rem;'
                f'text-decoration:underline dotted;text-decoration-color:var(--warn);'
                f'cursor:help">GP VPN STATUS UNKNOWN</span>'
            )

            # Aegis shift-card look (mockup .shift): card bg + line border + r12
            st.markdown(
                f'<div style="background:#0e182a;border:1px solid #223149;'
                f'border-radius:12px;padding:13px">'
                f'<div style="font-family:var(--mono);font-size:0.68rem">'
                f'<span class="dot dot-green"></span>Connected ✓</div>'
                f'<div style="margin-top:4px">{_gp_status}</div>'
                f'<div style="font-family:var(--mono);font-size:0.58rem;'
                f'color:var(--muted);margin-top:3px">{st.session_state.nw_msg}</div>'
                f'<div style="font-family:var(--mono);font-size:0.56rem;'
                f'color:#1A4A62;margin-top:6px">'
                f'Synced {last_str} &nbsp;·&nbsp; refresh in {remaining}s</div>'
                f'<div style="font-family:var(--mono);font-size:0.55rem;'
                f'color:#2A5A78;margin-top:4px">'
                f'{st.session_state.nw_incidents_path} · {st.session_state.nw_auth_style}</div>'
                f'<div style="background:#060E1A;border-radius:2px;height:3px;'
                f'width:100%;margin-top:7px;overflow:hidden">'
                f'<div style="width:{pct*100:.0f}%;height:100%;border-radius:2px;'
                f'background:{bar_color};transition:width 1s linear"></div>'
                f'</div></div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#0A0608;border:1px solid #2A1010;'
                'border-radius:7px;padding:11px 13px;font-family:var(--mono);font-size:0.68rem">'
                '<span class="dot dot-red"></span>Not Connected<br>'
                '<span style="font-size:0.57rem;color:var(--muted)">'
                'Please enter your login details below</span></div>',
                unsafe_allow_html=True,
            )

        # ── NetWitness credentials ─────────────────────────────────
        st.markdown('<div class="sec-label">  Connection</div>', unsafe_allow_html=True)

        # If already auto-connected from .env, show a clean status + update option
        if st.session_state.nw_verified and _env["username"]:
            st.markdown(
                '<div style="background:#10231c;border:1px solid #2a6146;border-radius:10px;'
                'padding:9px 12px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
                '<span class="dot dot-green"></span>'
                '<strong style="color:var(--green)">AUTO-CONNECTED FROM .ENV</strong><br>'
                '<span style="color:var(--muted);font-size:0.58rem">'
                f'Logged in as <strong>{st.session_state.nw_username}</strong> · '
                'token refreshed on startup.</span>'
                '</div>',
                unsafe_allow_html=True,
            )
            with st.expander("Update credentials"):
                host_in = st.text_input("Host URL", value=st.session_state.nw_host, key="sb_host")
                user_in = st.text_input("Username", value=st.session_state.nw_username, key="sb_user")
                pass_in = st.text_input("Password", value="", type="password",
                                         placeholder="Enter password…", key="sb_pass")
                st.markdown(
                    '<div style="font-family:var(--mono);font-size:0.6rem;'
                    'color:var(--muted);margin:6px 0 2px">— or paste token directly —</div>',
                    unsafe_allow_html=True,
                )
                token_paste_in = st.text_area(
                    "Paste Token",
                    value="",
                    placeholder="Paste a fresh accessToken (eyJ…) to skip re-login",
                    height=80,
                    key="sb_token_paste",
                    label_visibility="collapsed",
                )
                cv, cs, cd = st.columns(3)
                if cv.button("Login", use_container_width=True, key="sb_verify"):
                    st.session_state.nw_host     = host_in
                    st.session_state.nw_username = user_in
                    st.session_state.nw_password = pass_in

                    raw_paste = token_paste_in.strip()
                    if raw_paste:
                        # Use pasted token directly
                        st.session_state.nw_token = raw_paste
                        with st.spinner("Verifying token…"):
                            ok, msg = nw_verify_token()
                        st.session_state.nw_verified = ok
                        st.session_state.nw_msg      = msg
                        if not ok:
                            st.session_state.nw_token = ""
                    else:
                        # Fall back to username/password
                        with st.spinner("Logging in…"):
                            ok, msg, token = nw_login(host_in, user_in, pass_in)
                        st.session_state.nw_verified = ok
                        st.session_state.nw_msg      = msg
                        if ok:
                            st.session_state.nw_token = token

                    if st.session_state.nw_verified:
                        ok2, items, _diag = nw_fetch_incidents()
                        if ok2:
                            st.session_state.incidents  = items
                            st.session_state.last_fetch = datetime.now()
                            db_upsert_incidents(items)
                    st.rerun()
                if cs.button("Save", use_container_width=True, key="sb_save"):
                    if pass_in:
                        env_save(host_in, user_in, pass_in)
                        st.success("Saved to .env")
                    else:
                        st.warning("Enter password to save.")
                with cd:
                    if _confirm_action(
                            "sb_clear", "✕ Clear",
                            "This clears your saved NetWitness host, username, "
                            "and password and disconnects the session.",
                            confirm_label="Clear credentials"):
                        env_clear()
                        st.session_state.update(
                            nw_token="", nw_username="", nw_password="",
                            nw_host="", nw_verified=False, nw_msg="",
                            incidents=[], last_fetch=None, _startup_done=False,
                        )
                        st.rerun()
        else:
            # Show last login error if there is one
            _last_err = st.session_state.get("nw_msg", "")
            if _last_err and not st.session_state.nw_verified:
                st.markdown(
                    f'<div style="background:#1A0505;border:1px solid #5A1010;border-radius:5px;'
                    f'padding:8px 11px;font-family:var(--mono);font-size:0.6rem;'
                    f'color:#FF6B6B;margin-bottom:8px">'
                    f'Last error: {_last_err}</div>',
                    unsafe_allow_html=True,
                )
            if DOTENV_OK and not _env["username"]:
                st.markdown(
                    '<div style="background:#0A0800;border:1px solid #3A3000;border-radius:5px;'
                    'padding:8px 11px;font-family:var(--mono);font-size:0.6rem;'
                    'color:#FFB700;margin-bottom:8px">'
                    'No credentials in .env<br>'
                    '<span style="color:var(--muted)">Enter below & click Save</span></div>',
                    unsafe_allow_html=True,
                )
            host_in = st.text_input("Host URL", value=st.session_state.nw_host,
                                     placeholder="https://192.168.x.x")
            if host_in != st.session_state.nw_host:
                st.session_state.nw_host     = host_in
                st.session_state.nw_verified = False

            # ── Login method toggle ────────────────────────────────
            login_method = st.radio(
                "Login method",
                ["Username / Password", "Paste Token"],
                horizontal=True,
                label_visibility="collapsed",
            )

            if login_method == "Username / Password":
                user_in = st.text_input("Username", value=st.session_state.nw_username,
                                         placeholder="admin")
                pass_in = st.text_input("Password", value="", type="password",
                                         placeholder="Enter password…")
                token_in = ""
            else:
                user_in  = st.session_state.nw_username
                pass_in  = ""
                token_in = st.text_area(
                    "NetWitness Token",
                    value="",
                    placeholder="Paste your accessToken (eyJ…) here",
                    height=100,
                    label_visibility="collapsed",
                )
                st.markdown(
                    '<div style="font-family:var(--mono);font-size:0.58rem;'
                    'color:var(--muted);margin-top:-8px;margin-bottom:6px">'
                    'Tokens expire — re-paste when you get a 401</div>',
                    unsafe_allow_html=True,
                )

            cv, cs, cd = st.columns(3)
            if cv.button("Login", use_container_width=True):
                st.session_state.nw_username = user_in
                st.session_state.nw_password = pass_in

                # ── Token-paste path ──────────────────────────────
                if login_method == "Paste Token":
                    raw_token = token_in.strip()
                    if not raw_token:
                        st.error("Paste a token first.")
                    elif not host_in.strip():
                        st.error("Enter the Host URL.")
                    else:
                        st.session_state.nw_token = raw_token
                        with st.spinner("Verifying token…"):
                            ok, msg = nw_verify_token()
                        st.session_state.nw_verified = ok
                        st.session_state.nw_msg      = msg
                        if ok:
                            st.success(f"{msg}")
                            ok2, items, _diag = nw_fetch_incidents()
                            if ok2:
                                st.session_state.incidents  = items
                                st.session_state.last_fetch = datetime.now()
                                db_upsert_incidents(items)
                                st.success(f"Fetched {len(items)} incidents")
                            else:
                                st.warning("Token accepted but no incidents fetched yet.")
                        else:
                            st.session_state.nw_token = ""
                            st.error(f"{msg}")
                    st.rerun()

                # ── Username / Password path ───────────────────────
                else:
                    with st.spinner("Logging in…"):
                        try:
                            ok, msg, token = nw_login(host_in, user_in, pass_in)
                        except Exception as _ve:
                            ok, msg, token = False, f"Exception: {_ve}", ""
                    st.session_state.nw_verified = ok
                    st.session_state.nw_msg      = msg
                    if ok:
                        st.session_state.nw_token = token
                        st.success(f"{msg}")
                        ok2, items, _diag = nw_fetch_incidents()
                        if ok2:
                            st.session_state.incidents  = items
                            st.session_state.last_fetch = datetime.now()
                            db_upsert_incidents(items)
                            st.success(f"Fetched {len(items)} incidents")
                        else:
                            st.warning("Connected but no incidents fetched yet.")
                    else:
                        st.error(f"{msg}")
                    st.rerun()

            if cs.button("Save", use_container_width=True,
                         help="Save to .env — auto-connects on next startup"):
                if host_in and user_in and pass_in:
                    env_save(host_in, user_in, pass_in)
                    st.success("Saved — will auto-connect on next startup")
                else:
                    st.warning("Enter host, username and password first.")

            with cd:
                if _confirm_action(
                        "sb_clear_main", "✕ Clear",
                        "This clears your saved NetWitness host, username, "
                        "and password and disconnects the session.",
                        confirm_label="Clear credentials"):
                    env_clear()
                    st.session_state.update(
                        nw_token="", nw_username="", nw_password="",
                        nw_host="", nw_verified=False, nw_msg="",
                        incidents=[], last_fetch=None, _startup_done=False,
                    )
                    st.rerun()

        # ── TLS Certificate (Option B — verified HTTPS) ─────────────
        st.markdown('<div class="sec-label">  Security Certificate</div>', unsafe_allow_html=True)

        # Auto-clear bad cert if last message was an SSL error
        _last_msg = st.session_state.get("nw_msg", "")
        if "SSL error" in _last_msg and st.session_state.get("nw_cert_path", ""):
            st.warning(
                f"The uploaded cert caused an SSL error — it has been removed. "
                f"Revert to browser export or try again.\n\n`{_last_msg}`"
            )
            st.session_state.nw_cert_path = ""

        _cert_active = st.session_state.get("nw_cert_path", "").strip()
        _cert_valid  = bool(_cert_active) and Path(_cert_active).is_file()

        if _cert_valid:
            st.markdown(
                f'<div style="font-family:var(--mono);font-size:0.62rem;color:var(--green);'
                f'margin-bottom:6px"><span class="dot dot-green"></span>'
                f'Verifying against: {Path(_cert_active).name}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--warn);'
                'margin-bottom:6px"><span class="dot dot-yellow"></span>'
                'No cert configured — TLS verification skipped (insecure)</div>',
                unsafe_allow_html=True,
            )

        cert_upload = st.file_uploader(
            "Upload server/CA certificate (.pem / .crt)",
            type=["pem", "crt", "cer"],
            key="cert_uploader",
            label_visibility="collapsed",
        )
        if cert_upload is not None:
            certs_dir = Path(__file__).parent / "certs"
            certs_dir.mkdir(exist_ok=True)
            cert_dest = certs_dir / cert_upload.name
            cert_dest.write_bytes(cert_upload.getvalue())
            st.session_state.nw_cert_path = str(cert_dest)
            nw_cert_env_save(str(cert_dest))
            st.success(f"Saved {cert_upload.name} — re-login to apply verified TLS")
            st.rerun()

        cert_path_in = st.text_input(
            "…or enter an existing cert path",
            value=st.session_state.get("nw_cert_path", ""),
            placeholder="/path/to/netwitness-ca.pem",
            key="cert_path_text",
        )
        ccert1, ccert2 = st.columns(2)
        if ccert1.button("Use this path", use_container_width=True):
            if cert_path_in.strip() and Path(cert_path_in.strip()).is_file():
                st.session_state.nw_cert_path = cert_path_in.strip()
                nw_cert_env_save(cert_path_in.strip())
                st.success("Cert path set — re-login to apply")
            else:
                st.error("File not found at that path.")
            st.rerun()
        if ccert2.button("✕ Remove cert", use_container_width=True):
            st.session_state.nw_cert_path = ""
            nw_cert_env_clear()
            st.info("Reverted to verify=False (insecure)")
            st.rerun()

        # ── Foundation LLM (OpenAI) ────────────────────────────────
        st.markdown('<div class="sec-label">  AI Settings</div>', unsafe_allow_html=True)

        # Connection status indicator
        if st.session_state.cisco_connected:
            st.markdown(
                '<div style="background:#041208;border:1px solid #0A3020;border-radius:5px;'
                'padding:8px 11px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
                '<span class="dot dot-green"></span>'
                '<strong style="color:var(--green)">LLM CONFIGURED</strong><br>'
                f'<span style="color:var(--muted);font-size:0.58rem">'
                f'Model: {st.session_state.cisco_model or "—"}</span>'
                '</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#0A0608;border:1px solid #2A1010;border-radius:5px;'
                'padding:8px 11px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
                '<span class="dot dot-red"></span>'
                '<span style="color:var(--danger)">NOT CONFIGURED</span><br>'
                '<span style="font-size:0.57rem;color:var(--muted)">Enter your OpenAI API key below, '
                'or set OPENAI_API_KEY in .env</span>'
                '</div>',
                unsafe_allow_html=True,
            )

        cisco_url_in = st.text_input(
            "Endpoint URL",
            value=st.session_state.cisco_url,
            placeholder="https://api.openai.com/v1",
            key="sb_cisco_url",
        )
        cisco_key_in = st.text_input(
            "OpenAI API Key",
            value="",
            type="password",
            placeholder="sk-xxxxxxxxxxxxxxxxxxxx",
            key="sb_cisco_key",
        )
        cisco_model_in = st.text_input(
            "Model name",
            value=st.session_state.cisco_model,
            placeholder="gpt-4o-mini",
            key="sb_cisco_model",
        )

        cl1, cl2, cl3 = st.columns(3)

        if cl1.button("Apply", use_container_width=True, key="cisco_apply"):
            if not cisco_url_in.strip():
                st.error("Enter the endpoint URL.")
            elif not cisco_key_in.strip():
                st.error("Enter your OpenAI API key.")
            else:
                st.session_state.cisco_url       = cisco_url_in.strip()
                st.session_state.cisco_key       = cisco_key_in.strip()
                st.session_state.cisco_model     = (
                    cisco_model_in.strip() or "gpt-4o-mini"
                )
                st.session_state.cisco_connected = True
                st.success("LLM configured!")
                st.rerun()

        if cl2.button("Save", use_container_width=True, key="cisco_save"):
            if cisco_url_in.strip() and cisco_key_in.strip():
                st.session_state.cisco_url       = cisco_url_in.strip()
                st.session_state.cisco_key       = cisco_key_in.strip()
                st.session_state.cisco_model     = (
                    cisco_model_in.strip() or "gpt-4o-mini"
                )
                st.session_state.cisco_connected = True
                cisco_env_save(
                    cisco_url_in.strip(),
                    cisco_key_in.strip(),
                    st.session_state.cisco_model,
                )
                st.success("Saved to .env")
                st.rerun()
            else:
                st.warning("Enter URL and OpenAI API key first.")

        with cl3:
            if _confirm_action(
                    "cisco_clear", "✕ Clear",
                    "This clears your saved LLM endpoint URL, key, and model name.",
                    confirm_label="Clear LLM settings"):
                st.session_state.cisco_url       = ""
                st.session_state.cisco_key       = ""
                st.session_state.cisco_model     = ""
                st.session_state.cisco_connected = False
                cisco_env_clear()
                st.rerun()



        st.markdown("---")
        st.markdown(
            f'<div style="font-family:var(--mono);font-size:0.54rem;'
            f'color:#1A3A52;text-align:center;line-height:1.9">'
            f'v4 · AUTO-REFRESH {REFRESH_INTERVAL}s<br>'
            f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>',
            unsafe_allow_html=True,
        )


# ══════════════════════════════════════════════════════════════════════════════
# BACKGROUND WORKFLOW POLL — last statement so every tab renders first.
# While the worker thread runs, rerun every ~1.5s so the agent board (and any
# open detail panel) refreshes with the worker's latest thinking/output.
# ══════════════════════════════════════════════════════════════════════════════
try:
    _wf_poll = _workflow_store().get("run")
    if _wf_poll is not None and not _wf_poll.get("done"):
        time.sleep(1.5)
        st.rerun()
except Exception:
    pass
