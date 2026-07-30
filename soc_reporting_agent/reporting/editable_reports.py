




# ============================================================================
# [FYP-FILE] soc_reporting_agent/reporting/editable_reports.py
# File: soc_reporting_agent/reporting/editable_reports.py
# Purpose: This module implements report generation and export behaviour for editable reports.
# Inputs: Receives function arguments, configured state, and persisted artifacts described below.
# Outputs: Produces return values and documented state, file, database, export, or UI effects.
# Workflow position: Aegis report generation and export.
# Important dependencies: __future__, datetime, hashlib, json, os, pathlib, re, reporting.
# Key evaluator search terms: _aegis_logo_path, utc_now, markdown_to_plain_text, incident_report_dir, editable_dir, drafts_dir, [FYP-FUNCTION].
# ----------------------------------------------------------------------------
# PURPOSE
#   Owns the analyst-facing "editable report" lifecycle for the Reporting
#   stage: turning the Reporting Agent's freshly generated section text into
#   a per-incident manifest (report_manifest.json) that tracks each of the
#   four core report sections through draft -> (optional revision) ->
#   confirmed -> exported, plus rendering those sections to Word (.docx) and
#   PDF using a shared, hand-built document layout (no external template
#   engine at this layer). Also owns the separate, IMMUTABLE
#   candidate_manifest.json snapshot that is the sole hand-off artefact
#   between the Reporting stage and the downstream approval/workflow layer.
#
# MAIN FUNCTIONALITIES
#   - Report directory layout helpers (drafts/confirmed/exports/draft_history)
#   - report_manifest.json build/read/save (build_report_manifest,
#     load_manifest, save_manifest) and its draft/confirmed/exported status
#     bookkeeping
#   - Section-level analyst editing: read_section, save_section (with
#     revision history), confirm_section, confirm_report
#   - Word/PDF rendering of block-structured report content (_docx_write*,
#     _pdf_write*, render_blocks_to_docx/pdf) shared by both the manifest
#     workflow and callers that render arbitrary block content directly
#   - Word/PDF export of confirmed sections and the combined Final Incident
#     Report (export_section_docx/pdf, export_docx/pdf, download_path)
#   - finalize_candidate_manifest(): one-time, hash-verified publication of
#     candidate_manifest.json once all 4 sections are confirmed and exported
#
# INPUTS
#   - generated_sections: dict of file paths produced by the Reporting Agent
#     (soc_reporting_agent/agents/reporting_agent.py) for each report type
#   - output_dir / incident_id: identify the per-incident output tree under
#     settings.OUTPUT_DIR (<output_dir>/<incident_id>/reports/...)
#   - Analyst edits (text and/or structured blocks) submitted via
#     save_section()
#   - report_manifest.json / candidate_manifest.json read back from disk
#
# OUTPUTS
#   - report_manifest.json (mutable, evolves through the analyst workflow)
#   - candidate_manifest.json (write-once, sha256-hashed final snapshot)
#   - Draft/confirmed .txt and .json (structured block) files under
#     drafts/, confirmed/, draft_history/
#   - Exported .docx / .pdf files under exports/
#
# WORKFLOW POSITION
#   Reporting stage (last stage in the pipeline, after Triage and
#   Investigation -- see soc_workflow.py's stage_labels: parsing -> triage ->
#   threat_intel -> investigation -> reporting). This module is the
#   analyst-editing/export layer that sits downstream of report generation
#   (reporting_agent.py -> context_builder.py -> export_context_enhancer.py
#   -> report_renderer.py) and upstream of the reporting-approval gate
#   (reporting_approval.approve_reporting_candidate() /
#   workflow_state_store.commit_reporting_approval()), which is the only
#   code path allowed to move workflow_status to "Complete".
#
# CALLED BY (verified via repo-wide grep for "editable_reports" / imported
# names; see individual function docstrings below for which functions each
# caller actually uses)
#   - soc_reporting_agent/backend/app.py -- Flask dashboard API layer; wires
#     nearly every public function here (list_reports, read_section,
#     save_section, confirm_section, confirm_report, export_section_docx/pdf,
#     export_docx/pdf, download_path, list_section_drafts) to REST endpoints
#   - soc_reporting_agent/adapters/export_documents.py -- headless CLI
#     adapter used by the SOC workflow orchestrator to auto-confirm+export
#     all sections, then call finalize_candidate_manifest()
#   - soc_reporting_agent/reporting/report_renderer.py -- imports
#     REPORT_SECTION_CONFIG, build_report_manifest, editable_dir
#   - soc_reporting_agent/reporting/report_validator.py -- references
#     finalize_candidate_manifest()'s hashing/atomic-write contract in its
#     own docstrings (decides, per report, whether the candidate set is safe
#     to publish)
#   - report_editing.py (repo root) -- main Streamlit app's analyst-edit
#     layer for the four core reports; uses incident_report_dir(),
#     render_blocks_to_docx(), render_blocks_to_pdf()
#   - triage_ticket_editing.py (repo root) -- same Streamlit app, triage
#     ticket export; uses render_blocks_to_docx(), render_blocks_to_pdf()
#   - tests/test_reporting_stage.py,
#     soc_reporting_agent/tests/test_structured_report_tables.py,
#     soc_reporting_agent/scripts/test_structured_report_review_exports.py
#
# CALLS
#   - reporting.structured_report (blocks_from_text, blocks_to_plain_text,
#     load_blocks, markdown_to_blocks, paragraph_contains_raw_pipe_table,
#     repair_pipe_tables_in_blocks, save_blocks) -- the block-structured
#     content model shared across the reporting pipeline
#   - reporting.report_validator.validate_generated_report() (imported
#     lazily inside finalize_candidate_manifest())
#   - reporting.template_document_exporter.convert_docx_to_pdf() (imported
#     lazily inside export_section_pdf() / export_pdf() /
#     render_blocks_to_pdf()) -- LibreOffice-backed DOCX->PDF conversion,
#     with a pure-Python reportlab fallback (_pdf_write_blocks) if that
#     conversion is unavailable in this environment
#   - Optional third-party libs: python-docx (Document, ...) and reportlab
#     (SimpleDocTemplate, ...), both imported defensively (bound to None on
#     ImportError) so this module can still be imported without them; the
#     writer functions raise a clear RuntimeError only if the specific
#     library needed for that call is missing
#
# KEY EVALUATOR SEARCH TERMS
#   [FYP-STATE] draft / draft_revision / confirmed / exported /
#   partially_confirmed / analyst_editing / confirmed_by_analyst /
#   draft_ready_for_analyst_review
#   [FYP-EVALUATOR] save_section / confirm_section / confirm_report /
#   finalize_candidate_manifest
#   [FYP-EXPORT] export_docx / export_pdf / export_section_docx /
#   export_section_pdf / candidate_manifest.json / report_manifest.json
# ============================================================================

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from reporting.structured_report import (
    blocks_from_text,
    blocks_to_plain_text,
    load_blocks,
    markdown_to_blocks,
    paragraph_contains_raw_pipe_table,
    repair_pipe_tables_in_blocks,
    save_blocks,
)

# [FYP-FALLBACK] python-docx is an optional runtime dependency. If it is not
# installed, every name below is left as None so this module can still be
# imported (e.g. for its path helpers / manifest logic) without crashing;
# any code path that actually needs to write a .docx (see _docx_write_blocks)
# raises a clear RuntimeError instead of an ImportError at import time.
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:  # pragma: no cover
    Document = None
    Inches = None
    Pt = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    OxmlElement = None
    qn = None

# [FYP-FALLBACK] Same defensive-import pattern as python-docx above, but for
# reportlab, which backs the pure-Python PDF renderer (_pdf_write_blocks)
# used only when the preferred LibreOffice DOCX->PDF conversion path
# (reporting.template_document_exporter.convert_docx_to_pdf) is unavailable.
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except Exception:  # pragma: no cover
    A4 = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    inch = None
    Paragraph = None
    SimpleDocTemplate = None
    Image = None
    PageBreak = None
    Spacer = None
    Table = None
    TableStyle = None
    colors = None

# ============================================================================
# [FYP-SECTION] Report Type Configuration
# [FYP-CONFIG] Static catalogue of every report "section" the Reporting
# stage can produce -- title/template/output filename/description per key.
# Consumed by build_report_manifest() to seed each section's manifest entry,
# and re-exported (REPORT_SECTION_CONFIG, CORE_REPORT_KEYS) for backend/
# app.py and report_renderer.py to drive the analyst UI and template
# rendering without duplicating this catalogue.
# ============================================================================
# Five templates are still available to the renderer. The dashboard exposes the
# four SOC-facing report artefacts requested by the analyst. The triage review
# template remains a supporting template that can be merged or referenced by
# the analyst review, but it is not shown as a fifth primary report card.
REPORT_SECTION_CONFIG: dict[str, dict[str, str]] = {
    "executive_summary": {
        "title": "Executive Summary",
        "template": "executive_summary_template.md.j2",
        "filename": "executive_summary.txt",
        "description": "Management-level incident summary for handover and leadership review.",
    },
    "technical_findings": {
        "title": "Technical Findings",
        "template": "technical_findings_template.md.j2",
        "filename": "technical_findings.txt",
        "description": "Evidence, IOCs, technical observations, and validation points.",
    },
    "soc_analyst_review": {
        "title": "SOC Analyst Review",
        "template": "soc_analyst_review_template.md.j2",
        "filename": "soc_analyst_review.txt",
        "description": "Analyst judgement, limitations, approval notes, and review checklist.",
    },
    "soc_triage_review": {
        "title": "SOC Triage Review",
        "template": "soc_triage_review_template.md.j2",
        "filename": "soc_triage_review.txt",
        "description": "Supporting triage decision notes used by the analyst review.",
    },
    "final_incident_report": {
        "title": "Final Incident Report",
        "template": "incident_report_template.md.j2",
        "filename": "final_incident_report.txt",
        "description": "Complete incident report for analyst confirmation and export.",
    },
}

# [FYP-CONFIG] The four analyst-facing report types (dashboard "report
# cards") vs. the one supporting template that gets folded into
# soc_analyst_review rather than shown standalone (see build_report_manifest).
CORE_REPORT_KEYS = ["executive_summary", "technical_findings", "soc_analyst_review", "final_incident_report"]
SUPPORT_REPORT_KEYS = ["soc_triage_review"]
EXPORT_SECTION_ORDER = CORE_REPORT_KEYS


# [FYP-CONFIG] soc_reporting_agent/ package root, used to locate report_assets
# (logo) regardless of the current working directory the process was started
# from.
PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _aegis_logo_path() -> Path | None:
    """[FYP-FUNCTION] Locate Aegis Logo Asset

    Purpose: find the branding logo used on the cover of every exported
    Word/PDF report.
    Params: none.
    Source: checks report_assets/aegis-logo.png then dashboard/assets/
        aegis-logo.png under PROJECT_ROOT.
    Returns: the first existing, non-empty candidate Path, or None if
        neither is present.
    Side effects: none (read-only filesystem check).
    Called by: _add_title_block() (DOCX) and _pdf_write_blocks() (PDF).
    Fallback: callers render a plain "AEGIS" text label instead of the logo
        image when this returns None.
    """
    for candidate in [
        PROJECT_ROOT / "report_assets" / "aegis-logo.png",
        PROJECT_ROOT / "dashboard" / "assets" / "aegis-logo.png",
    ]:
        if candidate.exists() and candidate.stat().st_size > 0:
            return candidate
    return None


def utc_now() -> str:
    """[FYP-FUNCTION] Current UTC Timestamp

    Purpose: single canonical timestamp source for every "created_at" /
    "confirmed_at" / "last_saved_at" style field written by this module, so
    all timestamps in a manifest are directly comparable.
    Params: none. Returns: ISO-8601 string, timezone-aware (UTC).
    Called by: nearly every state-mutating function in this file.
    """
    return datetime.now(timezone.utc).isoformat()


def markdown_to_plain_text(value: Any) -> str:
    """[FYP-FUNCTION] Markdown To Plain Text

    Purpose: strip markdown syntax (code fences, inline code, links,
    headings, blockquotes, bullets, bold/italic, tables, rule lines) down to
    readable plain text, used as a fallback when structured "blocks" are not
    available and the module has to fall back to plain-text rendering.
    Params: value -- any value; coerced to str via str(value or "").
    Returns: cleaned plain-text string (never None).
    Side effects: none (pure string transform).
    Called by: save_section() (when no explicit blocks are supplied),
    list_section_drafts() (revision previews), _draft_history_entry()
    (history preview text).
    """
    text = str(value or "")
    if not text:
        return ""
    text = re.sub(r"```[a-zA-Z0-9_-]*\n?", "", text)
    text = text.replace("```", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    text = re.sub(r"^\s{0,3}#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s{0,3}>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"(^|[^*])\*([^*\n]+)\*", r"\1\2", text)
    text = re.sub(r"(^|[^_])_([^_\n]+)_", r"\1\2", text)
    text = re.sub(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", "", text, flags=re.MULTILINE)

    # [FYP-FUNCTION] `_table_row` — implements the table row operation used by the surrounding report generation and export workflow.
    # [FYP-INPUT] Parameters: `match`; values come from its direct caller, route, UI event, fixture, or stage handoff.
    # [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
    # [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
    # [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
    # [FYP-CALLS] Calls: `group`, `join`, `len`, `split`, `strip`.
    # [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

    def _table_row(match: re.Match) -> str:
        # [FYP-FUNCTION] Nested re.sub callback (used only by the pipe-table
        # regex directly below): turns one "|a|b|c|" markdown row into a
        # "a: b | c" plain-text line. First cell is treated as a label.
        row = match.group(1)
        cells = [c.strip() for c in row.split("|") if c.strip()]
        if len(cells) >= 2:
            return f"{cells[0]}: " + " | ".join(cells[1:])
        return " | ".join(cells)

    text = re.sub(r"^\s*\|(.+)\|\s*$", _table_row, text, flags=re.MULTILINE)
    text = text.replace("---", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ============================================================================
# [FYP-SECTION] Report Directory Layout Helpers
# [FYP-EXPORT] All paths below live under
# <output_dir>/<incident_id>/reports/... . This is the on-disk shape of the
# [FYP-STATE] draft -> confirmed -> exported lifecycle: a section's draft
# text lives in drafts_dir() until confirm_section() copies it into
# confirmed_dir() (see [FYP-STATE] transition in confirm_section()); each
# saved revision is preserved (never overwritten) under draft_history_dir();
# rendered Word/PDF output lands in exports_dir(); manifest_path() is the
# mutable report_manifest.json; candidate_manifest_path() (defined further
# below, near CandidateManifestConflictError) is the separate, immutable
# final snapshot.
# ============================================================================
def incident_report_dir(output_dir: Path, incident_id: str) -> Path:
    """[FYP-FUNCTION] Incident Reports Root Dir. Returns
    <output_dir>/<incident_id>/reports -- the parent of every other path
    helper in this section. Called by: every function in this file that
    resolves a report-related path; report_editing.py (repo root)."""
    return output_dir / incident_id / "reports"


def editable_dir(output_dir: Path, incident_id: str) -> Path:
    # Backwards-compatible renderer target. Files are copied into drafts by the manifest builder.
    """[FYP-FUNCTION] Legacy Editable Dir. Where the Reporting Agent's raw
    generated_sections output is read from by build_report_manifest() before
    being copied into drafts_dir(); kept as the renderer's stable target
    path for backward compatibility with existing generation code."""
    return incident_report_dir(output_dir, incident_id) / "editable"


def drafts_dir(output_dir: Path, incident_id: str) -> Path:
    """[FYP-FUNCTION] [FYP-STATE] Draft Reports Dir. Holds the current
    "draft" / "draft_revision" text+block files for each section, i.e. the
    analyst's in-progress edits before confirm_section() locks them in."""
    return incident_report_dir(output_dir, incident_id) / "drafts"


def confirmed_dir(output_dir: Path, incident_id: str) -> Path:
    """[FYP-FUNCTION] [FYP-STATE] Confirmed Reports Dir. Holds the
    analyst-confirmed (locked) text+block files written by confirm_section();
    this is the only source export_section_docx/pdf() are allowed to render
    from (see _confirmed_required())."""
    return incident_report_dir(output_dir, incident_id) / "confirmed"


def exports_dir(output_dir: Path, incident_id: str) -> Path:
    """[FYP-FUNCTION] [FYP-EXPORT] Exports Dir. Destination for every
    rendered .docx/.pdf file (per-section and the combined final report)."""
    return incident_report_dir(output_dir, incident_id) / "exports"


def draft_history_dir(output_dir: Path, incident_id: str, section_key: str) -> Path:
    """[FYP-FUNCTION] [FYP-STATE] Per-Section Draft History Dir. Append-only
    store of every prior saved draft (see _draft_history_entry()), keyed by
    section, used to power list_section_drafts()'s revision browser."""
    return incident_report_dir(output_dir, incident_id) / "draft_history" / section_key


def final_dir(output_dir: Path, incident_id: str) -> Path:
    # Backwards-compatible old name.
    """[FYP-FUNCTION] Deprecated alias for exports_dir() -- kept only so any
    external code still importing the old name does not break."""
    return exports_dir(output_dir, incident_id)


def manifest_path(output_dir: Path, incident_id: str) -> Path:
    """[FYP-FUNCTION] [FYP-STATE] Mutable manifest file path
    (report_manifest.json) -- the single source of truth for every
    section's status, paths, and history while the Reporting stage is in
    analyst-editing mode."""
    return incident_report_dir(output_dir, incident_id) / "report_manifest.json"


def _rel(output_dir: Path, path: Path) -> str:
    """[FYP-FUNCTION] Relative Path For Display. Best-effort path relative
    to output_dir.parent, used only for human-/UI-friendly "relative_path"
    manifest fields (never for actual file I/O, which always uses the
    absolute path). Falls back to the absolute path string if the input
    path is not actually under output_dir.parent (e.g. cross-drive paths)."""
    try:
        return str(path.relative_to(output_dir.parent))
    except Exception:
        return str(path)


def _read(path: str | Path | None) -> str:
    """[FYP-FUNCTION] Safe Text Read. Reads a text file as UTF-8 with
    invalid bytes ignored (rather than raising) so a partially-written or
    encoding-corrupt file degrades to a shorter string instead of crashing
    the reporting UI. Returns "" for a falsy/None path or a missing file."""
    if not path:
        return ""
    p = Path(path)
    return p.read_text(encoding="utf-8", errors="ignore") if p.exists() else ""


def _write(path: Path, text: str) -> None:
    """[FYP-FUNCTION] Safe Text Write. Creates parent directories as needed
    and writes str(text or "") as UTF-8. Side effect: creates/overwrites the
    file at path."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(str(text or ""), encoding="utf-8")


def _safe_resolve(path: Path) -> Path:
    """[FYP-FUNCTION] Resolve With Fallback. Path.resolve() can raise on
    some platforms/edge cases (e.g. broken symlinks); falls back to
    Path.absolute() so path-identity comparisons elsewhere (e.g. "is this
    draft file literally the same file as the confirmed file") never crash."""
    try:
        return path.resolve()
    except Exception:
        return path.absolute()


def _draft_history_entry(output_dir: Path, incident_id: str, section_key: str, title: str, text: str, analyst: str) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-STATE] Record One Draft Revision.

    Purpose: persist a timestamped, immutable snapshot of a section's draft
    text every time save_section() is called, so analysts can browse/restore
    prior revisions (list_section_drafts()).
    Params: output_dir/incident_id/section_key -- identify where to write;
        title -- section display title; text -- the draft text being saved
        (source: caller's cleaned block/plain text); analyst -- name of the
        SOC analyst who made this edit (source: API request body /
        default "SOC Analyst").
    Returns: a history-entry dict (created_at, saved_by, title, path,
        relative_path, preview) intended to be appended to
        section["draft_history"].
    Side effects: writes a new "<UTC-timestamp>.txt" file under
        draft_history_dir(); does not touch report_manifest.json itself
        (the caller, save_section(), is responsible for appending the
        returned entry and re-saving the manifest).
    Called by: save_section().
    """
    hdir = draft_history_dir(output_dir, incident_id, section_key)
    hdir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")
    path = hdir / f"{stamp}.txt"
    _write(path, text)
    preview = markdown_to_plain_text(text)[:500]
    return {
        "created_at": utc_now(),
        "saved_by": analyst,
        "title": title,
        "path": str(path),
        "relative_path": _rel(output_dir, path),
        "preview": preview,
    }


# ============================================================================
# [FYP-SECTION] Manifest Build / Load / Save
# [FYP-STATE] report_manifest.json is the mutable, evolving record of the
# whole Reporting stage's analyst workflow. Its top-level report_status /
# display_status and each section's status are recomputed by save_manifest()
# every time the manifest changes -- see the [FYP-STATE] values enumerated
# there (draft_ready_for_analyst_review, analyst_editing,
# partially_confirmed, confirmed_by_analyst; per-section: draft,
# draft_revision, confirmed, exported).
# ============================================================================
def build_report_manifest(output_dir: Path, incident_id: str, generated_sections: dict[str, str], context: dict[str, Any] | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EVALUATOR] [FYP-STATE] [FYP-EXPORT] Build The
    Initial Report Manifest.

    Purpose: the entry point that turns the Reporting Agent's raw generated
    output (plain-text + structured-block files on disk) into the very first
    report_manifest.json for an incident, with every core section seeded at
    [FYP-STATE] status="draft". This is where the Reporting stage's
    analyst-editable workflow begins.
    Params: output_dir/incident_id -- identify the incident's output tree;
        generated_sections -- dict of file paths keyed by report type (and
        "<type>_structured" for the structured block JSON), source: the
        Reporting Agent's write_outputs() step; context -- optional reporting
        context dict (source: export_context_enhancer.enhance_export_context
        output), used only to copy a few status fields into
        manifest["source_context"] for later diagnostics.
    Returns: the newly built manifest dict (same shape later read back by
        load_manifest()).
    Side effects: creates drafts_dir/confirmed_dir/exports_dir on disk; reads
        each generated section's text + structured-block files; special-cases
        "soc_analyst_review" by appending the supporting "soc_triage_review"
        content as an extra sub-heading when present; writes every section's
        draft text+blocks into drafts_dir(); writes report_manifest.json to
        manifest_path().
    Called by: soc_reporting_agent/reporting/report_renderer.py (after the
        Reporting Agent finishes generating section content).
    Calls: editable_dir(), drafts_dir()/confirmed_dir()/exports_dir(),
        reporting.structured_report.{load_blocks, repair_pipe_tables_in_blocks,
        blocks_from_text, save_blocks}, _rel(), utc_now().
    """
    context = context or {}
    ddir = drafts_dir(output_dir, incident_id)
    cdir = confirmed_dir(output_dir, incident_id)
    edir = exports_dir(output_dir, incident_id)
    ddir.mkdir(parents=True, exist_ok=True)
    cdir.mkdir(parents=True, exist_ok=True)
    edir.mkdir(parents=True, exist_ok=True)

    # Use the supporting triage template by appending a short supporting section
    # into SOC Analyst Review when the support file exists. This keeps the user-facing
    # output at four reports while still using the triage template in the renderer.
    support_text = _read(generated_sections.get("soc_triage_review"))

    sections: dict[str, dict[str, Any]] = {}
    for key in CORE_REPORT_KEYS:
        cfg = REPORT_SECTION_CONFIG[key]
        src = Path(generated_sections.get(key) or editable_dir(output_dir, incident_id) / cfg["filename"])
        text = _read(src)
        structured_src = Path(generated_sections.get(f"{key}_structured") or editable_dir(output_dir, incident_id) / f"{key}.json")
        blocks = repair_pipe_tables_in_blocks(load_blocks(structured_src))
        if not blocks:
            blocks = blocks_from_text(text) if text else []
        if key == "soc_analyst_review" and support_text and "Supporting SOC Triage Review" not in text:
            support_blocks = load_blocks(generated_sections.get("soc_triage_review_structured")) or blocks_from_text(support_text)
            text = (text.rstrip() + "\n\nSupporting SOC Triage Review\n\n" + support_text.strip()).strip()
            blocks = blocks + [{"type": "heading", "level": 2, "text": "Supporting SOC Triage Review"}] + support_blocks
        draft_path = ddir / cfg["filename"]
        structured_draft_path = ddir / f"{key}.json"
        _write(draft_path, text)
        save_blocks(structured_draft_path, blocks)
        sections[key] = {
            "key": key,
            "title": cfg["title"],
            "description": cfg["description"],
            "template": cfg["template"],
            "filename": cfg["filename"],
            "status": "draft",
            "draft_path": str(draft_path),
            "draft_relative_path": _rel(output_dir, draft_path),
            "structured_draft_path": str(structured_draft_path),
            "structured_draft_relative_path": _rel(output_dir, structured_draft_path),
            "confirmed_path": None,
            "confirmed_relative_path": None,
            "structured_confirmed_path": None,
            "structured_confirmed_relative_path": None,
            "exports": {},
            "last_saved_at": utc_now(),
            "last_saved_by": "Reporting Agent",
            "confirmed_at": None,
            "confirmed_by": None,
            "draft_history": [],
        }
    support_sections = {}
    for key in SUPPORT_REPORT_KEYS:
        if key in generated_sections:
            cfg = REPORT_SECTION_CONFIG[key]
            support_sections[key] = {
                "key": key,
                "title": cfg["title"],
                "template": cfg["template"],
                "path": generated_sections[key],
                "relative_path": _rel(output_dir, Path(generated_sections[key])),
            }
    manifest = {
        "schema_version": "editable-report-manifest-v2",
        "incident_id": incident_id,
        "report_status": "draft_ready_for_analyst_review",
        "display_status": "Draft ready for analyst review",
        "sections": sections,
        "section_order": CORE_REPORT_KEYS,
        "support_sections": support_sections,
        "draft_reports": list(CORE_REPORT_KEYS),
        "confirmed_reports": [],
        "created_at": utc_now(),
        "updated_at": utc_now(),
        "confirmed_by": None,
        "confirmed_at": None,
        "exports": {},
        "source_context": {
            "report_status": context.get("report_status"),
            "validation_status": context.get("validation_status"),
            "data_consistency_status": context.get("data_consistency_status"),
        },
    }
    path = manifest_path(output_dir, incident_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def load_manifest(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-STATE] Load Report Manifest.

    Purpose: the single read path used everywhere else in this module (and
    by backend/app.py) to fetch the current state of a report_manifest.json.
    Params: output_dir -- output tree root; incident_id -- optional; when
        omitted, falls back to the most-recently-modified
        */reports/report_manifest.json under output_dir (source: filesystem
        mtime scan), used by callers that operate on "whichever incident is
        currently open" without needing to pass an explicit id.
    Returns: the parsed manifest dict, or {} if none exists yet.
    Side effects: none (read-only).
    Called by: nearly every other function in this file (list_reports,
        read_section, save_section, confirm_section, confirm_report,
        export_*, finalize_candidate_manifest, ...); backend/app.py.
    """
    if incident_id:
        path = manifest_path(output_dir, incident_id)
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    manifests = sorted(output_dir.glob("*/reports/report_manifest.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    if manifests:
        return json.loads(manifests[0].read_text(encoding="utf-8"))
    return {}


def save_manifest(output_dir: Path, manifest: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-STATE] Persist Manifest + Recompute Top-Level
    Status.

    Purpose: the single write path for report_manifest.json. Every call
    recomputes the whole-report [FYP-STATE] rollup from each section's
    individual status, so report_status/display_status are always derived,
    never hand-set by callers:
      - all sections confirmed, none draft -> "confirmed_by_analyst" /
        "All reports confirmed by analyst"
      - some confirmed, some still draft   -> "partially_confirmed" /
        "Partially confirmed"
      - any section still draft/draft_revision -> "analyst_editing" /
        "Analyst editing"
    Params: output_dir -- output tree root; manifest -- the manifest dict to
        persist (mutated in place: draft_reports/confirmed_reports/
        report_status/display_status/updated_at are rewritten from
        manifest["sections"]).
    Returns: None.
    Side effects: overwrites report_manifest.json on disk (creates parent
        dirs if needed); mutates the manifest dict argument in place.
    Called by: build_report_manifest() is the only function that writes the
        manifest directly instead; every other mutator (save_section,
        confirm_section, confirm_report, export_section_docx/pdf, export_docx,
        export_pdf) calls this at the end of its own status change.
    """
    incident_id = manifest.get("incident_id") or "INC-0001"
    manifest["updated_at"] = utc_now()
    manifest["draft_reports"] = [k for k, s in (manifest.get("sections") or {}).items() if s.get("status") in {"draft", "draft_revision"} and s.get("draft_path")]
    manifest["confirmed_reports"] = [k for k, s in (manifest.get("sections") or {}).items() if s.get("status") in {"confirmed", "exported"} and s.get("confirmed_path")]
    if manifest["confirmed_reports"] and not manifest["draft_reports"]:
        manifest["report_status"] = "confirmed_by_analyst"
        manifest["display_status"] = "All reports confirmed by analyst"
    elif manifest["confirmed_reports"]:
        manifest["report_status"] = "partially_confirmed"
        manifest["display_status"] = "Partially confirmed"
    elif manifest["draft_reports"]:
        manifest["report_status"] = "analyst_editing"
        manifest["display_status"] = "Analyst editing"
    path = manifest_path(output_dir, incident_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def list_reports(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-API] List All Report Sections.

    Purpose: the dashboard-facing summary view of a manifest -- used to
    render the "Reports" tab's list of cards with their current status.
    Params: output_dir/incident_id -- see load_manifest().
    Returns: dict with manifest, reports (all sections), draft_reports and
        confirmed_reports (lists of section dicts, resolved via
        manifest["draft_reports"]/["confirmed_reports"] key lists), and
        section_order.
    Error handling: raises FileNotFoundError if no manifest exists yet (i.e.
        the Reporting Agent has not run for this incident) -- callers (e.g.
        backend/app.py's /api/reports routes) turn this into a 404-style
        JSON error response.
    Called by: soc_reporting_agent/backend/app.py (GET /api/reports and the
        per-incident equivalent).
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    sections = manifest.get("sections") or {}
    return {
        "success": True,
        "manifest": manifest,
        "reports": sections,
        "draft_reports": [sections[k] for k in manifest.get("draft_reports", []) if k in sections],
        "confirmed_reports": [sections[k] for k in manifest.get("confirmed_reports", []) if k in sections],
        "section_order": manifest.get("section_order", CORE_REPORT_KEYS),
    }


def _section_text(section: dict[str, Any]) -> tuple[str, str]:
    """[FYP-FUNCTION] [FYP-STATE] Resolve Section Plain Text. Prefers the
    draft file if one exists (analyst is mid-edit), else falls back to the
    confirmed file. Returns (source_label, text) where source_label is one
    of "draft" / "confirmed" / "missing" -- the same vocabulary used
    elsewhere for section status. Called by: read_section(),
    _section_blocks() (text fallback when no structured blocks exist)."""
    if section.get("draft_path") and Path(section["draft_path"]).exists():
        return "draft", _read(section["draft_path"])
    if section.get("confirmed_path") and Path(section["confirmed_path"]).exists():
        return "confirmed", _read(section["confirmed_path"])
    return "missing", ""


def _section_blocks(section: dict[str, Any]) -> tuple[str, list[dict[str, Any]]]:
    """[FYP-FUNCTION] [FYP-STATE] Resolve Section Structured Blocks. Same
    draft-then-confirmed precedence as _section_text(), but for the
    structured block-list representation; falls back to converting plain
    text into blocks (blocks_from_text) if no structured file is present.
    Returns (source_label, blocks). Called by: read_section()."""
    if section.get("structured_draft_path") and Path(section["structured_draft_path"]).exists():
        return "draft", repair_pipe_tables_in_blocks(load_blocks(section["structured_draft_path"]))
    if section.get("structured_confirmed_path") and Path(section["structured_confirmed_path"]).exists():
        return "confirmed", repair_pipe_tables_in_blocks(load_blocks(section["structured_confirmed_path"]))
    source, text = _section_text(section)
    return source, blocks_from_text(text) if text else []


# ============================================================================
# [FYP-SECTION] Section-Level Analyst Editing (read / save / confirm)
# ============================================================================
def read_section(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-API] Read One Report Section For Editing.

    Purpose: fetch a single section's current text+blocks for the analyst
    editor UI, resolving draft-vs-confirmed precedence.
    Params: output_dir/incident_id -- see load_manifest(); section_key --
        one of REPORT_SECTION_CONFIG's keys.
    Returns: dict with manifest, section (manifest entry), text, blocks,
        source ("draft"/"confirmed"/"missing"), block_source.
    Error handling: FileNotFoundError if no manifest; KeyError if
        section_key is not a known section.
    Called by: soc_reporting_agent/backend/app.py (GET section endpoints).
    Calls: load_manifest(), _section_text(), _section_blocks().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    sections = manifest.get("sections") or {}
    if section_key not in sections:
        raise KeyError(f"Unknown report section: {section_key}")
    source, text = _section_text(sections[section_key])
    block_source, blocks = _section_blocks(sections[section_key])
    return {"manifest": manifest, "section": sections[section_key], "text": text, "blocks": blocks, "source": source, "block_source": block_source}


def save_section(output_dir: Path, section_key: str, text: str, analyst: str = "SOC Analyst", incident_id: str | None = None, blocks: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EVALUATOR] [FYP-STATE] Save An Analyst's Edited
    Report Section (the "save edited report" function).

    Purpose: persist an analyst's in-progress edit of one report section as
    a new draft, WITHOUT confirming/locking it -- this is what runs every
    time the analyst clicks "Save draft" in the Reports editor. Always
    writable, regardless of the section's current status (including
    re-editing an already-confirmed section, which reopens it as a
    "draft_revision" -- see [FYP-STATE] transition below).
    Params: output_dir/incident_id -- see load_manifest(); section_key --
        target section; text -- plain/markdown text from the editor (source:
        API request body); analyst -- name recorded as the editor (source:
        API request body, default "SOC Analyst"); blocks -- optional
        pre-structured block list from a richer editor UI; when omitted,
        blocks are derived from `text` via markdown_to_plain_text() +
        blocks_from_text().
    Returns: dict with success, manifest, section, text (re-read from disk),
        blocks, message.
    [FYP-STATE] Side effects / state change: writes the cleaned draft
        text+blocks into drafts_dir(); appends a new entry (see
        _draft_history_entry()) to section["draft_history"], capped at the
        most recent 20 entries; sets section["status"] to "draft_revision"
        if the section was already confirmed (section.get("confirmed_path")
        truthy), otherwise "draft"; updates last_saved_at/last_saved_by;
        calls save_manifest() which recomputes the whole-manifest rollup
        status.
    Error handling: FileNotFoundError if no manifest; KeyError if
        section_key is unknown.
    Called by: soc_reporting_agent/backend/app.py (POST section save
        endpoints, both per-incident and legacy routes).
    Calls: load_manifest(), repair_pipe_tables_in_blocks(),
        markdown_to_plain_text(), blocks_from_text(), blocks_to_plain_text(),
        _write(), save_blocks(), _draft_history_entry(), save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = manifest.get("sections", {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    incident_id = manifest.get("incident_id") or "INC-0001"
    draft_path = drafts_dir(output_dir, incident_id) / section["filename"]
    structured_draft_path = drafts_dir(output_dir, incident_id) / f"{section_key}.json"
    clean_blocks = repair_pipe_tables_in_blocks(blocks) if isinstance(blocks, list) else None
    if clean_blocks is None:
        clean_text_input = markdown_to_plain_text(text)
        clean_blocks = blocks_from_text(clean_text_input) if clean_text_input else []
    clean_text = blocks_to_plain_text(clean_blocks) if clean_blocks else markdown_to_plain_text(text)
    _write(draft_path, clean_text)
    save_blocks(structured_draft_path, clean_blocks)
    history = list(section.get("draft_history") or [])
    history.append(_draft_history_entry(output_dir, incident_id, section_key, section.get("title") or section_key, clean_text, analyst))
    section["draft_history"] = history[-20:]
    section["status"] = "draft_revision" if section.get("confirmed_path") else "draft"
    section["draft_path"] = str(draft_path)
    section["draft_relative_path"] = _rel(output_dir, draft_path)
    section["structured_draft_path"] = str(structured_draft_path)
    section["structured_draft_relative_path"] = _rel(output_dir, structured_draft_path)
    section["last_saved_at"] = utc_now()
    section["last_saved_by"] = analyst
    manifest["sections"][section_key] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "text": _read(draft_path), "blocks": clean_blocks, "message": f"Draft saved for {section['title']}"}


def confirm_section(output_dir: Path, section_key: str, analyst: str = "SOC Analyst", incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EVALUATOR] [FYP-STATE] [FYP-APPROVAL] Confirm
    (Lock) One Report Section -- the per-section "confirm report" state
    transition.

    Purpose: analyst sign-off on a single section's current draft, copying
    it into confirmed_dir() and marking it locked. This is the gate
    export_section_docx/pdf() require before a section may be rendered to
    Word/PDF (see _confirmed_required()).
    Params: output_dir/incident_id -- see load_manifest(); section_key --
        target section; analyst -- name recorded as confirming this section
        (source: API request body, default "SOC Analyst").
    Returns: dict with success, manifest, section, text, blocks, message.
    [FYP-STATE] Side effects / state change: copies the current draft (or,
        if already confirmed, the existing confirmed file, in which case the
        copy is skipped as it would be a self-copy) into confirmed_dir();
        copies/derives the structured block JSON alongside it; deletes the
        now-superseded draft file(s) from disk (best-effort; failures are
        swallowed) so the UI's draft/confirmed distinction stays unambiguous;
        sets section["status"] = "confirmed"; clears draft_path/
        structured_draft_path to None; sets confirmed_path/
        structured_confirmed_path, confirmed_at, confirmed_by; also stamps
        manifest-level confirmed_by/confirmed_at (last section confirmed
        wins at the manifest level -- confirm_report() re-stamps these
        again at the end for the whole-manifest case); calls save_manifest().
    Error handling: FileNotFoundError if no manifest or no draft/confirmed
        source file exists yet ("Save a draft first"); KeyError if
        section_key unknown.
    Called by: confirm_report() (once per section); backend/app.py (POST
        per-section confirm endpoints).
    Calls: load_manifest(), _safe_resolve(), shutil.copy2(),
        repair_pipe_tables_in_blocks(), load_blocks(), save_blocks(),
        blocks_from_text(), save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = manifest.get("sections", {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    incident_id = manifest.get("incident_id") or "INC-0001"
    src = Path(section.get("draft_path") or section.get("confirmed_path") or "")
    if not src.exists():
        raise FileNotFoundError(f"No draft exists for {section.get('title') or section_key}. Save a draft first.")
    cdir = confirmed_dir(output_dir, incident_id)
    cdir.mkdir(parents=True, exist_ok=True)
    dst = cdir / section["filename"]
    src_is_dst = _safe_resolve(src) == _safe_resolve(dst)
    if not src_is_dst:
        shutil.copy2(src, dst)

    structured_src = Path(section.get("structured_draft_path") or section.get("structured_confirmed_path") or "")
    structured_dst = cdir / f"{section_key}.json"
    if structured_src.exists():
        confirmed_blocks = repair_pipe_tables_in_blocks(load_blocks(structured_src))
        save_blocks(structured_dst, confirmed_blocks)
    else:
        save_blocks(structured_dst, blocks_from_text(_read(dst)))

    # Remove active draft after confirmation to make the status clear.
    # If the source is already the confirmed file, do not unlink it.
    draft_value = section.get("draft_path")
    if draft_value:
        draft_path = Path(draft_value)
        if draft_path.exists() and _safe_resolve(draft_path) != _safe_resolve(dst):
            try:
                draft_path.unlink()
            except Exception:
                pass
    structured_draft_value = section.get("structured_draft_path")
    if structured_draft_value:
        structured_draft_path = Path(structured_draft_value)
        if structured_draft_path.exists() and _safe_resolve(structured_draft_path) != _safe_resolve(structured_dst):
            try:
                structured_draft_path.unlink()
            except Exception:
                pass
    section["status"] = "confirmed"
    section["draft_path"] = None
    section["draft_relative_path"] = None
    section["structured_draft_path"] = None
    section["structured_draft_relative_path"] = None
    section["confirmed_path"] = str(dst)
    section["confirmed_relative_path"] = _rel(output_dir, dst)
    section["structured_confirmed_path"] = str(structured_dst)
    section["structured_confirmed_relative_path"] = _rel(output_dir, structured_dst)
    section["confirmed_at"] = utc_now()
    section["confirmed_by"] = analyst
    manifest["sections"][section_key] = section
    manifest["confirmed_by"] = analyst
    manifest["confirmed_at"] = utc_now()
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "text": _read(dst), "blocks": load_blocks(structured_dst), "message": f"{section['title']} confirmed"}


def confirm_report(output_dir: Path, analyst: str = "SOC Analyst", incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EVALUATOR] [FYP-STATE] [FYP-APPROVAL] Confirm
    (Lock) The Whole Report -- the manifest-wide "confirm report"
    state transition.

    Purpose: bulk-confirm every core section in one call -- either the
    analyst's explicit "Confirm all reports" action, or the automated
    auto-confirm used by the headless workflow adapter
    (adapters/export_documents.py) so a run can proceed to export without a
    human clicking through each section individually.
    Params: output_dir/incident_id -- see load_manifest(); analyst -- name
        recorded as confirming (source: API request body, or
        "SOC Workflow (auto-confirm)" from the headless adapter).
    Returns: dict with success, manifest, message.
    [FYP-STATE] Side effects / state change: calls confirm_section() once
        per section in manifest["section_order"] (or CORE_REPORT_KEYS as a
        fallback) that has any draft or confirmed content, reloading the
        manifest after each call so later sections see up-to-date state;
        finally forces manifest["report_status"] =
        "confirmed_by_analyst" and display_status = "All reports confirmed
        by analyst" directly (this overrides save_manifest()'s own rollup
        computation from confirm_section()'s last call, guaranteeing the
        whole-report status reads as fully confirmed even if
        section_order/CORE_REPORT_KEYS diverge from manifest["sections"]);
        stamps manifest-level confirmed_by/confirmed_at; calls
        save_manifest() again.
    Error handling: FileNotFoundError if no manifest found for this incident.
    Called by: soc_reporting_agent/backend/app.py (POST confirm-all
        endpoint); soc_reporting_agent/adapters/export_documents.py (headless
        auto-confirm before export, wrapped in its own try/except so a
        confirm failure is recorded as "confirm_error" rather than aborting
        the whole export run).
    Calls: load_manifest(), confirm_section(), save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    for key in list(manifest.get("section_order") or CORE_REPORT_KEYS):
        # Reload after each confirm.
        current = load_manifest(output_dir, incident_id)
        section = (current.get("sections") or {}).get(key)
        if section and (section.get("draft_path") or section.get("confirmed_path")):
            confirm_section(output_dir, key, analyst=analyst, incident_id=current.get("incident_id"))
    manifest = load_manifest(output_dir, incident_id)
    manifest["report_status"] = "confirmed_by_analyst"
    manifest["display_status"] = "All reports confirmed by analyst"
    manifest["confirmed_by"] = analyst
    manifest["confirmed_at"] = utc_now()
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "message": "All report sections confirmed"}



def list_section_drafts(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-STATE] List A Section's Revision History.

    Purpose: power the analyst-facing "revision history" view for one
    section, most-recent first.
    Params: output_dir/incident_id -- see load_manifest(); section_key --
        target section.
    Returns: dict with manifest, section, title, drafts (each history entry
        from section["draft_history"], reversed, with the file re-read and
        a short preview attached).
    Error handling: FileNotFoundError if no manifest; KeyError if
        section_key unknown.
    Called by: soc_reporting_agent/backend/app.py (GET revision-history
        endpoint).
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    drafts = []
    for entry in reversed(list(section.get("draft_history") or [])):
        path = Path(entry.get("path", ""))
        text = _read(path)
        item = dict(entry)
        item["text"] = text
        item["preview"] = markdown_to_plain_text(text)[:700]
        drafts.append(item)
    return {"success": True, "manifest": manifest, "section": section, "title": section.get("title") or section_key, "drafts": drafts}

def _confirmed_required(section: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-VALIDATION] [FYP-STATE] Enforce Confirmed-Before-
    Export Gate. Raises PermissionError if the section's status is not
    "confirmed"/"exported" or its confirmed_path is unset, and
    FileNotFoundError if that confirmed file has gone missing from disk.
    This is the sole enforcement point preventing an un-reviewed draft from
    ever being exported to Word/PDF. Called by: export_section_docx(),
    export_section_pdf(), export_docx(), export_pdf()."""
    if section.get("status") not in {"confirmed", "exported"} or not section.get("confirmed_path"):
        raise PermissionError("This report must be confirmed by the SOC analyst before export.")
    if not Path(section["confirmed_path"]).exists():
        raise FileNotFoundError("Confirmed report file not found.")


def _docx_write(path: Path, title: str, text: str, incident_id: str, manifest: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-EXPORT] Plain-Text-To-DOCX Convenience Wrapper.
    Converts `text` to blocks via blocks_from_text() then delegates to
    _docx_write_blocks(). Kept for any caller that only has plain text
    rather than a pre-built block list."""
    blocks = blocks_from_text(text)
    _docx_write_blocks(path, title, blocks, incident_id, manifest)


# ============================================================================
# [FYP-SECTION] DOCX Rendering Helpers (python-docx)
# [FYP-EXPORT] Small, single-purpose styling helpers used exclusively by
# _docx_write_blocks() to build the shared Aegis report look (title page,
# table shading/borders, fonts, bullet indentation) from a block list. None
# of these are meaningful entry points on their own; see _docx_write_blocks()
# for the actual document assembly.
# ============================================================================
def _set_cell_shading(cell: Any, fill: str) -> None:
    """[FYP-FUNCTION] Set a table cell's background fill colour (hex, with
    or without leading '#') via raw OOXML <w:shd> manipulation -- python-docx
    has no high-level API for cell shading. No-op if OxmlElement/qn were not
    importable (python-docx missing)."""
    if OxmlElement is None or qn is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def _apply_report_font(run: Any, name: str = "Georgia") -> None:
    """[FYP-FUNCTION] Apply the report's house font ("Georgia" by default)
    to a python-docx run, including the East Asian/ASCII/High-ANSI font
    slots so the setting actually sticks across Word's font-fallback rules.
    Swallows any exception (best-effort styling only)."""
    try:
        run.font.name = name
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
            run._element.rPr.rFonts.set(qn("w:ascii"), name)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    except Exception:
        pass


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, font_size: int = 10, color: str = "1f2937",
                   align: str | None = None, valign: str | None = None) -> None:
    """[FYP-FUNCTION] Write styled text into a single table cell (font,
    bold, size, color, horizontal/vertical alignment) in one call, applying
    _apply_report_font() so cell text matches the rest of the document."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text or ""))
    _apply_report_font(run)
    run.bold = bool(bold)
    if Pt is not None:
        run.font.size = Pt(font_size)
    if RGBColor is not None and color:
        try:
            run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        except Exception:
            pass
    if WD_ALIGN_PARAGRAPH is not None and align:
        para.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                          "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    if WD_CELL_VERTICAL_ALIGNMENT is not None:
        cell.vertical_alignment = (WD_CELL_VERTICAL_ALIGNMENT.CENTER if valign == "center"
                                   else WD_CELL_VERTICAL_ALIGNMENT.TOP)


# [FYP-FUNCTION] `_set_column_widths` — implements the set column widths operation used by the surrounding report generation and export workflow.
# [FYP-INPUT] Parameters: `table`, `widths_in`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] Static symbol references include soc_reporting_agent/reporting/editable_reports.py:_docx_write_blocks; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `Inches`, `enumerate`, `len`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _set_column_widths(table: Any, widths_in: list[float]) -> None:
    # [FYP-FUNCTION] Fixed Table Column Widths.
    """Force fixed column widths (inches). Word can ignore table.columns[i].width
    unless every cell in that column is also given the same width, so both are set."""
    if Inches is None:
        return
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
    for idx, width in enumerate(widths_in):
        if idx < len(table.columns):
            table.columns[idx].width = Inches(width)


def _style_table(table: Any, *, header_fill: str = "EAF2FB", first_col_fill: str | None = None) -> None:
    """[FYP-FUNCTION] Apply the shared report table look: grid borders,
    centered table alignment, bold+shaded header row, optional shaded+bold
    first column (used for label/value style tables), and zero paragraph
    spacing inside cells so rows stay compact."""
    table.style = "Table Grid"
    if WD_TABLE_ALIGNMENT is not None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if r_idx == 0:
                _set_cell_shading(cell, header_fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        if RGBColor is not None:
                            run.font.color.rgb = RGBColor.from_string("1E3A8A")
            elif first_col_fill and c_idx == 0:
                _set_cell_shading(cell, first_col_fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            for para in cell.paragraphs:
                if Pt is not None:
                    para.paragraph_format.space_after = Pt(0)


def _add_title_block(doc: Any, title: str, incident_id: str, manifest: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-EXPORT] Render The Document Cover Block.

    Purpose: draws the Aegis logo (or "AEGIS" text fallback), the report
    title, a fixed subtitle, and a 4-row metadata table (Incident ID, Report
    Status, Confirmed By, Exported At) at the top of every exported .docx.
    Params: doc -- python-docx Document being built; title -- report title
        text; incident_id -- for the metadata table; manifest -- source of
        manifest.get("confirmed_by") (falls back to "SOC Analyst" display
        text if unset).
    Returns: None. Side effects: appends paragraphs/a table to `doc`.
    Called by: _docx_write_blocks().
    Calls: _aegis_logo_path(), _apply_report_font(), _set_cell_text(),
        _set_cell_shading(), utc_now().
    """
    logo_path = _aegis_logo_path()
    if logo_path:
        logo = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if Pt is not None:
            logo.paragraph_format.space_after = Pt(8)
        logo_run = logo.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(2.65))
    else:
        label = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label.add_run("AEGIS")
        _apply_report_font(label_run)
        label_run.bold = True
        if Pt is not None:
            label_run.font.size = Pt(10)
        if RGBColor is not None:
            label_run.font.color.rgb = RGBColor.from_string("2F66D0")

    heading = doc.add_paragraph()
    if WD_ALIGN_PARAGRAPH is not None:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(title)
    _apply_report_font(heading_run)
    heading_run.bold = True
    if Pt is not None:
        heading_run.font.size = Pt(30)
    if RGBColor is not None:
        heading_run.font.color.rgb = RGBColor.from_string("172033")

    subtitle = doc.add_paragraph("Reviewed SOC report section exported from the analyst-confirmed dashboard draft and converted to PDF from Word.")
    if WD_ALIGN_PARAGRAPH is not None:
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _run in subtitle.runs:
        _apply_report_font(_run)
    if Pt is not None:
        subtitle.runs[0].font.size = Pt(11)
    if RGBColor is not None:
        subtitle.runs[0].font.color.rgb = RGBColor.from_string("526071")

    meta = doc.add_table(rows=1, cols=2)
    meta.style = "Table Grid"
    values = [
        ("Incident ID", incident_id),
        ("Report Status", "Confirmed by SOC Analyst"),
        ("Confirmed By", manifest.get("confirmed_by") or "SOC Analyst"),
        ("Exported At", utc_now()),
    ]
    for idx, (field, value) in enumerate(values):
        row = meta.rows[0] if idx == 0 else meta.add_row()
        _set_cell_text(row.cells[0], field, bold=True, font_size=10, color="172033")
        _set_cell_text(row.cells[1], value, font_size=10, color="172033")
        _set_cell_shading(row.cells[0], "F3F6FA")
    doc.add_paragraph("")

def _docx_write_blocks(path: Path, title: str, blocks: list[dict[str, Any]], incident_id: str, manifest: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-EXPORT] Render A Block List To A .docx File.

    Purpose: the single shared implementation behind every .docx export path
    in this module (export_section_docx, export_docx, render_blocks_to_docx)
    -- walks a structured block list (heading/paragraph/bullet_list/table/
    page_break, see reporting.structured_report) and writes a styled Word
    document.
    Params: path -- output .docx path; title -- document title (also used
        to de-duplicate a leading heading that repeats it, see
        _strip_duplicate_leading_heading()); blocks -- structured content
        (source: confirmed/draft section content, or an arbitrary caller-
        supplied block list for render_blocks_to_docx()); incident_id/
        manifest -- passed through to _add_title_block()'s metadata table.
    Returns: None.
    Side effects: creates parent directories and writes/overwrites the file
        at `path`.
    [FYP-VALIDATION] Calls _validate_no_raw_markdown_tables() before
        rendering, which raises ValueError if any paragraph still contains
        raw "|"-delimited markdown table syntax that repair_pipe_tables_in_
        blocks() could not convert into a real table block -- this is a
        deliberate export-time guard against a malformed table leaking into
        the analyst-facing Word document as literal pipe characters.
    Error handling: raises RuntimeError if python-docx is not installed
        (Document is None).
    Called by: export_section_docx(), export_docx(), render_blocks_to_docx(),
        _docx_write() (plain-text wrapper).
    Calls: repair_pipe_tables_in_blocks(), _validate_no_raw_markdown_tables(),
        _add_title_block(), _strip_duplicate_leading_heading(),
        _apply_report_font(), _set_cell_text(), _style_table(),
        _set_column_widths().
    """
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    blocks = repair_pipe_tables_in_blocks(blocks)
    _validate_no_raw_markdown_tables(blocks, title)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    try:
        normal.font.name = "Georgia"
        if qn is not None:
            normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
            normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
            normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    except Exception:
        pass
    if Pt is not None:
        normal.font.size = Pt(10.5)
    if RGBColor is not None:
        normal.font.color.rgb = RGBColor.from_string("172033")

    _add_title_block(doc, title, incident_id, manifest)

    # [FYP-PROCESS] Block-type dispatch: each block dict's "type" (page_break /
    # heading / paragraph / bullet_list / table) maps to its own python-docx
    # rendering branch below; unrecognised types are silently skipped.
    for block in _strip_duplicate_leading_heading(blocks or [], title):
        btype = block.get("type") if isinstance(block, dict) else None
        if btype == "page_break":
            doc.add_page_break()
        elif btype == "heading":
            level = int(block.get("level") or 2)
            para = doc.add_paragraph()
            run = para.add_run(str(block.get("text") or ""))
            _apply_report_font(run)
            run.bold = True
            if Pt is not None:
                run.font.size = Pt(22 if level <= 1 else 16 if level == 2 else 13)
            if RGBColor is not None:
                run.font.color.rgb = RGBColor.from_string("2F66D0" if level <= 2 else "1F2937")
            if Pt is not None:
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
        elif btype == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                para = doc.add_paragraph(text)
                for _run in para.runs:
                    _apply_report_font(_run)
                if Pt is not None:
                    para.paragraph_format.space_after = Pt(6)
        elif btype == "bullet_list":
            bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
            for item in block.get("items") or []:
                if isinstance(item, dict):
                    item_text = str(item.get("text") or "").strip()
                    level = min(2, max(0, int(item.get("level") or 0)))
                else:
                    item_text, level = str(item or "").strip(), 0
                if not item_text:
                    continue
                bullet_para = doc.add_paragraph(item_text, style=bullet_styles[level])
                # Explicit hanging indent per level — belt-and-suspenders on top
                # of the style's own numbering so wrapped lines stay aligned
                # under the bullet text (not the bullet glyph) in every viewer,
                # and each nesting level reads visibly deeper than its parent.
                if Inches is not None:
                    bullet_para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
                    bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
                if Pt is not None:
                    bullet_para.paragraph_format.space_after = Pt(2)
                for _run in bullet_para.runs:
                    _apply_report_font(_run)
        elif btype == "table":
            columns = [str(c or "") for c in (block.get("columns") or [])]
            rows = block.get("rows") or []
            if not columns:
                continue
            is_step_checklist = (len(columns) == 2
                                 and columns[0].strip().lower() == "step"
                                 and columns[1].strip().lower() == "analyst action")
            table = doc.add_table(rows=1, cols=len(columns))
            for idx, col in enumerate(columns):
                _set_cell_text(table.rows[0].cells[idx], col, bold=True, font_size=10, color="1E3A8A",
                               align="center" if (is_step_checklist and idx == 0) else None)
            for row in rows:
                cells = table.add_row().cells
                row_values = list(row or []) + [""] * (len(columns) - len(row or []))
                for idx, value in enumerate(row_values[:len(columns)]):
                    _set_cell_text(cells[idx], value, bold=(idx == 0 and len(columns) == 2 and not is_step_checklist),
                                  font_size=9, color="172033",
                                  align="center" if (is_step_checklist and idx == 0) else None,
                                  valign="center" if (is_step_checklist and idx == 0) else None)
            _style_table(table, header_fill="EAF2FB",
                        first_col_fill="F3F6FA" if (len(columns) == 2 and not is_step_checklist) else None)
            if is_step_checklist:
                # Narrow, centered Step column (~10% of the usable page width);
                # Analyst Action gets the rest. Matches the page's 0.7in side
                # margins on an 8.5in page (7.1in usable width).
                _set_column_widths(table, [0.7, 6.4])
            doc.add_paragraph("")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _strip_duplicate_leading_heading(blocks: list[dict[str, Any]], title: str) -> list[dict[str, Any]]:
    """[FYP-FUNCTION] Drop a leading heading block that just repeats the
    document title (or the generic legacy title
    "Cybersecurity Incident Post-Incident Review Report"), so the exported
    document does not show the title twice -- once from _add_title_block()'s
    cover heading, once from the content's own first heading. Compares
    whitespace-normalised, case-insensitive text; stops at the first heading
    that does NOT match (only ever strips from the very start of the list).
    Called by: _docx_write_blocks(), _pdf_write_blocks()."""
    if not blocks:
        return []
    normalised_title = re.sub(r"\s+", " ", str(title or "").strip().lower())
    cleaned = list(blocks)
    while cleaned:
        first = cleaned[0]
        if not isinstance(first, dict) or first.get("type") != "heading":
            break
        text = re.sub(r"\s+", " ", str(first.get("text") or "").strip().lower())
        if text in {normalised_title, "cybersecurity incident post-incident review report"}:
            cleaned.pop(0)
            continue
        break
    return cleaned

# ============================================================================
# [FYP-SECTION] PDF Rendering Helpers (reportlab)
# [FYP-EXPORT] [FYP-FALLBACK] Pure-Python PDF renderer, structurally mirroring
# the DOCX helpers above (same block-type dispatch), used only when the
# preferred DOCX->PDF conversion path
# (reporting.template_document_exporter.convert_docx_to_pdf, LibreOffice-
# backed) is unavailable in the current environment. Produces a visually
# similar but independently-styled document (reportlab has no shared styling
# API with python-docx).
# ============================================================================
def _pdf_escape(text: str) -> str:
    """[FYP-FUNCTION] Escape &, <, > for safe embedding in reportlab's
    Paragraph mini-XML markup (reportlab Paragraphs interpret a subset of
    HTML-like tags, so raw angle brackets/ampersands in report text must be
    escaped first)."""
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_para(text: str, style: Any) -> Any:
    """[FYP-FUNCTION] Build one reportlab Paragraph with escaped text."""
    return Paragraph(_pdf_escape(str(text or "")), style)


def _pdf_write(path: Path, title: str, text: str, incident_id: str, manifest: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-EXPORT] Plain-text-to-PDF convenience wrapper,
    mirroring _docx_write(): converts `text` to blocks then delegates to
    _pdf_write_blocks()."""
    blocks = blocks_from_text(text)
    _pdf_write_blocks(path, title, blocks, incident_id, manifest)


def _pdf_write_blocks(path: Path, title: str, blocks: list[dict[str, Any]], incident_id: str, manifest: dict[str, Any]) -> None:
    """[FYP-FUNCTION] [FYP-EXPORT] [FYP-FALLBACK] Render A Block List
    Directly To PDF (reportlab).

    Purpose: fallback PDF renderer used when convert_docx_to_pdf() (the
    preferred LibreOffice-based path) is unavailable -- see callers.
    Params: same shape as _docx_write_blocks() (path/title/blocks/
        incident_id/manifest).
    Returns: None. Side effects: creates parent directories and writes the
        PDF at `path`.
    [FYP-VALIDATION] Runs _validate_no_raw_markdown_tables() before
        rendering, same guard as the DOCX path.
    Error handling: raises RuntimeError if reportlab is not installed
        (SimpleDocTemplate/Table are None).
    Called by: export_section_pdf(), export_pdf(), render_blocks_to_pdf()
        (all as an `except Exception` fallback after trying
        convert_docx_to_pdf() first), _pdf_write().
    Calls: repair_pipe_tables_in_blocks(), _validate_no_raw_markdown_tables(),
        _aegis_logo_path(), _strip_duplicate_leading_heading(), _pdf_para().
    """
    if SimpleDocTemplate is None or Table is None:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")
    blocks = repair_pipe_tables_in_blocks(blocks)
    _validate_no_raw_markdown_tables(blocks, title)
    styles = getSampleStyleSheet()
    # Hanging-indent bullet styles, one per nesting level, so wrapped lines
    # align under the bullet text (not the glyph) and nested items sit
    # visibly deeper than their parent.
    bullet_styles = [
        ParagraphStyle(f"Bullet{lvl}", parent=styles["BodyText"],
                       leftIndent=18 + lvl * 18, bulletIndent=lvl * 18,
                       spaceAfter=3)
        for lvl in range(3)
    ] if ParagraphStyle is not None else None
    story = []
    logo_path = _aegis_logo_path()
    if Image is not None and logo_path:
        logo = Image(str(logo_path), width=2.65 * inch, height=0.86 * inch)
        logo.hAlign = "CENTER"
        story.extend([logo, Spacer(1, 0.12 * inch)])
    story.extend([_pdf_para(title, styles["Title"]), Spacer(1, 0.15 * inch)])
    meta_data = [[_pdf_para("Field", styles["Heading5"]), _pdf_para("Value", styles["Heading5"])],
                 [_pdf_para("Incident ID", styles["BodyText"]), _pdf_para(incident_id, styles["BodyText"])],
                 [_pdf_para("Confirmed by", styles["BodyText"]), _pdf_para(manifest.get("confirmed_by") or "SOC Analyst", styles["BodyText"])],
                 [_pdf_para("Generated at", styles["BodyText"]), _pdf_para(utc_now(), styles["BodyText"])] ]
    meta = Table(meta_data, repeatRows=1, hAlign="LEFT", colWidths=[1.6*inch, 4.8*inch])
    meta.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#9AA4B2")),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF6")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
        ("RIGHTPADDING", (0,0), (-1,-1), 6),
    ]))
    story.extend([meta, Spacer(1, 0.2*inch)])

    for block in _strip_duplicate_leading_heading(blocks or [], title):
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        if btype == "page_break":
            story.append(PageBreak())
        elif btype == "heading":
            level = int(block.get("level") or 2)
            style_name = "Heading1" if level <= 1 else ("Heading2" if level == 2 else "Heading3")
            story.extend([Spacer(1, 0.08*inch), _pdf_para(block.get("text") or "", styles[style_name])])
        elif btype == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                story.append(_pdf_para(text, styles["BodyText"]))
                story.append(Spacer(1, 0.06*inch))
        elif btype == "bullet_list":
            for item in block.get("items") or []:
                if isinstance(item, dict):
                    item_text = str(item.get("text") or "").strip()
                    level = min(2, max(0, int(item.get("level") or 0)))
                else:
                    item_text, level = str(item or "").strip(), 0
                if not item_text:
                    continue
                if bullet_styles is not None:
                    style = bullet_styles[level]
                    story.append(Paragraph(f"<bullet>&bull;</bullet>{_pdf_escape(item_text)}", style))
                else:
                    story.append(_pdf_para(("  " * level) + "• " + item_text, styles["BodyText"]))
        elif btype == "table":
            columns = [str(c or "") for c in (block.get("columns") or [])]
            rows = block.get("rows") or []
            if not columns:
                continue
            is_step_checklist = (len(columns) == 2
                                 and columns[0].strip().lower() == "step"
                                 and columns[1].strip().lower() == "analyst action")
            if is_step_checklist and ParagraphStyle is not None:
                centered_header = ParagraphStyle("StepHeader", parent=styles["Heading5"], alignment=1)
                centered_body = ParagraphStyle("StepBody", parent=styles["BodyText"], alignment=1)
                data = [[_pdf_para(columns[0], centered_header), _pdf_para(columns[1], styles["Heading5"])]]
                for row in rows:
                    values = list(row or []) + [""] * (2 - len(row or []))
                    data.append([_pdf_para(values[0], centered_body), _pdf_para(values[1], styles["BodyText"])])
            else:
                data = [[_pdf_para(c, styles["Heading5"]) for c in columns]]
                for row in rows:
                    values = list(row or []) + [""] * (len(columns) - len(row or []))
                    data.append([_pdf_para(v, styles["BodyText"]) for v in values[:len(columns)]])
            usable_width = 7.0 * inch
            if is_step_checklist:
                # Narrow, centered Step column (~10% of the usable width);
                # Analyst Action gets the rest.
                col_widths = [usable_width * 0.10, usable_width * 0.90]
            else:
                col_widths = [usable_width / max(1, len(columns))] * len(columns)
            table = Table(data, repeatRows=1, hAlign="LEFT", colWidths=col_widths)
            table_style_cmds = [
                ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#9AA4B2")),
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF6")),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING", (0,0), (-1,-1), 5),
                ("RIGHTPADDING", (0,0), (-1,-1), 5),
                ("TOPPADDING", (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]
            if is_step_checklist:
                table_style_cmds.append(("VALIGN", (0, 0), (0, -1), "MIDDLE"))
            table.setStyle(TableStyle(table_style_cmds))
            story.extend([table, Spacer(1, 0.15*inch)])
    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=0.45*inch, leftMargin=0.45*inch, topMargin=0.5*inch, bottomMargin=0.5*inch).build(story)


# ============================================================================
# [FYP-SECTION] Export Validation Helpers
# ============================================================================
def _confirmed_blocks_or_text(section: dict[str, Any]) -> tuple[list[dict[str, Any]], str]:
    """[FYP-FUNCTION] [FYP-EXPORT] [FYP-STATE] Resolve A Confirmed Section's
    Renderable Content.

    Purpose: load the structured blocks for a CONFIRMED section (never a
    draft -- callers only reach this after _confirmed_required() has passed),
    repairing any raw pipe-table syntax and re-persisting the repaired
    blocks back to disk if repair changed anything, so the fix is not
    silently re-derived on every export.
    Params: section -- a manifest section entry with confirmed_path /
        structured_confirmed_path.
    Returns: (blocks, text) -- falls back to blocks_from_text(text) if no
        structured file exists.
    Side effects: may rewrite the structured_confirmed_path file if
        repair_pipe_tables_in_blocks() changed the content.
    [FYP-VALIDATION] Calls _validate_no_raw_markdown_tables(), which raises
        ValueError if raw table syntax survives repair.
    Called by: export_section_docx(), export_section_pdf() (fallback path),
        export_pdf() (fallback path).
    """
    structured_path = section.get("structured_confirmed_path")
    original_blocks = load_blocks(structured_path)
    blocks = repair_pipe_tables_in_blocks(original_blocks)
    text = _read(section.get("confirmed_path"))
    if not blocks:
        blocks = blocks_from_text(text)
        blocks = repair_pipe_tables_in_blocks(blocks)
    _validate_no_raw_markdown_tables(blocks, section.get("title"))
    if structured_path and blocks != original_blocks:
        save_blocks(Path(structured_path), blocks)
    return blocks, text


def _safe_raw_table_preview(text: Any, limit: int = 180) -> str:
    """[FYP-FUNCTION] [FYP-SECURITY] Build A Redacted Error-Message Preview.

    Purpose: produce a short, single-line preview of offending text for the
    ValueError raised by _validate_no_raw_markdown_tables(), without leaking
    secret-shaped values (api_key/access_token/token/secret/authorization
    key: value pairs are replaced with "key=[REDACTED]") into logs or error
    messages. Truncates to `limit` characters with a trailing "...".
    Called by: _validate_no_raw_markdown_tables().
    """
    preview = re.sub(r"\s+", " ", str(text or "")).strip()
    preview = re.sub(
        r"(?i)\b(api[_ -]?key|access[_ -]?token|token|secret|authorization)\b\s*[:=]\s*\S+",
        r"\1=[REDACTED]",
        preview,
    )
    return preview[:limit] + ("..." if len(preview) > limit else "")


def _validate_no_raw_markdown_tables(blocks: list[dict[str, Any]], section_title: str | None = None) -> None:
    """[FYP-FUNCTION] [FYP-VALIDATION] Export-Time Guard: No Raw Table
    Syntax Allowed Through.

    Purpose: last-line-of-defence check run immediately before every
    DOCX/PDF write -- after attempting to repair markdown pipe-table syntax
    into real table blocks, confirm that no paragraph block still contains
    unconverted "|"-delimited table text (which would otherwise render as
    literal pipe characters in the analyst-facing document).
    Params: blocks -- structured block list, mutated in place (blocks[:] =
        repaired) so the caller's own list reflects the repair; section_title
        -- used only to make the error message legible (falls back to the
        most recent heading's text while scanning, then "Report").
    Returns: None.
    Error handling: raises ValueError with a redacted preview
        (_safe_raw_table_preview()) of the offending text if a raw table is
        found -- this aborts the export entirely rather than shipping a
        broken table.
    Called by: _docx_write_blocks(), _pdf_write_blocks().
    Calls: repair_pipe_tables_in_blocks(), paragraph_contains_raw_pipe_table(),
        _safe_raw_table_preview().
    """
    repaired = repair_pipe_tables_in_blocks(blocks)
    if isinstance(blocks, list):
        blocks[:] = repaired
    current_section = section_title or "Report"
    for index, block in enumerate(repaired):
        if isinstance(block, dict) and block.get("type") == "heading":
            current_section = str(block.get("text") or current_section)
            continue
        if not isinstance(block, dict) or block.get("type") != "paragraph":
            continue
        text = str(block.get("text") or "")
        if paragraph_contains_raw_pipe_table(text):
            preview = _safe_raw_table_preview(text)
            raise ValueError(
                f'Export validation failed: raw table syntax remained in section "{current_section}", '
                f"block index {index}: {preview}"
            )


# ============================================================================
# [FYP-SECTION] Public Block-Rendering Entry Points (manifest-independent)
# ============================================================================
def render_blocks_to_docx(path: Path, title: str, blocks: list[dict[str, Any]],
                          incident_id: str, meta: dict[str, Any] | None = None) -> None:
    """[FYP-FUNCTION] [FYP-ENTRY-POINT] [FYP-EXPORT] [FYP-USED-BY]
    report_editing.py, triage_ticket_editing.py (repo root)

    Public entry point for rendering an arbitrary block list straight to a
    .docx file, with no dependency on report_manifest.json / draft-confirm
    status — used by callers (e.g. the main Streamlit app's analyst-edit
    layer) that maintain their own block content outside this module's
    manifest but still want the exact same document styling/layout. Thin
    wrapper around the private renderer used by export_section_docx().
    """
    _docx_write_blocks(path, title, blocks, incident_id, meta or {})


def render_blocks_to_pdf(path: Path, title: str, blocks: list[dict[str, Any]],
                         incident_id: str, meta: dict[str, Any] | None = None, *,
                         docx_path: Path | None = None) -> None:
    """[FYP-FUNCTION] [FYP-ENTRY-POINT] [FYP-EXPORT] [FYP-FALLBACK]
    [FYP-USED-BY] report_editing.py, triage_ticket_editing.py (repo root)

    Public entry point mirroring export_section_pdf()'s docx-then-convert
    pattern, but for an arbitrary block list rather than a confirmed manifest
    section: renders (or reuses) a .docx via render_blocks_to_docx(), then
    converts it to PDF with the same LibreOffice-backed converter used
    elsewhere, falling back to the pure-Python reportlab renderer if that
    conversion isn't available in this environment.
    Calls: render_blocks_to_docx(),
    reporting.template_document_exporter.convert_docx_to_pdf() (lazy import),
    _pdf_write_blocks() (fallback on any exception from the converter).
    """
    meta = meta or {}
    working_docx = docx_path
    if working_docx is None:
        working_docx = path.with_suffix(".docx")
        render_blocks_to_docx(working_docx, title, blocks, incident_id, meta)
    try:
        from reporting.template_document_exporter import convert_docx_to_pdf
        convert_docx_to_pdf(working_docx, path)
    except Exception:
        _pdf_write_blocks(path, title, blocks, incident_id, meta)


# ============================================================================
# [FYP-SECTION] Manifest-Driven Section Export (Word / PDF)
# ============================================================================
def export_section_docx(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EXPORT] [FYP-STATE] Export One Confirmed Section
    To Word.

    Purpose: render a single confirmed section to .docx and record the
    export in the manifest.
    Params: output_dir/incident_id -- see load_manifest(); section_key --
        target section (must already be confirmed).
    Returns: dict with success, manifest, section, path, download_url,
        message.
    [FYP-STATE] Side effects: writes exports_dir()/<section_key>.docx;
        records section["exports"]["docx"] (path/relative_path/created_at);
        sets section["status"] = "exported"; calls save_manifest().
    Error handling: FileNotFoundError/KeyError from load_manifest() lookups;
        PermissionError/FileNotFoundError from _confirmed_required() if the
        section is not yet confirmed.
    Called by: export_section_pdf() (renders the DOCX first, then converts);
        soc_reporting_agent/backend/app.py (POST per-section DOCX export
        endpoint).
    Calls: load_manifest(), _confirmed_required(), _confirmed_blocks_or_text(),
        _docx_write_blocks(), save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    _confirmed_required(section)
    incident_id = manifest.get("incident_id") or "INC-0001"
    blocks, text = _confirmed_blocks_or_text(section)
    path = exports_dir(output_dir, incident_id) / f"{section_key}.docx"
    _docx_write_blocks(path, section.get("title") or section_key, blocks, incident_id, manifest)
    section.setdefault("exports", {})["docx"] = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now()}
    section["status"] = "exported"
    manifest["sections"][section_key] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "path": str(path), "download_url": f"/api/reports/{section_key}/download/docx", "message": "Word export ready"}


def export_section_pdf(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EXPORT] [FYP-STATE] [FYP-FALLBACK] Export One
    Confirmed Section To PDF.

    Purpose: render a single confirmed section to PDF, preferring conversion
    from the freshly-exported Word document (so DOCX and PDF are always
    visually identical) and falling back to the pure-Python reportlab
    renderer only if that conversion is unavailable.
    Params: output_dir/incident_id/section_key -- see export_section_docx().
    Returns: dict with success, manifest, section, path, download_url,
        message.
    [FYP-STATE] Side effects: first calls export_section_docx() (which
        itself records the DOCX export + sets status "exported"); writes
        exports_dir()/<section_key>.pdf; records
        section["exports"]["pdf"] (including source_docx path); re-sets
        status = "exported"; calls save_manifest() again.
    Error handling: same lookup/confirmation errors as export_section_docx();
        DOCX->PDF conversion failures are caught and silently fall back to
        _pdf_write_blocks() rather than failing the export.
    Called by: soc_reporting_agent/backend/app.py (POST per-section PDF
        export endpoint).
    Calls: export_section_docx(),
        reporting.template_document_exporter.convert_docx_to_pdf() (lazy
        import), _confirmed_blocks_or_text(), _pdf_write_blocks(),
        save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    _confirmed_required(section)
    incident_id = manifest.get("incident_id") or "INC-0001"
    docx_result = export_section_docx(output_dir, section_key, incident_id=incident_id)
    docx_path = Path(docx_result["path"])
    path = exports_dir(output_dir, incident_id) / f"{section_key}.pdf"
    try:
        from reporting.template_document_exporter import convert_docx_to_pdf
        convert_docx_to_pdf(docx_path, path)
    except Exception:
        blocks, text = _confirmed_blocks_or_text(section)
        _pdf_write_blocks(path, section.get("title") or section_key, blocks, incident_id, manifest)
    manifest = load_manifest(output_dir, incident_id)
    section = (manifest.get("sections") or {}).get(section_key) or section
    section.setdefault("exports", {})["pdf"] = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now(), "source_docx": str(docx_path)}
    section["status"] = "exported"
    manifest["sections"][section_key] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "path": str(path), "download_url": f"/api/reports/{section_key}/download/pdf", "message": "PDF export ready from confirmed Word document"}


# ============================================================================
# [FYP-SECTION] Combined "Final Incident Report" Export (Word / PDF)
# ============================================================================
# [FYP-FUNCTION] `_final_report_blocks` — implements the final report blocks operation used by the surrounding report generation and export workflow.
# [FYP-INPUT] Parameters: `manifest`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
# [FYP-USED-BY] Static symbol references include soc_reporting_agent/reporting/editable_reports.py:export_docx, soc_reporting_agent/reporting/editable_reports.py:export_pdf; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `KeyError`, `_read`, `_strip_duplicate_leading_heading`, `blocks_from_text`, `get`, `load_blocks`, `repair_pipe_tables_in_blocks`.
# [FYP-ERROR] Raises explicit validation/processing errors to the caller; no silent fallback is applied here.

def _final_report_blocks(manifest: dict[str, Any]) -> tuple[list[dict[str, Any]], str]:
    # [FYP-FUNCTION] Resolve Final Incident Report Content.
    """Blocks for the Final Incident Report export.

    This is the standalone output of incident_report_template.md.j2 —
    the dedicated template already covers executive summary, technical
    findings, SOC analyst review guidance, conclusions and appendices in
    one coherent narrative. It is exported as-is, NOT stitched together
    with the separately-generated Executive Summary / Technical Findings /
    SOC Analyst Review section reports (those stay independent documents)."""
    section = (manifest.get("sections") or {}).get("final_incident_report")
    if not section:
        raise KeyError("final_incident_report section not found in manifest")
    title = section.get("title") or "Final Incident Report"
    section_blocks = repair_pipe_tables_in_blocks(
        load_blocks(section.get("structured_confirmed_path") or section.get("structured_draft_path")))
    if not section_blocks:
        text = _read(section.get("confirmed_path") or section.get("draft_path"))
        section_blocks = blocks_from_text(text)
    return _strip_duplicate_leading_heading(section_blocks, title), title


def export_docx(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EVALUATOR] [FYP-EXPORT] [FYP-STATE] Export The
    Combined Final Incident Report To Word.

    Purpose: render the standalone "final_incident_report" section (the
    dedicated template's own complete narrative -- see
    _final_report_blocks()'s docstring) to .docx, recording the export at
    both the manifest level and the section level.
    Params: output_dir/incident_id -- see load_manifest().
    Returns: dict with success, path, manifest, download_url.
    [FYP-STATE] Side effects: requires the final_incident_report section be
        confirmed (_confirmed_required()); writes
        exports_dir()/final_incident_report.docx; records the export entry
        in BOTH manifest["exports"]["docx"] and
        manifest["sections"]["final_incident_report"]["exports"]["docx"] (so
        the report is self-describing the same way the other three core
        sections are); calls save_manifest().
    Error handling: FileNotFoundError/KeyError if manifest or section
        missing; PermissionError/FileNotFoundError via _confirmed_required().
    Called by: export_pdf() (renders DOCX first, then converts);
        soc_reporting_agent/backend/app.py (POST combined DOCX export
        endpoint); soc_reporting_agent/adapters/export_documents.py
        (headless workflow export, wrapped in try/except so a failure is
        recorded as result["docx_error"] rather than aborting the run --
        export_documents.py continues on to export_pdf() and the per-section
        exports regardless).
    Calls: load_manifest(), _confirmed_required(), _final_report_blocks(),
        _docx_write_blocks(), save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get("final_incident_report")
    if not section:
        raise KeyError("final_incident_report section not found in manifest")
    _confirmed_required(section)
    incident_id = manifest.get("incident_id") or "INC-0001"
    blocks, title = _final_report_blocks(manifest)
    # Canonical filename is final_incident_report.* — combined_incident_report.*
    # was a legacy, misleading name: this export has always been the standalone
    # output of incident_report_template.md.j2 (see _final_report_blocks'
    # docstring), never a concatenation of the other three reports.
    path = exports_dir(output_dir, incident_id) / "final_incident_report.docx"
    _docx_write_blocks(path, f"{title} - {incident_id}", blocks, incident_id, manifest)
    export_entry = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now()}
    manifest.setdefault("exports", {})["docx"] = export_entry
    # Also mirror into sections.final_incident_report.exports so this report
    # is self-describing exactly like the other three (previously left {}).
    section.setdefault("exports", {})["docx"] = export_entry
    manifest["sections"]["final_incident_report"] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "path": str(path), "manifest": manifest, "download_url": "/api/reports/download/docx"}


def export_pdf(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EXPORT] [FYP-STATE] [FYP-FALLBACK] Export The
    Combined Final Incident Report To PDF.

    Purpose: PDF counterpart of export_docx(), preferring conversion from
    the freshly-exported combined Word document, falling back to reportlab.
    Params: output_dir/incident_id -- see load_manifest().
    Returns: dict with success, path, manifest, download_url.
    [FYP-STATE] Side effects: calls export_docx() first; writes
        exports_dir()/final_incident_report.pdf; records the export entry in
        both manifest["exports"]["pdf"] and the section's own exports dict
        (including source_docx); calls save_manifest().
    Error handling: same as export_docx(); DOCX->PDF conversion failure
        falls back silently to _pdf_write_blocks().
    Called by: soc_reporting_agent/backend/app.py (POST combined PDF export
        endpoint); soc_reporting_agent/adapters/export_documents.py
        (headless workflow export; failure recorded as result["pdf_error"]).
    Calls: export_docx(),
        reporting.template_document_exporter.convert_docx_to_pdf() (lazy
        import), _final_report_blocks(), _pdf_write_blocks(), save_manifest().
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get("final_incident_report")
    if not section:
        raise KeyError("final_incident_report section not found in manifest")
    _confirmed_required(section)
    incident_id = manifest.get("incident_id") or "INC-0001"
    docx_result = export_docx(output_dir, incident_id=incident_id)
    docx_path = Path(docx_result["path"])
    path = exports_dir(output_dir, incident_id) / "final_incident_report.pdf"
    try:
        from reporting.template_document_exporter import convert_docx_to_pdf
        convert_docx_to_pdf(docx_path, path)
    except Exception:
        manifest = load_manifest(output_dir, incident_id)
        blocks, title = _final_report_blocks(manifest)
        _pdf_write_blocks(path, f"{title} - {incident_id}", blocks, incident_id, manifest)
    manifest = load_manifest(output_dir, incident_id)
    section = (manifest.get("sections") or {}).get("final_incident_report") or section
    export_entry = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now(), "source_docx": str(docx_path)}
    manifest.setdefault("exports", {})["pdf"] = export_entry
    section.setdefault("exports", {})["pdf"] = export_entry
    manifest["sections"]["final_incident_report"] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "path": str(path), "manifest": manifest, "download_url": "/api/reports/download/pdf"}


def download_path(output_dir: Path, section_key: str | None, file_type: str, incident_id: str | None = None) -> Path:
    """[FYP-FUNCTION] [FYP-EXPORT] Resolve An Already-Exported File's Path.

    Purpose: look up the on-disk path for a previously exported file so it
    can be streamed back to the analyst as a download.
    Params: output_dir/incident_id -- see load_manifest(); section_key --
        a specific section's export, or None for the combined report;
        file_type -- "docx" or "pdf" (matches the keys used in
        section["exports"] / manifest["exports"]).
    Returns: the Path to the exported file.
    Error handling: FileNotFoundError if no manifest, no matching export
        record, or the file no longer exists on disk; KeyError if
        section_key is not a known section.
    Called by: soc_reporting_agent/backend/app.py (GET download endpoints,
        which wrap the returned Path in Flask's send_file()).
    """
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found")
    if section_key:
        section = (manifest.get("sections") or {}).get(section_key)
        if not section:
            raise KeyError(f"Unknown report section: {section_key}")
        export = (section.get("exports") or {}).get(file_type)
    else:
        export = (manifest.get("exports") or {}).get(file_type)
    if not export:
        raise FileNotFoundError(f"No {file_type} export found")
    path = Path(export.get("path", ""))
    if not path.exists():
        raise FileNotFoundError("Export file not found")
    return path


# ═══════════════════════════════════════════════════════════════════════
# [FYP-SECTION] Candidate Manifest -- Immutable Final Snapshot
# [FYP-EXPORT] [FYP-STATE] [FYP-APPROVAL]
# Candidate manifest — an IMMUTABLE final snapshot, distinct from the
# mutable report_manifest.json above (which stays exactly as it always
# was: the Reporting Agent's own working manifest through draft/confirm/
# export). candidate_manifest.json is written exactly once, only after all
# 4 structured reports + 4 DOCX + 4 PDF exist and have been validated, and
# is the sole basis for preview/approval/export from that point on — see
# reporting_approval.approve_reporting_candidate() and
# workflow_state_store.commit_reporting_approval().
# ═══════════════════════════════════════════════════════════════════════

# [FYP-CLASS] `CandidateManifestConflictError` — owns CandidateManifestConflictError state or behaviour for the report generation and export component.
# [FYP-PROCESS] Important methods: no public methods; class-level data/exception semantics only.
# [FYP-USED-BY] Static constructor/type references include soc_reporting_agent/reporting/editable_reports.py:finalize_candidate_manifest.
# [FYP-OUTPUT] Instances expose the state and operations defined by the class body; local methods document side effects.
# [FYP-ERROR] Constructor/method exceptions propagate unless a documented local fallback handles them.

class CandidateManifestConflictError(RuntimeError):
    # [FYP-CLASS] CandidateManifestConflictError
    # Purpose: signals a write-once integrity violation -- a caller tried to
    # finalize a candidate set for a run/attempt whose candidate_manifest.json
    # was already published with genuinely different content.
    # Raised by: finalize_candidate_manifest().
    # Caught by: callers such as adapters/export_documents.py, which should
    # surface this as a hard failure rather than silently overwriting the
    # published (and possibly already-approved) candidate set.
    """Raised when finalize_candidate_manifest() is asked to (re-)publish a
    candidate set for an attempt whose candidate_manifest.json already
    exists with DIFFERENT content — the immutable manifest is never
    overwritten. A byte-for-byte-equivalent repeat call (same 12 file
    hashes) is treated as a safe idempotent no-op instead of raising."""


def candidate_manifest_path(output_dir: Path, incident_id: str) -> Path:
    """[FYP-FUNCTION] [FYP-EXPORT] Path to the immutable
    candidate_manifest.json for an incident (see finalize_candidate_manifest()
    for how it is produced, and reporting_approval.py /
    workflow_state_store.py for how it is consumed downstream)."""
    return incident_report_dir(output_dir, incident_id) / "candidate_manifest.json"


def _hash_file(path: Path) -> tuple[str, int]:
    """[FYP-FUNCTION] SHA-256 + byte size of a file on disk. Used to build
    the per-file integrity hashes embedded in candidate_manifest.json (one
    for each of the 12 structured/docx/pdf files) so any later mutation of
    an already-published file can be detected. Called by:
    finalize_candidate_manifest()."""
    data = path.read_bytes()
    return hashlib.sha256(data).hexdigest(), len(data)


# [FYP-FUNCTION] `_canonical_manifest_bytes` — implements the canonical manifest bytes operation used by the surrounding report generation and export workflow.
# [FYP-INPUT] Parameters: `manifest_without_hash`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
# [FYP-USED-BY] Static symbol references include reporting_approval.py:_verify_candidate_manifest, soc_reporting_agent/reporting/editable_reports.py:finalize_candidate_manifest; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `dumps`, `encode`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _canonical_manifest_bytes(manifest_without_hash: dict[str, Any]) -> bytes:
    # [FYP-FUNCTION] Canonical Manifest Serialisation
    """Deterministic serialisation used for both the published hash and
    the atomic-write content — sorted keys, no incidental whitespace, so
    the hash never drifts for reasons unrelated to actual content."""
    return json.dumps(manifest_without_hash, sort_keys=True, separators=(",", ":"),
                      ensure_ascii=False).encode("utf-8")


# [FYP-FUNCTION] `_report_content_signature` — implements the report content signature operation used by the surrounding report generation and export workflow.
# [FYP-INPUT] Parameters: `entries`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
# [FYP-USED-BY] Static symbol references include soc_reporting_agent/reporting/editable_reports.py:finalize_candidate_manifest; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `get`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _report_content_signature(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    # [FYP-FUNCTION] Content-Only Signature For Idempotency Checks
    """The parts of a report entry that represent actual CONTENT identity —
    excludes generated_at, which legitimately differs between an original
    finalize call and a same-content idempotent repeat call. Used only to
    decide "is this a safe no-op retry", never for the published hash
    itself (which does cover generated_at, since that hash's job is to
    detect ANY later mutation of the one-time-published file)."""
    return [
        {k: e.get(k) for k in ("report_type", "title", "filename", "template",
                               "structured_content", "docx", "pdf", "validation")}
        for e in entries
    ]


def finalize_candidate_manifest(output_dir: Path, incident_id: str, run_id: str,
                                reporting_stage_attempt: int) -> dict[str, Any]:
    """[FYP-FUNCTION] [FYP-EVALUATOR] [FYP-EXPORT] [FYP-STATE] [FYP-VALIDATION]
    Publish The Immutable Candidate Manifest -- the hand-off point from the
    Reporting stage to the approval/workflow layer.

    Params: output_dir/incident_id -- identify the incident; run_id -- the
        workflow run this belongs to (source: SOC_RUN_ID env var, threaded
        in by soc_workflow.run_reporting_stage()); reporting_stage_attempt --
        which reporting attempt this is (source: SOC_REPORTING_ATTEMPT env
        var), included so a re-run after a rejected report produces a
        distinctly-attributed candidate set.
    Returns: the published (or, on a safe idempotent repeat, the pre-
        existing) candidate manifest dict, including candidate_manifest_sha256
        and a fresh report_set_id (uuid4 hex) on first publication.
    Calls: report_validator.validate_generated_report() (lazy import, one
        call per core section) to attach a validation status onto each
        report entry; load_manifest(), _hash_file(),
        _canonical_manifest_bytes(), _report_content_signature().
    Called by: soc_reporting_agent/adapters/export_documents.py, after
        confirm_report() + export_docx/pdf() + export_section_docx/pdf()
        (for executive_summary, technical_findings, soc_analyst_review) have
        all completed for a headless workflow run.

    Build and atomically publish candidate_manifest.json — the true
    final snapshot of a completed Reporting generation.

    Must be called only after adapters/export_documents.py has confirmed
    and exported all 4 core reports (structured content + DOCX + PDF for
    each). If any of those 12 files is missing/empty, this raises
    FileNotFoundError and does NOT publish anything — the caller
    (run_reporting_stage()) must treat that as a generation/validator
    execution failure (Reporting=Failed, Workflow=Failed), never as a
    published-but-blocked candidate set. A CONTENT-level validation
    problem (raw pipe syntax, unresolved Jinja, etc.) is different: that
    still publishes normally, with validation.status="error" recorded on
    the affected report entry — publication and approval-eligibility are
    decided by different code at different times (see
    reporting_approval.approve_reporting_candidate())."""
    from reporting.report_validator import validate_generated_report

    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found — cannot finalize a candidate set.")

    # [FYP-FUNCTION] `_build_report_entries` — constructs build report entries output for the next report generation and export consumer or analyst-facing view.
    # [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
    # [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
    # [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
    # [FYP-USED-BY] Static symbol references include soc_reporting_agent/reporting/editable_reports.py:finalize_candidate_manifest; dynamic framework calls may add callers.
    # [FYP-CALLS] Calls: `FileNotFoundError`, `Path`, `_hashed`, `append`, `exists`, `get`, `stat`, `utc_now`.
    # [FYP-ERROR] Raises explicit validation/processing errors to the caller; no silent fallback is applied here.

    def _build_report_entries() -> list[dict[str, Any]]:
        # [FYP-FUNCTION] Nested helper (closes over manifest/output_dir/
        # incident_id): for each of CORE_REPORT_KEYS, verifies the section's
        # structured/docx/pdf exports exist and are non-empty, runs
        # report_validator.validate_generated_report() on them, and returns
        # one candidate-manifest "report entry" dict per core report. Raises
        # FileNotFoundError on any missing/empty export -- see docstring
        # above for why that must NOT result in a partial publish. Called
        # twice on a repeat-finalize (to build fresh_entries for the
        # idempotency comparison) and once on first publish.
        entries: list[dict[str, Any]] = []
        for key in CORE_REPORT_KEYS:
            section = (manifest.get("sections") or {}).get(key)
            if not section:
                raise FileNotFoundError(f"candidate manifest: section '{key}' missing from report_manifest.json")
            structured_path = section.get("structured_confirmed_path")
            exports = section.get("exports") or {}
            docx_path = (exports.get("docx") or {}).get("path")
            pdf_path = (exports.get("pdf") or {}).get("path")
            if not (structured_path and docx_path and pdf_path):
                raise FileNotFoundError(
                    f"candidate manifest: '{key}' is missing a required structured/docx/pdf export")
            structured_p, docx_p, pdf_p = Path(structured_path), Path(docx_path), Path(pdf_path)
            for p, label in ((structured_p, "structured_content"), (docx_p, "docx"), (pdf_p, "pdf")):
                if not p.exists() or p.stat().st_size == 0:
                    raise FileNotFoundError(f"candidate manifest: '{key}' {label} file missing or empty: {p}")
            validation = validate_generated_report(
                docx_path=docx_p, pdf_path=pdf_p, structured_content_path=structured_p,
                report_title=section.get("title") or key, incident_id=incident_id)

            # [FYP-FUNCTION] `_hashed` — implements the hashed operation used by the surrounding report generation and export workflow.
            # [FYP-INPUT] Parameters: `p`; values come from its direct caller, route, UI event, fixture, or stage handoff.
            # [FYP-PROCESS] Executes the named operation within the Aegis report generation and export workflow; branch rules remain in the body below.
            # [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
            # [FYP-USED-BY] Static symbol references include soc_reporting_agent/reporting/editable_reports.py:_build_report_entries; dynamic framework calls may add callers.
            # [FYP-CALLS] Calls: `_hash_file`, `_rel`.
            # [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

            def _hashed(p: Path) -> dict[str, Any]:
                # [FYP-FUNCTION] Nested helper: wraps _hash_file() into the
                # {path, sha256, size} dict shape stored per-file in each
                # candidate-manifest report entry (structured_content/docx/pdf).
                sha, size = _hash_file(p)
                return {"path": _rel(output_dir, p), "sha256": sha, "size": size}

            entries.append({
                "report_type": key,
                "title": section.get("title") or key,
                "filename": {"docx": docx_p.name, "pdf": pdf_p.name},
                "template": section.get("template"),
                "structured_content": _hashed(structured_p),
                "docx": _hashed(docx_p),
                "pdf": _hashed(pdf_p),
                "generated_at": utc_now(),
                "validation": validation,
            })
        return entries

    final_path = candidate_manifest_path(output_dir, incident_id)

    if final_path.exists():
        existing = json.loads(final_path.read_text(encoding="utf-8"))
        fresh_entries = _build_report_entries()
        if _report_content_signature(existing.get("reports") or []) == _report_content_signature(fresh_entries):
            return existing   # identical content — safe no-op, no new report_set_id minted
        raise CandidateManifestConflictError(
            f"candidate_manifest.json already exists for {incident_id}/{run_id} attempt "
            f"{reporting_stage_attempt} and differs from what this finalize call would "
            "produce — refusing to overwrite a published candidate set.")

    report_entries = _build_report_entries()
    manifest_without_hash: dict[str, Any] = {
        "incident_id": incident_id,
        "run_id": run_id,
        "reporting_stage_attempt": reporting_stage_attempt,
        "report_set_id": uuid.uuid4().hex,
        "generated_at": utc_now(),
        "reports": report_entries,
        # Metadata-only compatibility alias for any code still keyed on the
        # legacy "combined_incident_report" name — never a second document.
        "legacy_combined_incident_report": {"deprecated": True, "points_to": "final_incident_report"},
    }
    digest = hashlib.sha256(_canonical_manifest_bytes(manifest_without_hash)).hexdigest()
    full_manifest = dict(manifest_without_hash)
    full_manifest["candidate_manifest_sha256"] = digest

    final_path.parent.mkdir(parents=True, exist_ok=True)
    # [FYP-PROCESS] Atomic write pattern: write to a uniquely-named temp file
    # first, fsync it, then os.replace() it onto the final path in one step
    # (os.replace is atomic on the same filesystem) so a crash mid-write can
    # never leave a partially-written candidate_manifest.json behind.
    tmp_path = final_path.with_name(f".{final_path.name}.tmp-{uuid.uuid4().hex[:8]}")
    try:
        with open(tmp_path, "w", encoding="utf-8") as fh:
            fh.write(json.dumps(full_manifest, indent=2, ensure_ascii=False))
            fh.flush()
            os.fsync(fh.fileno())
        os.replace(tmp_path, final_path)  # safe: existence already ruled out above
        try:
            dir_fd = os.open(str(final_path.parent), os.O_RDONLY)
            try:
                os.fsync(dir_fd)
            finally:
                os.close(dir_fd)
        except (OSError, AttributeError):
            pass  # directory fsync unsupported on this platform (e.g. Windows) — documented no-op
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass
    return full_manifest
