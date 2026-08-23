# =============================================================================
# [FYP-FILE] FILE OVERVIEW
# Important dependencies: __future__, case_view, hashlib, json, pathlib, pytest, reporting_approval, soc_workflow.
# =============================================================================
# File: tests/test_reporting_stage.py
# Purpose: This module implements test and validation behaviour for test reporting stage.
# Main functionality: _isolated_db, _isolated_artifact_root, _run_awaiting_reporting_approval, _write_minimal_docx, _write_minimal_pdf, _build_candidate_set.
# Inputs: Function parameters, configured environment values, persisted artifacts,
#   or framework callbacks identified by the documented entry points below.
# Outputs: Return values and documented file, database, workflow-state, or UI
#   side effects consumed by the next stage or analyst-facing component.
# Workflow position: Part of the Aegis test and validation component.
# Called by: Direct callers are identified on each function/class annotation;
#   framework and command-line entry points are marked explicitly.
# Calls / important dependencies: __future__, case_view, hashlib, json, pathlib, pytest, reporting_approval, soc_workflow.
# Important side effects: See [FYP-OUTPUT], [FYP-STATE], [FYP-DATABASE],
#   [FYP-EXPORT], and [FYP-UI] annotations on the affected operations.
# Error and fallback behaviour: Local try/except and fallback paths are marked
#   per function; otherwise failures propagate to the documented caller.
# Key evaluator search terms: _isolated_db, _isolated_artifact_root, _run_awaiting_reporting_approval, _write_minimal_docx, _write_minimal_pdf, _build_candidate_set, [FYP-FUNCTION], [FYP-EVALUATOR].
# =============================================================================

from __future__ import annotations

import hashlib
import json
from pathlib import Path

import pytest

from workflow import state_store as wss
from workflow import engine as sw
from agents.reporting import reporting_approval as ra
import backend.services.case_view_service as cv


# =============================================================================
# [FYP-SECTION] TEST SETUP, FIXTURES, AND ASSERTIONS
# =============================================================================

# [FYP-FUNCTION] `_isolated_db` — implements the isolated db operation used by the surrounding test and validation workflow.
# [FYP-INPUT] Parameters: `tmp_path`, `monkeypatch`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `db_init`, `setattr`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

@pytest.fixture(autouse=True)
def _isolated_db(tmp_path, monkeypatch):
    monkeypatch.setattr(wss, "DB_FILE", tmp_path / "reporting.db")
    wss.db_init()


# [FYP-FUNCTION] `_isolated_artifact_root` — implements the isolated artifact root operation used by the surrounding test and validation workflow.
# [FYP-INPUT] Parameters: `tmp_path`, `monkeypatch`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `setattr`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

@pytest.fixture(autouse=True)
def _isolated_artifact_root(tmp_path, monkeypatch):
    """Every run-scoped path (reporting_attempt_dir / _artifact_dir) is
    computed from soc_workflow._TRUSTED_OUTPUT_ROOT — isolate it to a tmp
    dir so tests never touch the real soc_reporting_agent/outputs tree."""
    trusted_root = tmp_path / "trusted_outputs"
    monkeypatch.setattr(sw, "_TRUSTED_OUTPUT_ROOT", trusted_root)
    monkeypatch.setattr(ra, "_TRUSTED_OUTPUT_ROOT", trusted_root)
    return trusted_root


# [FYP-FUNCTION] `_run_awaiting_reporting_approval` — orchestrates the run awaiting reporting approval entry point and its ordered test and validation operations.
# [FYP-INPUT] Parameters: `incident_id`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
# [FYP-USED-BY] Static symbol references include tests/test_reporting_stage.py:test_approve_reporting_candidate_end_to_end, tests/test_reporting_stage.py:test_approve_reporting_candidate_fails_on_identity_mismatch, tests/test_reporting_stage.py:test_approve_reporting_candidate_fails_when_docx_tampered_after_generation; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `_guarded_update`, `start_run`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _run_awaiting_reporting_approval(incident_id: str = "INC-1") -> str:
    run_id = wss.start_run(incident_id)
    wss._guarded_update(incident_id, run_id, {
        "triage_status": "Approved",
        "threat_intel_status": "Complete",
        "investigation_status": "Approved",
        "reporting_status": "Awaiting Approval",
        "workflow_status": "Awaiting Approval",
        "approval_stage": "reporting",
    })
    return run_id


# [FYP-FUNCTION] `_write_minimal_docx` — persists or updates write minimal docx state used by the surrounding test and validation workflow.
# [FYP-INPUT] Parameters: `path`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] Static symbol references include tests/test_reporting_stage.py:_build_candidate_set; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `Document`, `add_paragraph`, `mkdir`, `save`, `str`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _write_minimal_docx(path: Path) -> None:
    from docx import Document
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("Test report content.")
    doc.save(str(path))


# [FYP-FUNCTION] `_write_minimal_pdf` — persists or updates write minimal pdf state used by the surrounding test and validation workflow.
# [FYP-INPUT] Parameters: `path`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] Static symbol references include tests/test_reporting_stage.py:_build_candidate_set; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `Paragraph`, `SimpleDocTemplate`, `build`, `getSampleStyleSheet`, `mkdir`, `str`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _write_minimal_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    SimpleDocTemplate(str(path), pagesize=A4).build(
        [Paragraph("Test report content.", styles["BodyText"])])


_ALL_CORE_REPORT_TYPES = ("executive_summary", "technical_findings",
                         "soc_analyst_review", "final_incident_report")


# [FYP-FUNCTION] `_build_candidate_set` — constructs build candidate set output for the next test and validation consumer or analyst-facing view.
# [FYP-INPUT] Parameters: `incident_id`, `run_id`, `attempt`, `trusted_root`, `report_types`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns the explicit value(s) from its decision paths for the documented caller to consume.
# [FYP-USED-BY] Static symbol references include tests/test_reporting_stage.py:test_approve_reporting_candidate_end_to_end, tests/test_reporting_stage.py:test_approve_reporting_candidate_fails_on_identity_mismatch, tests/test_reporting_stage.py:test_approve_reporting_candidate_fails_when_docx_tampered_after_generation; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `Path`, `_write_minimal_docx`, `_write_minimal_pdf`, `candidate_manifest_path`, `confirmed_dir`, `dumps`, `exports_dir`, `finalize_candidate_manifest`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _build_candidate_set(incident_id: str, run_id: str, attempt: int, trusted_root: Path,
                         *, report_types=_ALL_CORE_REPORT_TYPES) -> tuple[Path, dict]:
    """Builds a real, on-disk candidate set (structured content + real
    DOCX + real PDF for each report type) and finalizes it via
    editable_reports.finalize_candidate_manifest()'s own hashing/atomic-
    write logic reused here directly for test setup speed — instead we
    hand-construct a report_manifest.json shaped exactly like
    editable_reports.build_report_manifest() would produce, then call the
    real finalize_candidate_manifest()."""
    import sys
    rep_dir = str(Path(__file__).resolve().parent.parent / "agents" / "reporting")
    if rep_dir not in sys.path:
        sys.path.insert(0, rep_dir)
    from reporting import editable_reports as er

    attempt_dir = sw.reporting_attempt_dir(incident_id, run_id, attempt)
    output_dir = attempt_dir / "outputs"
    incident_report_dir = er.incident_report_dir(output_dir, incident_id)
    confirmed_dir = er.confirmed_dir(output_dir, incident_id)
    exports_dir = er.exports_dir(output_dir, incident_id)
    confirmed_dir.mkdir(parents=True, exist_ok=True)
    exports_dir.mkdir(parents=True, exist_ok=True)

    sections = {}
    for key in report_types:
        structured_path = confirmed_dir / f"{key}.json"
        structured_path.write_text(json.dumps(
            [{"type": "heading", "level": 1, "text": key},
             {"type": "paragraph", "text": "Body text."}]), encoding="utf-8")
        docx_path = exports_dir / f"{key}.docx"
        pdf_path = exports_dir / f"{key}.pdf"
        _write_minimal_docx(docx_path)
        _write_minimal_pdf(pdf_path)
        sections[key] = {
            "key": key, "title": key.replace("_", " ").title(), "template": f"{key}.md.j2",
            "structured_confirmed_path": str(structured_path),
            "exports": {"docx": {"path": str(docx_path)}, "pdf": {"path": str(pdf_path)}},
        }
    manifest = {"schema_version": "editable-report-manifest-v2", "incident_id": incident_id,
               "sections": sections, "section_order": list(report_types)}
    (incident_report_dir / "report_manifest.json").write_text(
        json.dumps(manifest), encoding="utf-8")

    candidate = er.finalize_candidate_manifest(output_dir, incident_id, run_id, attempt)
    return er.candidate_manifest_path(output_dir, incident_id), candidate


# [FYP-FUNCTION] `_approve_via_state` — applies the human-in-the-loop approve via state decision and returns or persists the resulting workflow state.
# [FYP-INPUT] Parameters: `incident_id`, `run_id`, `attempt`, `candidate_manifest_path`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] Static symbol references include tests/test_reporting_stage.py:test_approve_reporting_candidate_end_to_end, tests/test_reporting_stage.py:test_approve_reporting_candidate_fails_on_identity_mismatch, tests/test_reporting_stage.py:test_approve_reporting_candidate_fails_when_docx_tampered_after_generation; dynamic framework calls may add callers.
# [FYP-CALLS] Calls: `_guarded_update`, `dumps`, `str`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def _approve_via_state(incident_id: str, run_id: str, attempt: int,
                       candidate_manifest_path: Path) -> None:
    """Stamps reporting_result_json the way run_reporting_stage() would,
    then transitions to Awaiting Approval so approve_reporting_candidate()
    has something real to validate."""
    result = {"status": "completed",
             "document_exports": {"candidate_manifest_path": str(candidate_manifest_path)}}
    wss._guarded_update(incident_id, run_id, {
        "reporting_attempt": attempt,
        "reporting_result_json": json.dumps(result),
        "reporting_status": "Awaiting Approval",
        "workflow_status": "Awaiting Approval",
        "approval_stage": "reporting",
    })


# [FYP-FUNCTION] `test_run_scoped_handoff_includes_threat_intel_in_reporting_inputs` — verifies run scoped handoff includes threat intel in reporting inputs behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `tmp_path`, `monkeypatch`, `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `Path`, `dumps`, `handoff_to_reporting`, `loads`, `mkdir`, `read_text`, `reporting_attempt_dir`, `setattr`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_run_scoped_handoff_includes_threat_intel_in_reporting_inputs(
        tmp_path, monkeypatch, _isolated_artifact_root):
    incident_id = "INC-HANDOFF-1"
    run_id = "run-handoff-1"
    attempt = 2
    monkeypatch.setattr(sw, "REP_DIR", tmp_path)

    parsing_dir = tmp_path / "outputs" / incident_id / run_id / "parsing"
    parsing_dir.mkdir(parents=True)
    (parsing_dir / "processed_alert.json").write_text(json.dumps({
        "incident_id": incident_id,
        "alert_id": "ALERT-HANDOFF-1",
    }), encoding="utf-8")

    threat_intel = {
        "incident_id": incident_id,
        "run_id": run_id,
        "stage": "threat_intelligence",
        "status": "completed",
        "threat_intelligence": {"iocs": {}},
    }
    sw.handoff_to_reporting(
        {
            "metakeys_payload": {"incident_id": incident_id},
            "ticket": {"incident_id": incident_id, "unc": "#HANDOFF-1"},
        },
        {"id": incident_id, "title": "Handoff regression test"},
        {"incident_id": incident_id, "status": "completed", "summary": "Done"},
        threat_intel_result=threat_intel,
        incident_id=incident_id,
        run_id=run_id,
        reporting_stage_attempt=attempt,
    )

    attempt_dir = sw.reporting_attempt_dir(incident_id, run_id, attempt)
    input_path = attempt_dir / "inputs" / "threat_intel_result.json"
    output_path = attempt_dir / "outputs" / "threat_intel_result.json"
    assert json.loads(input_path.read_text(encoding="utf-8")) == threat_intel
    assert json.loads(output_path.read_text(encoding="utf-8")) == threat_intel

    manifest = json.loads(
        (attempt_dir / "inputs" / "handoff_manifest.json").read_text(encoding="utf-8"))
    assert str(Path("inputs") / "threat_intel_result.json") in manifest["files"]
    assert str(Path("outputs") / "threat_intel_result.json") in manifest["files"]


# ══════════════════════════════════════════════════════════════════════
# workflow_state_store.py — durable approval binding
# ══════════════════════════════════════════════════════════════════════

# [FYP-FUNCTION] `test_commit_reporting_approval_is_only_approve_reporting_workflow_status_setter` — verifies commit reporting approval is only approve reporting workflow status setter behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_run_awaiting_reporting_approval`, `commit_reporting_approval`, `get_state`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_commit_reporting_approval_is_only_approve_reporting_workflow_status_setter():
    run_id = _run_awaiting_reporting_approval()
    state = wss.get_state("INC-1")
    result = wss.commit_reporting_approval(
        "INC-1", run_id, expected_reporting_attempt=state["reporting_attempt"],
        expected_reporting_result_json=state["reporting_result_json"],
        metadata={"report_set_id": "rs-1"}, approved_by="analyst")
    assert result["incident_id"] == "INC-1"
    state = wss.get_state("INC-1")
    assert state["reporting_status"] == "Approved"
    assert state["workflow_status"] == "Complete"


# [FYP-FUNCTION] `test_commit_reporting_approval_fails_when_reporting_result_json_changed` — verifies commit reporting approval fails when reporting result json changed behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_guarded_update`, `_run_awaiting_reporting_approval`, `commit_reporting_approval`, `dumps`, `get_state`, `raises`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_commit_reporting_approval_fails_when_reporting_result_json_changed():
    run_id = _run_awaiting_reporting_approval()
    state = wss.get_state("INC-1")
    captured_json = state["reporting_result_json"]
    # Simulate a concurrent rerun changing reporting_result_json between
    # the caller's validation pass and this call.
    wss._guarded_update("INC-1", run_id, {"reporting_result_json": json.dumps({"x": 1})})
    with pytest.raises(wss.ApprovalConflictError):
        wss.commit_reporting_approval(
            "INC-1", run_id, expected_reporting_attempt=state["reporting_attempt"],
            expected_reporting_result_json=captured_json,
            metadata={}, approved_by="analyst")


# [FYP-FUNCTION] `test_commit_reporting_approval_fails_when_attempt_changed` — verifies commit reporting approval fails when attempt changed behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_run_awaiting_reporting_approval`, `commit_reporting_approval`, `get_state`, `raises`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_commit_reporting_approval_fails_when_attempt_changed():
    run_id = _run_awaiting_reporting_approval()
    state = wss.get_state("INC-1")
    with pytest.raises(wss.ApprovalConflictError):
        wss.commit_reporting_approval(
            "INC-1", run_id, expected_reporting_attempt=state["reporting_attempt"] + 1,
            expected_reporting_result_json=state["reporting_result_json"],
            metadata={}, approved_by="analyst")


# [FYP-FUNCTION] `test_get_approved_reporting_sets_reads_durable_metadata_survives_rerun_clear` — verifies get approved reporting sets reads durable metadata survives rerun clear behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_run_awaiting_reporting_approval`, `commit_reporting_approval`, `get_latest_approved_reporting_set`, `get_state`, `rerun_stage`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_get_approved_reporting_sets_reads_durable_metadata_survives_rerun_clear():
    run_id = _run_awaiting_reporting_approval()
    state = wss.get_state("INC-1")
    wss.commit_reporting_approval(
        "INC-1", run_id, expected_reporting_attempt=state["reporting_attempt"],
        expected_reporting_result_json=state["reporting_result_json"],
        metadata={"report_set_id": "rs-durable", "candidate_manifest_path": "x",
                 "candidate_manifest_sha256": "abc", "reporting_stage_attempt": 1,
                 "validation_status": "valid", "warning_count": 0},
        approved_by="analyst")

    # rerun clears the working reporting_result_json — the durable
    # approval record must not be affected.
    wss.rerun_stage("INC-1", run_id, "reporting")
    state = wss.get_state("INC-1")
    assert state["reporting_result_json"] is None

    latest = wss.get_latest_approved_reporting_set("INC-1", run_id)
    assert latest is not None
    assert latest["report_set_id"] == "rs-durable"
    assert latest["candidate_manifest_sha256"] == "abc"


# [FYP-FUNCTION] `test_rejected_reporting_attempt_can_be_rerun` — verifies rejected reporting attempt can be rerun behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_run_awaiting_reporting_approval`, `get_state`, `reject_reporting`, `rerun_stage`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_rejected_reporting_attempt_can_be_rerun():
    run_id = _run_awaiting_reporting_approval()
    wss.reject_reporting("INC-1", run_id, rejected_by="analyst", reason="needs fixes")
    assert wss.get_state("INC-1")["reporting_status"] == "Rejected"
    wss.rerun_stage("INC-1", run_id, "reporting")
    state = wss.get_state("INC-1")
    assert state["reporting_status"] == "Processing"
    assert state["reporting_attempt"] == 2


# [FYP-FUNCTION] `test_complete_stage_expected_stage_attempt_rejects_stale_worker` — verifies complete stage expected stage attempt rejects stale worker behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_guarded_update`, `claim_stage`, `complete_stage`, `start_run`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_complete_stage_expected_stage_attempt_rejects_stale_worker():
    run_id = wss.start_run("INC-1")
    wss._guarded_update("INC-1", run_id, {"reporting_status": "Processing"})
    worker_id, attempt = wss.claim_stage(
        "INC-1", run_id, stage="reporting", status_column="reporting_status",
        expect_status="Processing")
    # Simulate the attempt having moved on before this (late) worker saves.
    wss._guarded_update("INC-1", run_id, {"reporting_attempt": attempt + 1})
    ok = wss.complete_stage(
        "INC-1", run_id, worker_id, stage="reporting", result_column="reporting_result_json",
        result={"status": "completed"},
        status_updates={"reporting_status": "Awaiting Approval"},
        expected_stage_attempt=attempt)
    assert ok is False


# ══════════════════════════════════════════════════════════════════════
# reporting_approval.py — layering + validation
# ══════════════════════════════════════════════════════════════════════

# [FYP-FUNCTION] `test_workflow_state_store_has_no_report_specific_validation_imports` — verifies workflow state store has no report specific validation imports behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: no explicit parameters; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `getsource`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_workflow_state_store_has_no_report_specific_validation_imports():
    import inspect
    src = inspect.getsource(wss)
    for banned in ("import docx", "import pypdf", "from docx", "hashlib.sha256(Path"):
        assert banned not in src, f"workflow_state_store.py must not import/do report validation: {banned}"


# [FYP-FUNCTION] `test_approve_reporting_candidate_end_to_end` — verifies approve reporting candidate end to end behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `tmp_path`, `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_approve_via_state`, `_build_candidate_set`, `_run_awaiting_reporting_approval`, `approve_reporting_candidate`, `get_latest_approved_reporting_set`, `get_state`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_approve_reporting_candidate_end_to_end(tmp_path, _isolated_artifact_root):
    run_id = _run_awaiting_reporting_approval()
    manifest_path, candidate = _build_candidate_set(
        "INC-1", run_id, 1, _isolated_artifact_root)
    _approve_via_state("INC-1", run_id, 1, manifest_path)

    result = ra.approve_reporting_candidate("INC-1", run_id, analyst="analyst")
    assert result["incident_id"] == "INC-1"
    state = wss.get_state("INC-1")
    assert state["reporting_status"] == "Approved"
    assert state["workflow_status"] == "Complete"

    latest = wss.get_latest_approved_reporting_set("INC-1", run_id)
    assert latest["report_set_id"] == candidate["report_set_id"]


# [FYP-FUNCTION] `test_approve_reporting_candidate_fails_when_docx_tampered_after_generation` — verifies approve reporting candidate fails when docx tampered after generation behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `tmp_path`, `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_approve_via_state`, `_build_candidate_set`, `_run_awaiting_reporting_approval`, `approve_reporting_candidate`, `get_state`, `raises`, `read_bytes`, `reporting_attempt_dir`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_approve_reporting_candidate_fails_when_docx_tampered_after_generation(
        tmp_path, _isolated_artifact_root):
    run_id = _run_awaiting_reporting_approval()
    manifest_path, candidate = _build_candidate_set(
        "INC-1", run_id, 1, _isolated_artifact_root)
    _approve_via_state("INC-1", run_id, 1, manifest_path)

    attempt_dir = sw.reporting_attempt_dir("INC-1", run_id, 1)
    docx_path = attempt_dir / candidate["reports"][0]["docx"]["path"]
    docx_path.write_bytes(docx_path.read_bytes() + b"tampered")

    with pytest.raises(ra.ReportValidationError):
        ra.approve_reporting_candidate("INC-1", run_id, analyst="analyst")
    assert wss.get_state("INC-1")["reporting_status"] == "Awaiting Approval"


# [FYP-FUNCTION] `test_approve_reporting_candidate_fails_on_identity_mismatch` — verifies approve reporting candidate fails on identity mismatch behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `tmp_path`, `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `_approve_via_state`, `_build_candidate_set`, `_guarded_update`, `_run_awaiting_reporting_approval`, `approve_reporting_candidate`, `raises`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_approve_reporting_candidate_fails_on_identity_mismatch(tmp_path, _isolated_artifact_root):
    run_id = _run_awaiting_reporting_approval()
    manifest_path, _ = _build_candidate_set("INC-1", run_id, 1, _isolated_artifact_root)
    _approve_via_state("INC-1", run_id, 1, manifest_path)
    # Reporting attempt moved on without the manifest being re-finalized.
    wss._guarded_update("INC-1", run_id, {"reporting_attempt": 2})
    with pytest.raises(ra.ReportValidationError):
        ra.approve_reporting_candidate("INC-1", run_id, analyst="analyst")


# ══════════════════════════════════════════════════════════════════════
# editable_reports.finalize_candidate_manifest — immutability
# ══════════════════════════════════════════════════════════════════════

# [FYP-FUNCTION] `test_finalize_candidate_manifest_refuses_to_overwrite_differing_content` — verifies finalize candidate manifest refuses to overwrite differing content behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `Path`, `_build_candidate_set`, `finalize_candidate_manifest`, `insert`, `raises`, `read_bytes`, `reporting_attempt_dir`, `resolve`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_finalize_candidate_manifest_refuses_to_overwrite_differing_content(
        _isolated_artifact_root):
    manifest_path, candidate = _build_candidate_set("INC-2", "run-a", 1, _isolated_artifact_root)

    import sys
    rep_dir = str(Path(__file__).resolve().parent.parent / "agents" / "reporting")
    if rep_dir not in sys.path:
        sys.path.insert(0, rep_dir)
    from reporting import editable_reports as er

    attempt_dir = sw.reporting_attempt_dir("INC-2", "run-a", 1)
    output_dir = attempt_dir / "outputs"
    # Mutate one exported DOCX so a re-finalize would produce different hashes.
    docx_path = attempt_dir / candidate["reports"][0]["docx"]["path"]
    docx_path.write_bytes(docx_path.read_bytes() + b"changed")

    with pytest.raises(er.CandidateManifestConflictError):
        er.finalize_candidate_manifest(output_dir, "INC-2", "run-a", 1)


# [FYP-FUNCTION] `test_finalize_candidate_manifest_idempotent_on_identical_repeat_call` — verifies finalize candidate manifest idempotent on identical repeat call behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `Path`, `_build_candidate_set`, `finalize_candidate_manifest`, `insert`, `reporting_attempt_dir`, `resolve`, `str`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_finalize_candidate_manifest_idempotent_on_identical_repeat_call(_isolated_artifact_root):
    manifest_path, candidate = _build_candidate_set("INC-3", "run-a", 1, _isolated_artifact_root)

    import sys
    rep_dir = str(Path(__file__).resolve().parent.parent / "agents" / "reporting")
    if rep_dir not in sys.path:
        sys.path.insert(0, rep_dir)
    from reporting import editable_reports as er

    attempt_dir = sw.reporting_attempt_dir("INC-3", "run-a", 1)
    output_dir = attempt_dir / "outputs"
    repeat = er.finalize_candidate_manifest(output_dir, "INC-3", "run-a", 1)
    assert repeat["report_set_id"] == candidate["report_set_id"]


# [FYP-FUNCTION] `test_final_incident_report_export_populates_section_exports` — verifies final incident report export populates section exports behaviour and protects the related test and validation code path.
# [FYP-INPUT] Parameters: `_isolated_artifact_root`; values come from its direct caller, route, UI event, fixture, or stage handoff.
# [FYP-PROCESS] Executes the named operation within the Aegis test and validation workflow; branch rules remain in the body below.
# [FYP-OUTPUT] Returns `None` implicitly or explicitly; its observable result is the documented side effect or assertion.
# [FYP-USED-BY] No direct caller confidently identified; this may be an entry point, callback, or test helper.
# [FYP-CALLS] Calls: `Path`, `_write`, `build_report_manifest`, `confirm_report`, `drafts_dir`, `export_docx`, `insert`, `load_manifest`.
# [FYP-ERROR] Does not define a local fallback; unexpected failures propagate to the caller/framework error boundary.

def test_final_incident_report_export_populates_section_exports(_isolated_artifact_root):
    """Final Incident Report's DOCX/PDF export must be mirrored into
    sections.final_incident_report.exports (previously left {}), and use
    the renamed final_incident_report.* filename."""
    import sys
    rep_dir = str(Path(__file__).resolve().parent.parent / "agents" / "reporting")
    if rep_dir not in sys.path:
        sys.path.insert(0, rep_dir)
    from reporting import editable_reports as er

    output_dir = _isolated_artifact_root / "final_report_test"
    incident_id = "INC-4"
    for key in er.CORE_REPORT_KEYS:
        cfg = er.REPORT_SECTION_CONFIG[key]
        er._write(er.drafts_dir(output_dir, incident_id) / cfg["filename"], f"{key} content")
        er.save_blocks(er.drafts_dir(output_dir, incident_id) / f"{key}.json",
                       [{"type": "paragraph", "text": f"{key} body"}])
    er.build_report_manifest(output_dir, incident_id, {
        key: str(er.drafts_dir(output_dir, incident_id) / er.REPORT_SECTION_CONFIG[key]["filename"])
        for key in er.CORE_REPORT_KEYS
    })
    er.confirm_report(output_dir, analyst="tester", incident_id=incident_id)
    result = er.export_docx(output_dir, incident_id=incident_id)
    assert Path(result["path"]).name == "final_incident_report.docx"
    manifest = er.load_manifest(output_dir, incident_id)
    assert manifest["sections"]["final_incident_report"]["exports"]["docx"]["path"] == result["path"]
